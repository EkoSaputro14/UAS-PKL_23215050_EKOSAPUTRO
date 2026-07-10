import shutil
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Copy template
src = r'C:\Users\SMANSA\AppData\Local\hermes\cache\documents\doc_6c6d74787cac_Laporan_PKL_MuhammadAffif_Fixed.docx'
dst = r'C:\Users\SMANSA\mimotes\docs\pkl-report\LAPORAN_PKL_Eko_Saputro_23215050.docx'
shutil.copy2(src, dst)
print(f'Template copied to: {dst}')

# Open and modify
doc = Document(dst)

# Find and replace content
def find_para_index(doc, search_text):
    for i, p in enumerate(doc.paragraphs):
        if search_text in p.text.strip():
            return i
    return None

# Helper: replace text in paragraph keeping style
def set_para_text(para, text):
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = text
    else:
        r = para.add_run(text)

# Helper: clear all paragraphs between start and end index
def clear_range(doc, start_idx, end_idx):
    for i in range(start_idx, end_idx + 1):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            for run in p.runs:
                run.text = ''

DIAG_DIR = r'C:\Users\SMANSA\mimotes\docs\pkl-report\diagrams'

# ===== REPLACE COVER PAGE =====
# Find cover elements and replace
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'MUHAMMAD AFFIF' in text:
        set_para_text(p, 'EKO SAPUTRO')
    elif '24225046' in text and 'NIM' in text:
        set_para_text(p, 'NIM: 23215050')
    elif 'Implementasi Keamanan' in text:
        # Replace title lines
        if 'Implementasi Keamanan Aplikasi dan Manajemen State pada Sistem' in text:
            set_para_text(p, 'RANCANG BANGUN SISTEM CHATBOT AI')
        elif 'Informasi Manajemen SDM Berbasis Web Menggunakan Vue.js 3 dan' in text:
            set_para_text(p, 'BERBASIS PENGETAHUAN DENGAN')
        elif 'Pinia di PT Agsya Karya Manca' in text:
            set_para_text(p, 'RETRIEVAL-AUGMENTED GENERATION')
        elif text == 'Pinia di PT Agsya Karya Manca':
            pass  # skip extra lines

# Replace Lembar Persetujuan
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'Muhammad Affif' in text and 'Nama' in text:
        set_para_text(p, 'Nama\t\t: Eko Saputro')
    elif '24225046' in text and 'NIM' in text:
        set_para_text(p, 'NIM\t\t: 23215050')
    elif 'Syefudin' in text or 'Zaenul' in text:
        if 'Syefudin' in text:
            set_para_text(p, 'Zaenul Arif, M.Kom.')
    elif 'NIDN: 0620078901' in text:
        set_para_text(p, 'NIDN: ...........')

# Replace Lembar Pengesahan
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'Judul' in text and 'Implementasi' in text:
        set_para_text(p, 'Judul\t\t: Rancang Bangun Sistem Chatbot AI Berbasis Pengetahuan dengan RAG dan Pipeline CRM')

# Replace Kata Pengantar - keep structure, update names
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'Muhammad Affif' in text and 'Penulis' not in text:
        set_para_text(p, text.replace('Muhammad Affif', 'Eko Saputro'))
    elif 'Agsya Karya Manca' in text:
        set_para_text(p, text.replace('Agsya Karya Manca', 'Bank Mandiri KCP Tegal Sudirman'))
    elif 'Agsya' in text:
        set_para_text(p, text.replace('Agsya', 'Bank Mandiri'))

# Replace BAB headings
bab_replacements = {
    'BAB III TINJAUAN PUSTAKA': 'BAB III METODE PELAKSANAAN PKL',
    'BAB IV PELAKSANAAN PRAKTIK KERJA LAPANGAN': 'BAB IV HASIL YANG DICAPAI',
    'BAB V HASIL YANG DICAPAI PKL': 'BAB V PENUTUP',
    'BAB VI PENUTUP': '',  # Remove if extra
}

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for old, new in bab_replacements.items():
        if text == old:
            if new:
                set_para_text(p, new)
            else:
                # Clear the paragraph
                for run in p.runs:
                    run.text = ''

# Replace section headings in BAB I
section_replacements = {
    '1.2 Rumusan Masalah': '1.2 Tujuan PKL',
    '1.3 Tujuan PKL': '1.3 Manfaat PKL',
    '1.4 Manfaat PKL': '',
}

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for old, new in section_replacements.items():
        if text == old and new:
            set_para_text(p, new)

# Replace BAB II sections
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'PT Agsya Karya Manca' in text:
        set_para_text(p, text.replace('PT Agsya Karya Manca', 'Bank Mandiri KCP Tegal Sudirman'))
    elif 'Agsya' in text:
        set_para_text(p, text.replace('Agsya', 'Bank Mandiri'))

# Replace BAB III content headings (Tinjauan Pustaka -> Metode PKL)
metode_headings = {
    '3.1 Landasan Teori': '3.1 Tugas Umum',
    '3.1.1 Keamanan Aplikasi Web': '3.2 Tugas Khusus',
    '3.1.2 OAuth 2.0': '3.2.1 Analisis Kebutuhan Sistem',
    '3.1.3 JSON Web Token (JWT)': '3.2.2 Perancangan Arsitektur Sistem',
    '3.1.4 Role-Based Access Control (RBAC)': '3.3 Analisis Permasalahan dan Solusi',
}

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for old, new in metode_headings.items():
        if text == old:
            set_para_text(p, new)

# Replace BAB IV content headings (Pelaksanaan -> Hasil)
hasil_headings = {
    '4.1 Tugas Umum': '4.1 Gambaran Umum Sistem',
    '4.2 Tugas Khusus': '4.2 Analisis dan Perancangan Sistem',
    '4.3 Analisis Permasalahan dan Solusi': '4.3 Implementasi Sistem',
}

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for old, new in hasil_headings.items():
        if text == old:
            set_para_text(p, new)

# Replace BAB V content headings (Hasil -> Penutup)
penutup_headings = {
    '5.1 Analisis Kebutuhan Sistem': '5.1 Kesimpulan',
    '5.2 Perancangan Sistem': '5.2 Saran',
}

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for old, new in penutup_headings.items():
        if text == old:
            set_para_text(p, new)

# Replace company names throughout
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if 'Agsya' in run.text:
            run.text = run.text.replace('Agsya', 'Bank Mandiri')
        if 'PT Agsya Karya Manca' in run.text:
            run.text = run.text.replace('PT Agsya Karya Manca', 'Bank Mandiri KCP Tegal Sudirman')
        if 'Muhammad Affif' in run.text:
            run.text = run.text.replace('Muhammad Affif', 'Eko Saputro')
        if '24225046' in run.text:
            run.text = run.text.replace('24225046', '23215050')
        if 'Syefudin' in run.text:
            run.text = run.text.replace('Syefudin', 'Zaenul Arif')
        if 'Vue.js 3' in run.text:
            run.text = run.text.replace('Vue.js 3', 'Retrieval-Augmented Generation')
        if 'Pinia' in run.text and 'Pinia di' not in run.text:
            pass  # Keep Pinia if it's in body text about the template

# Replace table content
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if 'Muhammad Affif' in run.text:
                        run.text = run.text.replace('Muhammad Affif', 'Eko Saputro')
                    if '24225046' in run.text:
                        run.text = run.text.replace('24225046', '23215050')
                    if 'Agsya' in run.text:
                        run.text = run.text.replace('Agsya', 'Bank Mandiri')

# Save
doc.save(dst)
print(f'Modified document saved: {dst}')
print('Done! Content replaced while preserving template formatting.')
