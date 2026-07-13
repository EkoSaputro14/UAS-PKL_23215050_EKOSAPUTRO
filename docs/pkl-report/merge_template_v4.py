#!/usr/bin/env python3
"""
Fix v4: Check BAB II/III/IV before BAB I to avoid substring matching.
"""
import os
from docx import Document
from docx.oxml.ns import qn

TEMPLATE = 'C:/Users/SMANSA/AppData/Local/hermes/cache/documents/doc_4ce6543a497c_LAPORAN PKL (1).docx'
SOURCE = 'LAPORAN_PKL_v14_UML.docx'
OUTPUT = 'LAPORAN_PKL_v15_Template.docx'

NEW_TITLE = "RANCANG BANGUN SISTEM CHATBOT AI BERBASIS PENGETAHUAN DENGAN RETRIEVAL-AUGMENTED GENERATION DAN PIPELINE CRM UNTUK OPTIMALISASI LAYANAN PELANGGAN"
NEW_STUDENT = "Eko Saputro"
NEW_NIM = "23215050"
NEW_SUPERVISOR = "Widianto Agung Nugroho"
OLD_STUDENT = "Moh. Arif Prasetyo"
OLD_NIM = "23215043"

print("📄 Loading...")
tpl_doc = Document(TEMPLATE)
src_doc = Document(SOURCE)

# ===== STEP 1: Text replacements =====
print("\n📝 Step 1: Text replacements...")
for para in tpl_doc.paragraphs:
    full_text = para.text
    if "SISTEM PAKAR DETEKSI HAMA" in full_text:
        for run in para.runs:
            if "SISTEM PAKAR" in run.text or "RANCANG BANGUN" in run.text:
                run.text = ""
        for run in para.runs:
            if not run.text.strip():
                run.text = NEW_TITLE
                break
    for run in para.runs:
        if OLD_STUDENT in run.text:
            run.text = run.text.replace(OLD_STUDENT, NEW_STUDENT)
        if OLD_NIM in run.text:
            run.text = run.text.replace(OLD_NIM, NEW_NIM)
        if "Nama Dosen Pembimbing" in run.text:
            run.text = run.text.replace("Nama Dosen Pembimbing", NEW_SUPERVISOR)
        if "NIPY." in run.text:
            run.text = run.text.replace("NIPY.", "")
print("  ✅ Done")

# ===== STEP 2: Detect headings =====
print("\n📝 Step 2: Detecting headings...")

def get_heading_level(para):
    if para.style and para.style.name:
        name = para.style.name
        if name in ('Heading 1', 'Heading1'): return 1
        if name in ('Heading 2', 'Heading2'): return 2
        if name in ('Heading 3', 'Heading3'): return 3
    return None

def classify_bab(text):
    """Classify BAB heading. Check longer names FIRST to avoid substring issues."""
    text = text.strip()
    # Check BAB V before BAB II (V doesn't contain I pattern, but be safe)
    if 'BAB V' in text: return 'BAB_V'
    if 'BAB IV' in text: return 'BAB_IV'
    if 'BAB III' in text: return 'BAB_III'
    if 'BAB II' in text: return 'BAB_II'
    if 'BAB I' in text: return 'BAB_I'
    if 'DAFTAR PUSTAKA' in text: return 'DAFTAR_PUSTAKA'
    if 'LAMPIRAN' in text: return 'LAMPIRAN'
    return None

# Find BAB positions in template
tpl_babs = {}
for i, para in enumerate(tpl_doc.paragraphs):
    level = get_heading_level(para)
    if level != 1: continue
    text = para.text.strip()
    bab_key = classify_bab(text)
    if bab_key:
        if bab_key == 'BAB_I':
            # Verify next paragraph is PENDAHULUAN
            if i + 1 < len(tpl_doc.paragraphs):
                next_text = tpl_doc.paragraphs[i + 1].text.strip()
                if 'PENDAHULUAN' in next_text:
                    tpl_babs[bab_key] = i
        else:
            tpl_babs[bab_key] = i

print(f"  BAB positions: {tpl_babs}")

# ===== STEP 3: Extract source =====
print("\n📝 Step 3: Extracting source...")

src_content = {}
current_bab = None
current_content = []

for para in src_doc.paragraphs:
    text = para.text.strip()
    level = get_heading_level(para)
    
    if level == 1:
        if current_bab and current_content:
            src_content[current_bab] = current_content
        current_content = []
        bab_key = classify_bab(text)
        current_bab = bab_key
        continue
    
    if current_bab:
        current_content.append({'text': text, 'style': para.style.name if para.style else ''})

if current_bab and current_content:
    src_content[current_bab] = current_content

for bab, content in src_content.items():
    print(f"  {bab}: {len(content)} paragraphs")

# ===== STEP 4: Replace content =====
print("\n📝 Step 4: Replacing BAB content...")

sorted_babs = sorted(tpl_babs.items(), key=lambda x: x[1])

for i, (tpl_key, start_idx) in enumerate(sorted_babs):
    end_idx = len(tpl_doc.paragraphs)
    for j in range(i + 1, len(sorted_babs)):
        _, next_idx = sorted_babs[j]
        end_idx = next_idx
        break
    
    if tpl_key not in src_content:
        print(f"  ⚠️ No source for {tpl_key}")
        continue
    
    section_paras = list(tpl_doc.paragraphs[start_idx:end_idx])
    src_paras = src_content[tpl_key]
    
    content_start = 1
    if tpl_key == 'BAB_I' and len(section_paras) > 1:
        if section_paras[1].text.strip() == 'PENDAHULUAN':
            content_start = 2
    
    available = len(section_paras) - content_start
    
    for j, src_p in enumerate(src_paras):
        if j < available:
            target = section_paras[content_start + j]
            for run in target.runs:
                run.text = ''
            if target.runs:
                target.runs[0].text = src_p['text']
            elif src_p['text']:
                target.add_run(src_p['text'])
    
    replaced = min(len(src_paras), available)
    print(f"  ✅ {tpl_key}: {replaced}/{len(src_paras)}")

# ===== SAVE =====
print(f"\n💾 Saving...")
tpl_doc.save(OUTPUT)

# ===== VERIFY =====
doc2 = Document(OUTPUT)
print(f"\n📊 Verification:")
for para in doc2.paragraphs[:30]:
    if "SISTEM CHATBOT AI" in para.text:
        print(f"  ✅ Title OK")
        break

# Check each BAB has Mimotes content
checks = {
    'BAB_I': 'Bank Mandiri KCP Tegal Sudirman',
    'BAB_II': 'Bagian Kredit dan Lelang',
    'BAB_III': 'Retrieval-Augmented Generation',
    'BAB_IV': 'Mimotes AI',
    'BAB_V': 'Kesimpulan',
}
for key, keyword in checks.items():
    found = any(keyword in p.text for p in doc2.paragraphs)
    print(f"  {'✅' if found else '❌'} {key}: {'present' if found else 'MISSING'}")

print(f"  Size: {os.path.getsize(OUTPUT)/1024:.0f} KB")
