"""
Comprehensive PKL report fix — template leaks, TOC, figure numbering, daftar gambar/tabel.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SRC = r'C:\Users\SMANSA\mimotes\docs\pkl-report\LAPORAN_PKL_Eko_Saputro_23215050.docx'
doc = Document(SRC)

# ============================================
# 1. FIX TEMPLATE DATA LEAKS
# ============================================
print("=== 1. Fixing template data leaks ===")

# Paragraph-level fixes (merge all runs, replace, write back)
PARA_FIXES = [
    # Cover page — name
    ("MUHAMMAD AFFIF", "EKO SAPUTRO"),
    # Kata Pengantar — pembimbing lapangan & company
    ("Bapak Nursalim, selaku Direktur Utama PT Bank Mandiri Karya Manca dan selaku Pembimbing Lapangan",
     "Bapak Widianto Agung Nugroho, selaku Pembimbing Lapangan di Bank Mandiri KCP Tegal Sudirman"),
    ("Seluruh tim PT Bank Mandiri Karya Manca atas bimbingan dan dukungannya.",
     "Seluruh tim Bank Mandiri KCP Tegal Sudirman atas bimbingan dan dukungannya."),
    ("PT Bank Mandiri Karya Manca", "Bank Mandiri KCP Tegal Sudirman"),
]

for p in doc.paragraphs:
    runs = p.runs
    if not runs:
        continue
    merged = ''.join(r.text for r in runs)
    original = merged
    for old, new in PARA_FIXES:
        if old in merged:
            merged = merged.replace(old, new)
    if merged != original:
        runs[0].text = merged
        for r in runs[1:]:
            r.text = ''
        print(f"  Fixed: {p.text[:60]}")

# Fix in tables too
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                runs = p.runs
                if not runs:
                    continue
                merged = ''.join(r.text for r in runs)
                original = merged
                for old, new in PARA_FIXES:
                    if old in merged:
                        merged = merged.replace(old, new)
                if merged != original:
                    runs[0].text = merged
                    for r in runs[1:]:
                        r.text = ''

# ============================================
# 2. FIX DAFTAR ISI — replace old TOC with correct one
# ============================================
print("\n=== 2. Fixing Daftar Isi ===")

# Find TOC range and collect indices
toc_indices = []
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name and p.style.name.startswith('toc'):
        toc_indices.append(i)

if toc_indices:
    # Map old TOC entries to new ones
    NEW_TOC = {
        'DAFTAR ISI': 'DAFTAR ISI\tvii',
        'DAFTAR GAMBAR': 'DAFTAR GAMBAR\txi',
        'Lihat halaman Daftar Gambar': 'Lihat halaman Daftar Gambar\txi',
        'DAFTAR TABEL': 'DAFTAR TABEL\txii',
        'Lihat halaman Daftar Tabel': 'Lihat halaman Daftar Tabel\txii',
        'BAB I PENDAHULUAN': 'BAB I PENDAHULUAN\t1',
        '1.1 Latar Belakang': '1.1 Latar Belakang\t1',
        '1.2 Rumusan Masalah': '1.2 Rumusan Masalah\t1',
        '1.3 Tujuan PKL': '1.3 Tujuan PKL\t2',
        '1.3.1 Tujuan Umum': '1.3.1 Tujuan Umum\t2',
        '1.3.2 Tujuan Khusus': '1.3.2 Tujuan Khusus\t2',
        '1.4 Manfaat PKL': '1.4 Manfaat PKL\t3',
        # BAB II
        'BAB II GAMBARAN UMUM PRKATIK KERJA LAPANGAN': 'BAB II GAMBARAN UMUM INSTANSI\t4',
        'BAB II GAMBARAN UMUM INSTANSI': 'BAB II GAMBARAN UMUM INSTANSI\t4',
        '2.1 Sejarah Perkembangan Perusahaan': '2.1 Sejarah Perkembangan Perusahaan\t4',
        '2.2 Visi, Misi, dan Tujuan': '2.2 Visi, Misi, dan Tujuan\t4',
        '2.2.1 Visi': '2.2.1 Visi\t4',
        '2.2.2 Misi': '2.2.2 Misi\t5',
        '2.3 Struktur Organisasi Perusahaan': '2.3 Struktur Organisasi\t5',
        '2.3 Struktur Organisasi': '2.3 Struktur Organisasi\t5',
        '2.4 Job Deskripsi': '2.4 Job Deskripsi\t6',
        # BAB III
        'BAB III TINJAUAN PUSTAKA': 'BAB III METODE PELAKSANAAN PKL\t7',
        'BAB III METODE PELAKSANAAN PKL': 'BAB III METODE PELAKSANAAN PKL\t7',
        '3.1 Landasan Teori': '3.1 Landasan Teori\t7',
        '3.1.1 Keamanan Aplikasi Web': '3.1.1 AI dan Chatbot\t7',
        '3.1.1 Artificial Intelligence dan Chatbot': '3.1.1 AI dan Chatbot\t7',
        '3.1.2 OAuth 2.0': '3.1.2 RAG\t7',
        '3.1.2 Retrieval-Augmented Generation (RAG)': '3.1.2 RAG\t7',
        '3.1.3 JSON Web Token (JWT)': '3.1.3 LLM\t8',
        '3.1.3 NextAuth (JWT)': '3.1.3 LLM\t8',
        '3.1.3 Large Language Model (LLM)': '3.1.3 LLM\t8',
        '3.1.4 Role-Based Access Control (RBAC)': '3.1.4 Embedding dan Vector DB\t8',
        '3.1.4 Embedding dan Vector Database': '3.1.4 Embedding dan Vector DB\t8',
        '3.1.5 Vue.js': '3.1.5 CRM\t8',
        '3.1.5 Next.js': '3.1.5 CRM\t8',
        '3.1.5 Customer Relationship Management (CRM)': '3.1.5 CRM\t8',
        '3.1.6 State Management': '3.1.6 Next.js dan PostgreSQL\t9',
        '3.1.6 Next.js': '3.1.6 Next.js dan PostgreSQL\t9',
        '3.1.6 Next.js dan PostgreSQL': '3.1.6 Next.js dan PostgreSQL\t9',
        '3.1.7 Clean Architecture': '3.1.7 WhatsApp Baileys\t9',
        '3.1.7 API Routes': '3.1.7 WhatsApp Baileys\t9',
        '3.1.7 WhatsApp Integration dengan Baileys': '3.1.7 WhatsApp Baileys\t9',
        '3.2 Kerangka Berpikir': '3.2 Kerangka Berpikir\t9',
        '3.3 Metode Pengembangan': '3.3 Metode Pengembangan\t9',
        # BAB IV
        'BAB IV PELAKSANAAN PRAKTIK KERJA LAPANGAN': 'BAB IV HASIL YANG DICAPAI\t10',
        'BAB IV HASIL YANG DICAPAI': 'BAB IV HASIL YANG DICAPAI\t10',
        '4.1 Tugas Umum': '4.1 Gambaran Umum Sistem\t10',
        '4.1 Gambaran Umum Sistem': '4.1 Gambaran Umum Sistem\t10',
        '4.2 Tugas Khusus': '4.2 Analisis dan Perancangan\t10',
        '4.2 Analisis dan Perancangan Sistem': '4.2 Analisis dan Perancangan\t10',
        '4.2.1 Kebutuhan Fungsional': '4.2.1 Kebutuhan Fungsional\t10',
        '4.2.2 Kebutuhan Non-Fungsional': '4.2.2 Kebutuhan Non-Fungsional\t10',
        '4.2.3 Use Case Diagram': '4.2.3 Use Case Diagram\t11',
        '4.2.4 Activity Diagram Upload Dokumen': '4.2.4 Activity Diagram Upload\t11',
        '4.2.5 Activity Diagram Proses Chat RAG': '4.2.5 Activity Diagram Chat\t11',
        '4.2.6 Entity Relationship Diagram': '4.2.6 ERD\t12',
        '4.2.7 Arsitektur Sistem': '4.2.7 Arsitektur Sistem\t12',
        '4.2.8 Arsitektur RAG Pipeline': '4.2.8 Arsitektur RAG Pipeline\t12',
        '4.2.9 Arsitektur CRM Pipeline': '4.2.9 Arsitektur CRM Pipeline\t12',
        '4.3 Analisis Permasalahan dan Solusi': '4.3 Implementasi Sistem\t13',
        '4.3 Implementasi Sistem': '4.3 Implementasi Sistem\t13',
        '4.3.1 Spesifikasi Perangkat Keras dan Lunak': '4.3.1 Spesifikasi\t13',
        '4.3.2 Implementasi RAG Pipeline': '4.3.2 RAG Pipeline\t13',
        '4.3.3 Implementasi Frontend': '4.3.3 Frontend\t14',
        '4.3.4 Implementasi Backend': '4.3.4 Backend\t14',
        '4.4 Pengujian Sistem': '4.4 Pengujian Sistem\t14',
        '4.5 Pembahasan Hasil': '4.5 Pembahasan Hasil\t15',
        '4.5.1 Keefektifan Pipeline RAG': '4.5.1 RAG\t15',
        '4.5.2 Integrasi CRM': '4.5.2 CRM\t15',
        '4.5.3 Integrasi WhatsApp': '4.5.3 WhatsApp\t15',
        '4.5.4 Multi-Tenancy dan Keamanan': '4.5.4 Multi-Tenancy\t16',
        '4.5.5 Kendala dan Pembelajaran': '4.5.5 Kendala\t16',
        '4.6 Tampilan Antarmuka Sistem': '4.6 Tampilan Antarmuka\t16',
        # BAB V
        'BAB V HASIL YANG DICAPAI PKL': 'BAB V PENUTUP\t18',
        'BAB V PENUTUP': 'BAB V PENUTUP\t18',
        '5.1 Analisis Kebutuhan Sistem': '5.1 Kesimpulan\t18',
        '5.1 Kesimpulan': '5.1 Kesimpulan\t18',
        '5.1.1 Kebutuhan Fungsional': '5.1.1 Kebutuhan Fungsional\t12',
        '5.1.2 Kebutuhan Non-Fungsional': '5.1.2 Kebutuhan Non-Fungsional\t12',
        '5.2 Perancangan Sistem': '5.2 Saran\t18',
        '5.2 Saran': '5.2 Saran\t18',
        '5.2.1 Arsitektur Sistem': '5.2.1 Arsitektur Sistem\t12',
        '5.2.2 Diagram Alur Autentikasi': '5.2.2 Diagram Alur Autentikasi\t13',
        '5.2.3 Diagram Alur OAuth 2.0': '5.2.3 Diagram Alur OAuth\t14',
        '5.2.4 Perancangan RBAC': '5.2.4 Perancangan RBAC\t15',
        '5.2.5 ERD Database': '5.2.5 ERD Database\t16',
        '5.2.6 Perancangan State Management': '5.2.6 Perancangan State Mgmt\t17',
        '5.3 Implementasi Sistem': '5.3 Implementasi Sistem\t18',
        '5.3.1 Implementasi Backend dengan Next.js API Routes': '5.3.1 Backend\t18',
        '5.3.2 Implementasi JWT pada Backend': '5.3.2 NextAuth\t18',
        '5.3.3 Implementasi Auth Middleware': '5.3.3 Auth Middleware\t19',
        '5.3.4 Implementasi Auth Store (Frontend)': '5.3.4 Auth Store\t19',
        '5.3.5 Implementasi API Security Layer': '5.3.5 API Security\t20',
        '5.3.6 Implementasi Route Guards': '5.3.6 Route Guards\t20',
        '5.4 Dokumentasi Tampilan Sistem': '5.4 Dokumentasi Tampilan\t21',
        '5.4.1 Halaman Login': '5.4.1 Login\t21',
        '5.4.2 Dashboard': '5.4.2 Dashboard\t21',
        '5.4.3 Manajemen Tugas': '5.4.3 Chat AI\t21',
        '5.4.4 KPI Tracking': '5.4.4 Knowledge Base\t21',
        '5.4.5 Staff Management': '5.4.5 Analytics\t21',
        '5.5 Pengujian Sistem': '5.5 Pengujian Sistem\t21',
        '5.5.1 Pengujian Fungsional': '5.5.1 Pengujian Fungsional\t21',
        '5.5.2 Pengujian Keamanan': '5.5.2 Pengujian Keamanan\t22',
        '5.6 Pembahasan Hasil': '5.6 Pembahasan Hasil\t23',
        '5.6.1 Keberhasilan Implementasi NextAuth': '5.6.1 RAG Pipeline\t23',
        '5.6.1 Keberhasilan Implementasi JWT': '5.6.1 RAG Pipeline\t23',
        '5.6.2 Efektivitas RBAC': '5.6.2 Multi-Tenancy\t23',
        '5.6.3 Manfaat State Management': '5.6.3 CRM Integration\t23',
        '5.6.4 Keunggulan App Router': '5.6.4 WhatsApp Integration\t24',
        '5.6.4 Keunggulan Clean Architecture': '5.6.4 WhatsApp Integration\t24',
        '5.6.5 Kontribusi Pribadi': '5.6.5 Kontribusi Pribadi\t24',
        '5.6.6 Kendala dan Pembelajaran': '5.6.6 Kendala\t24',
        # BAB VI → now BAB V
        'BAB VI PENUTUP': 'BAB V PENUTUP\t26',
        '6.1 Kesimpulan': '5.1 Kesimpulan\t26',
        '6.2 Saran': '5.2 Saran\t26',
        # DAFTAR PUSTAKA
        'DAFTAR PUSTAKA': 'DAFTAR PUSTAKA\t28',
        'LAMPIRAN': 'LAMPIRAN\t30',
        'Lampiran 1. Surat Keterangan PKL': 'Lampiran A: Logbook\t30',
        'Lampiran 2. Absensi Kehadiran PKL': 'Lampiran B: GitHub\t30',
        'Lampiran 3. Dokumentasi Kegiatan PKL': 'Lampiran C: Screenshots\t30',
    }

    for i in toc_indices:
        p = doc.paragraphs[i]
        runs = p.runs
        if not runs:
            # Try raw XML approach
            t_elems = doc.element.body.findall('.//' + qn('w:p'))[i].findall('.//' + qn('w:t'))
            text = ''.join(t.text or '' for t in t_elems).strip()
        else:
            text = ''.join(r.text for r in runs).strip()

        # Find matching entry (try exact match first, then prefix match)
        matched = False
        for old_text, new_text in NEW_TOC.items():
            if text == old_text or text.startswith(old_text):
                if runs:
                    runs[0].text = new_text
                    for r in runs[1:]:
                        r.text = ''
                else:
                    # XML approach
                    for t in t_elems:
                        if t.text and (t.text.strip() == old_text or t.text.strip().startswith(old_text)):
                            t.text = new_text
                matched = True
                break

        if not matched:
            # Print unmatched for debugging
            if text and len(text) > 3:
                pass  # Skip debug output

# ============================================
# 3. FIX DAFTAR GAMBAR
# ============================================
print("\n=== 3. Fixing Daftar Gambar ===")
# The old Daftar Gambar entries reference SDM/HRIS figures
# Replace with correct Mimotes AI figure references
DAFTAR_GAMBAR_FIXES = {
    'Gambar 5.1 Arsitektur Sistem Informasi Manajemen SDM': 'Gambar 4.2 Use Case Diagram',
    'Gambar 5.2 Alur Autentikasi NextAuth': 'Gambar 4.3 Activity Diagram Upload',
    'Gambar 5.3 Alur Autentikasi Google OAuth 2.0': 'Gambar 4.4 Activity Diagram Chat RAG',
    'Gambar 5.4 Model Role-Based Access Control': 'Gambar 4.5 Entity Relationship Diagram',
    'Gambar 5.5 Entity Relationship Diagram (ERD)': 'Gambar 4.6 Arsitektur Sistem',
}

body = doc.element.body
all_p = body.findall(qn('w:p'))
for elem in all_p:
    t_elems = elem.findall('.//' + qn('w:t'))
    text = ''.join(t.text or '' for t in t_elems)
    for old, new in DAFTAR_GAMBAR_FIXES.items():
        if old in text:
            for t in t_elems:
                if t.text and old in t.text:
                    t.text = t.text.replace(old, new)
            print(f"  Fixed Daftar Gambar: {old}")

# ============================================
# 4. FIX DAFTAR TABEL
# ============================================
print("\n=== 4. Fixing Daftar Tabel ===")
DAFTAR_TABEL_FIXES = {
    'Tabel 2.1 Informasi Perusahaan PT Bank Mandiri Karya Manca': 'Tabel 2.1 Informasi Perusahaan',
    'Tabel 5.1 Matrix Hak Akses Berdasarkan Peran': 'Tabel 4.1 Spesifikasi Perangkat',
    'Tabel 5.2 Daftar State Management Stores': 'Tabel 4.2 Hasil Pengujian',
    'Tabel 5.3 Hasil Pengujian Fungsional': 'Tabel 4.3 Hasil Pengujian Fungsional',
    'Tabel 5.4 Hasil Pengujian Keamanan': 'Tabel 4.4 Logbook Kegiatan',
    'Tabel 5.5 Kontribusi Pribadi Selama PKL': 'Tabel 4.5 Kontribusi Pribadi',
}

for elem in all_p:
    t_elems = elem.findall('.//' + qn('w:t'))
    text = ''.join(t.text or '' for t in t_elems)
    for old, new in DAFTAR_TABEL_FIXES.items():
        if old in text:
            for t in t_elems:
                if t.text and old in t.text:
                    t.text = t.text.replace(old, new)
            print(f"  Fixed Daftar Tabel: {old}")

# ============================================
# 5. FIX SCREENSHOT NUMBERING (4.1-4.8 → 4.9-4.16)
# ============================================
print("\n=== 5. Fixing screenshot numbering ===")
SCREENSHOT_NUMBERING = {
    'Gambar 4.1 Halaman Login': 'Gambar 4.9 Halaman Login',
    'Gambar 4.2 Dashboard Admin': 'Gambar 4.10 Dashboard Admin',
    'Gambar 4.3 Halaman Upload Dokumen': 'Gambar 4.11 Upload Dokumen',
    'Gambar 4.4 Daftar Dokumen': 'Gambar 4.12 Daftar Dokumen',
    'Gambar 4.5 Halaman Chat AI dengan Sumber': 'Gambar 4.13 Chat AI',
    'Gambar 4.6 Knowledge Search': 'Gambar 4.14 Knowledge Search',
    'Gambar 4.7 Analytics Chat': 'Gambar 4.15 Analytics Chat',
    'Gambar 4.8 Pengaturan AI Provider': 'Gambar 4.16 Pengaturan AI Provider',
}

for p in doc.paragraphs:
    runs = p.runs
    if not runs:
        continue
    merged = ''.join(r.text for r in runs)
    for old, new in SCREENSHOT_NUMBERING.items():
        if old in merged:
            merged = merged.replace(old, new)
            runs[0].text = merged
            for r in runs[1:]:
                r.text = ''
            print(f"  Renumbered: {new}")

# Also fix in XML
for elem in all_p:
    t_elems = elem.findall('.//' + qn('w:t'))
    text = ''.join(t.text or '' for t in t_elems)
    for old, new in SCREENSHOT_NUMBERING.items():
        if old in text:
            for t in t_elems:
                if t.text and old in t.text:
                    t.text = t.text.replace(old, new)

# ============================================
# 6. FIX Qdrant REFERENCE (should be pgvector)
# ============================================
print("\n=== 6. Fixing Qdrant → pgvector ===")
for p in doc.paragraphs:
    runs = p.runs
    if not runs:
        continue
    merged = ''.join(r.text for r in runs)
    if 'Qdrant' in merged and 'pgvector' not in merged:
        merged = merged.replace('Qdrant', 'pgvector')
        runs[0].text = merged
        for r in runs[1:]:
            r.text = ''
        print(f"  Fixed Qdrant reference")

# ============================================
# 7. ADD TABLE BORDERS
# ============================================
print("\n=== 7. Adding table borders ===")
for table in doc.tables:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is None:
        borders = OxmlElement('w:tblBorders')
        for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{bn}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')
            borders.append(b)
        tblPr.append(borders)

# ============================================
# SAVE
# ============================================
doc.save(SRC)
print(f"\n=== DONE === Saved: {SRC}")
print(f"Size: {os.path.getsize(SRC) / 1024:.0f} KB")

# Final leak check
doc2 = Document(SRC)
leaks = []
bad_terms = ['Muhammad Affif', '24225046', 'Syefudin', 'PT Agsya', 'Agsya Karya', 'PT Bank Mandiri Karya Manca',
             'Nursalim', 'Vue.js', 'Pinia', 'Go Fiber', 'Clean Architecture', 'Implementasi Keamanan',
             'Manajemen SDM', 'Informasi Manajemen', 'OAuth 2.0', 'KPI Tracking', 'Staff Management']
for i, p in enumerate(doc2.paragraphs):
    for term in bad_terms:
        if term in p.text:
            leaks.append(f"  para[{i}]: {term} — {p.text[:60]}")
            break
print(f"\nFinal leak check: {len(leaks)} leaks found")
for l in leaks:
    print(l)
