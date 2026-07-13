#!/usr/bin/env python3
"""
Complete rebuild: Take template formatting, fill with ALL v14 content.
No mixing - clean separation between template shell and v14 body.
"""
import os
from docx import Document
from docx.shared import Inches, Cm
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = 'C:/Users/SMANSA/AppData/Local/hermes/cache/documents/doc_4ce6543a497c_LAPORAN PKL (1).docx'
SOURCE = 'LAPORAN_PKL_v14_UML.docx'
OUTPUT = 'LAPORAN_PKL_v15_Template.docx'
DIAGRAMS_DIR = 'diagrams'

# New info
NEW_TITLE = "RANCANG BANGUN SISTEM CHATBOT AI BERBASIS PENGETAHUAN DENGAN RETRIEVAL-AUGMENTED GENERATION DAN PIPELINE CRM UNTUK OPTIMALISASI LAYANAN PELANGGAN"
NEW_STUDENT = "Eko Saputro"
NEW_NIM = "23215050"
NEW_SUPERVISOR = "Widianto Agung Nugroho"
OLD_STUDENT = "Moh. Arif Prasetyo"
OLD_NIM = "23215043"

print("📄 Step 1: Read template structure...")
tpl = Document(TEMPLATE)

# Find the LAST Heading1 before BAB I (this is the end of front matter)
front_matter_end = None
bab_i_start = None
for i, para in enumerate(tpl.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    if style == 'Heading 1' and 'BAB I' in text:
        # Check next para is PENDAHULUAN
        if i + 1 < len(tpl.paragraphs):
            nxt = tpl.paragraphs[i + 1].text.strip()
            if 'PENDAHULUAN' in nxt:
                bab_i_start = i
                break

# Find DAFTAR PUSTAKA position (before old content)
dafpus_before = None
for i, para in enumerate(tpl.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    if style == 'Heading 1' and 'Daftar Isi' in text:
        dafpus_before = i
        break

print(f"  Front matter: 0-{dafpus_before}")
print(f"  BAB I starts at: {bab_i_start}")

# Find ALL Heading1 positions to understand sections
print("\n  All Heading 1 sections:")
for i, para in enumerate(tpl.paragraphs):
    if para.style.name == 'Heading 1':
        print(f"    [{i:3d}] {para.text.strip()[:50]}")

print("\n📄 Step 2: Read source content...")
src = Document(SOURCE)

# Extract ALL content from source (everything)
src_paras = []
for para in src.paragraphs:
    text = para.text.strip()
    style = para.style.name if para.style else ''
    runs_data = []
    for run in para.runs:
        runs_data.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
        })
    src_paras.append({
        'text': text,
        'style': style,
        'runs': runs_data,
    })

print(f"  Source: {len(src_paras)} paragraphs")

# Find where content starts in source (after front matter)
src_content_start = 0
for i, p in enumerate(src_paras):
    if p['style'] in ('Heading1', 'Heading 1') and 'BAB I' in p['text']:
        # Check next para
        if i + 1 < len(src_paras) and 'PENDAHULUAN' in src_paras[i+1]['text']:
            src_content_start = i
            break

print(f"  Source content starts at: {src_content_start}")

print("\n📄 Step 3: Rebuild document...")

# Strategy: 
# 1. Keep template paragraphs 0 to (dafpus_before-1) = front matter
# 2. Replace everything from dafpus_before onwards with source content

# First, get the XML body
body = tpl.element.body
children = list(body)

# Find the Daftar Isi heading element
dai_heading_idx = None
for i, child in enumerate(children):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        pPr = child.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                style_val = pStyle.get(qn('w:val'), '')
                if style_val in ('Heading1', 'Judul1'):
                    text = ''.join(t.text or '' for t in child.iter(qn('w:t'))).strip()
                    if 'Daftar Isi' in text:
                        dai_heading_idx = i
                        break

print(f"  Daftar Isi heading element index: {dai_heading_idx}")

if dai_heading_idx is not None:
    # Remove everything from Daftar Isi onwards
    to_remove = children[dai_heading_idx:]
    for elem in to_remove:
        body.remove(elem)
    
    print(f"  Removed {len(to_remove)} elements from Daftar Isi onwards")
    
    # Now add source content starting from BAB I
    # But first, add Daftar Isi, Daftar Gambar, Daftar Tabel headings
    # (we'll populate them later)
    
    # Add source content paragraphs
    print("\n📄 Step 4: Adding source content...")
    
    # Get the last section properties (keep it)
    sectPr = body.find(qn('w:sectPr'))
    
    added = 0
    for p in src_paras[src_content_start:]:
        if not p['text'] and not p['style'].startswith('Heading'):
            continue  # Skip empty non-heading paragraphs
        
        new_p = etree.SubElement(body, qn('w:p'))
        
        # Set style
        if p['style']:
            pPr = etree.SubElement(new_p, qn('w:pPr'))
            pStyle = etree.SubElement(pPr, qn('w:pStyle'))
            # Map source style names to template style IDs
            style_map = {
                'Heading1': 'Judul1',
                'Heading 1': 'Judul1',
                'Heading2': 'Judul2',
                'Heading 2': 'Judul2',
                'Heading3': 'Judul3',
                'Heading 3': 'Judul3',
                'Normal': 'Normal',
                'Normal (Web)': 'Normal',
                'List Paragraph': 'ListParagraph',
            }
            mapped_style = style_map.get(p['style'], p['style'])
            pStyle.set(qn('w:val'), mapped_style)
        
        # Add text runs
        if p['text']:
            r = etree.SubElement(new_p, qn('w:r'))
            t = etree.SubElement(r, qn('w:t'))
            t.text = p['text']
            t.set(qn('xml:space'), 'preserve')
            
            # Apply formatting
            if p['runs']:
                rPr = etree.SubElement(r, qn('w:rPr'))
                if p['runs'][0].get('bold'):
                    b = etree.SubElement(rPr, qn('w:b'))
                if p['runs'][0].get('italic'):
                    i_elem = etree.SubElement(rPr, qn('w:i'))
        
        added += 1
    
    print(f"  Added {added} paragraphs")

print("\n📄 Step 5: Replace front matter text...")

# Re-load to get proper paragraph objects
tpl.save(OUTPUT)
doc2 = Document(OUTPUT)

for para in doc2.paragraphs:
    full_text = para.text
    
    # Fix title
    if "SISTEM PAKAR" in full_text:
        for run in para.runs:
            if "SISTEM PAKAR" in run.text or "RANCANG BANGUN" in run.text:
                run.text = ""
        for run in para.runs:
            if not run.text.strip():
                run.text = NEW_TITLE
                break
    
    # Fix student/NIM
    for run in para.runs:
        if OLD_STUDENT in run.text:
            run.text = run.text.replace(OLD_STUDENT, NEW_STUDENT)
        if OLD_NIM in run.text:
            run.text = run.text.replace(OLD_NIM, NEW_NIM)
        if "Nama Dosen Pembimbing" in run.text:
            run.text = run.text.replace("Nama Dosen Pembimbing", NEW_SUPERVISOR)
        if "NIPY." in run.text:
            run.text = run.text.replace("NIPY.", "")

doc2.save(OUTPUT)
print("  ✅ Front matter updated")

# ===== STEP 6: Insert images =====
print("\n📄 Step 6: Inserting images...")

doc3 = Document(OUTPUT)

# Find all Gambar 4.x captions and insert images before them
img_inserted = 0
# Caption to image mapping
CAPTION_IMG = {
    "Use Case": "usecase.png",
    "Activity Diagram Upload": "activity-upload.png",
    "Activity Diagram Proses Chat": "activity-chat.png",
    "Sequence Diagram Chat RAG": "sequence-chat-rag.png",
    "ERD Domain Identity": "erd-a-identity.png",
    "ERD Domain RAG": "erd-b-rag.png",
    "ERD Domain Chat": "erd-c-chat-crm.png",
    "ERD Domain Billing": "erd-d-billing.png",
    "ERD Ringkasan": "erd-summary.png",
    "Arsitektur Sistem": "architecture.png",
    "Arsitektur RAG": "rag-pipeline.png",
    "Arsitektur CRM": "crm-pipeline.png",
}

for i, para in enumerate(doc3.paragraphs):
    text = para.text.strip()
    if not text.startswith('Gambar 4.'):
        continue
    
    for keyword, img_file in CAPTION_IMG.items():
        if keyword in text:
            img_path = os.path.join(DIAGRAMS_DIR, img_file)
            if not os.path.exists(img_path):
                continue
            
            try:
                # Insert image paragraph before caption
                new_para = doc3.add_paragraph()
                new_para.alignment = 1  # CENTER
                new_para.add_run().add_picture(img_path, width=Inches(5.5))
                
                # Move before caption
                parent = para._element.getparent()
                pos = list(parent).index(para._element)
                new_elem = new_para._element
                parent.remove(new_elem)
                parent.insert(pos, new_elem)
                
                img_inserted += 1
                print(f"  ✅ {text[:50]} + {img_file}")
            except Exception as e:
                print(f"  ❌ {text[:30]}: {e}")
            break

doc3.save(OUTPUT)
print(f"\n  Images inserted: {img_inserted}")

# ===== VERIFY =====
doc4 = Document(OUTPUT)
total_imgs = sum(
    len(run._element.findall(qn('w:drawing')))
    for para in doc4.paragraphs
    for run in para.runs
)

print(f"\n📊 Final:")
print(f"  Paragraphs: {len(doc4.paragraphs)}")
print(f"  Images: {total_imgs}")
print(f"  Size: {os.path.getsize(OUTPUT)/1024:.0f} KB")

# Check for any remaining old content
old_content = False
for para in doc4.paragraphs:
    if 'Sistem Pakar Deteksi Hama' in para.text:
        old_content = True
        print(f"  ⚠️ OLD CONTENT FOUND: {para.text[:50]}")
        break

if not old_content:
    print(f"  ✅ No old template content remaining")
