#!/usr/bin/env python3
"""
Rebuild v4: CLEAR all section content first, THEN fill with source.
"""
import os
import copy
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = 'C:/Users/SMANSA/AppData/Local/hermes/cache/documents/doc_4ce6543a497c_LAPORAN PKL (1).docx'
SOURCE = 'LAPORAN_PKL_v14_UML.docx'
OUTPUT = 'LAPORAN_PKL_v15_Template.docx'
DIAGRAMS_DIR = 'diagrams'

NEW_TITLE = "RANCANG BANGUN SISTEM CHATBOT AI BERBASIS PENGETAHUAN DENGAN RETRIEVAL-AUGMENTED GENERATION DAN PIPELINE CRM UNTUK OPTIMALISASI LAYANAN PELANGGAN"
NEW_STUDENT = "Eko Saputro"
NEW_NIM = "23215050"
NEW_SUPERVISOR = "Widianto Agung Nugroho"
OLD_STUDENT = "Moh. Arif Prasetyo"
OLD_NIM = "23215043"

def get_heading_level(para):
    name = para.style.name if para.style else ''
    if name in ('Heading 1', 'Heading1'): return 1
    if name in ('Heading 2', 'Heading2'): return 2
    if name in ('Heading 3', 'Heading3'): return 3
    return None

def classify_bab(text):
    first_line = text.strip().split('\n')[0].strip()
    if 'BAB V' in first_line: return 'BAB_V'
    if 'BAB IV' in first_line: return 'BAB_IV'
    if 'BAB III' in first_line: return 'BAB_III'
    if 'BAB II' in first_line: return 'BAB_II'
    if 'BAB I' in first_line: return 'BAB_I'
    if 'DAFTAR PUSTAKA' in first_line: return 'DAFTAR_PUSTAKA'
    if 'LAMPIRAN' in first_line: return 'LAMPIRAN'
    return None

print("📄 Loading...")
tpl_doc = Document(TEMPLATE)
src_doc = Document(SOURCE)

# ===== Extract source sections =====
print("\n📝 Extracting source...")
src_sections = {}
current_key = None
current_paras = []

for para in src_doc.paragraphs:
    text = para.text.strip()
    level = get_heading_level(para)
    
    if level == 1:
        if current_key:
            src_sections[current_key] = current_paras
        current_paras = []
        current_key = classify_bab(text)
        continue
    
    if current_key and text:  # Skip empty paragraphs
        current_paras.append({
            'text': text,
            'style_name': para.style.name if para.style else 'Normal',
        })

if current_key:
    src_sections[current_key] = current_paras

for key, paras in src_sections.items():
    print(f"  {key}: {len(paras)} paragraphs")

# ===== Find template sections =====
print("\n📝 Finding template sections...")
tpl_sections = {}
for i, para in enumerate(tpl_doc.paragraphs):
    level = get_heading_level(para)
    if level == 1:
        key = classify_bab(para.text)
        if key:
            tpl_sections[key] = i

# ===== Replace content: CLEAR then FILL =====
print("\n📝 Replacing content (clear + fill)...")

sorted_sections = sorted(tpl_sections.items(), key=lambda x: x[1])

for sec_i, (key, start_idx) in enumerate(sorted_sections):
    if key not in src_sections:
        print(f"  ⚠️ No source for {key}")
        continue
    
    end_idx = len(tpl_doc.paragraphs)
    for j in range(sec_i + 1, len(sorted_sections)):
        _, next_idx = sorted_sections[j]
        end_idx = next_idx
        break
    
    section_paras = list(tpl_doc.paragraphs[start_idx:end_idx])
    src_content = src_sections[key]
    
    # Find where content starts (after heading)
    content_start = 1
    if key == 'BAB_I' and len(section_paras) > 1:
        nxt = section_paras[1].text.strip()
        if 'PENDAHULUAN' in nxt or nxt == '':
            content_start = 2
    
    # CLEAR all content paragraphs (set text to empty)
    for j in range(content_start, len(section_paras)):
        for run in section_paras[j].runs:
            run.text = ''
    
    # FILL with source content
    for j, src_p in enumerate(src_content):
        if j + content_start < len(section_paras):
            target = section_paras[j + content_start]
            if target.runs:
                target.runs[0].text = src_p['text']
            elif src_p['text']:
                target.add_run(src_p['text'])
    
    # If source has MORE paragraphs, insert extras
    excess = len(src_content) - (len(section_paras) - content_start)
    if excess > 0:
        last_para = section_paras[-1]
        last_elem = last_para._element
        parent = last_elem.getparent()
        insert_pos = list(parent).index(last_elem) + 1
        
        fmt_template = section_paras[-1]
        
        for k in range(excess):
            src_p = src_content[(len(section_paras) - content_start) + k]
            new_elem = copy.deepcopy(fmt_template._element)
            for t_elem in new_elem.iter(qn('w:t')):
                t_elem.text = ''
            runs = new_elem.findall(f'.//{qn("w:r")}')
            if runs:
                t_elems = runs[0].findall(qn('w:t'))
                if t_elems:
                    t_elems[0].text = src_p['text']
                else:
                    t = etree.SubElement(runs[0], qn('w:t'))
                    t.text = src_p['text']
                    t.set(qn('xml:space'), 'preserve')
            parent.insert(insert_pos + k, new_elem)
        
        print(f"  ✅ {key}: cleared + {len(src_content)} filled + {excess} inserted")
    else:
        print(f"  ✅ {key}: cleared + {len(src_content)} filled")

# ===== Fix front matter =====
print("\n📝 Fixing front matter...")
for para in tpl_doc.paragraphs:
    if "SISTEM PAKAR" in para.text:
        for run in para.runs:
            if "SISTEM PAKAR" in run.text or "RANCANG BANGUN" in run.text:
                run.text = ""
        for run in para.runs:
            if not run.text.strip():
                run.text = NEW_TITLE
                break
    for run in para.runs:
        if OLD_STUDENT in run.text:
            run.text = run.text.replace(OLD_STUDENT, NEW_STUDENT)
        if OLD_NIM in run.text:
            run.text = run.text.replace(OLD_NIM, NEW_NIM)
        if "Nama Dosen Pembimbing" in run.text:
            run.text = run.text.replace("Nama Dosen Pembimbing", NEW_SUPERVISOR)
        if "NIPY." in run.text:
            run.text = run.text.replace("NIPY.", "")

# ===== Save =====
print("\n💾 Saving...")
tpl_doc.save(OUTPUT)

# ===== Insert images =====
print("\n🖼️ Inserting images...")
doc2 = Document(OUTPUT)

CAPTION_IMG = {
    "Use Case Diagram": "usecase.png",
    "Activity Diagram Upload": "activity-upload.png",
    "Activity Diagram Proses Chat": "activity-chat.png",
    "Sequence Diagram Chat RAG": "sequence-chat-rag.png",
    "ERD Domain Identity": "erd-a-identity.png",
    "ERD Domain RAG": "erd-b-rag.png",
    "ERD Domain Chat": "erd-c-chat-crm.png",
    "ERD Domain Billing": "erd-d-billing.png",
    "ERD Ringkasan": "erd-summary.png",
    "Arsitektur Sistem Mimotes": "architecture.png",
    "Arsitektur RAG": "rag-pipeline.png",
    "Arsitektur CRM": "crm-pipeline.png",
}

img_count = 0
for i, para in enumerate(doc2.paragraphs):
    text = para.text.strip()
    if not text.startswith('Gambar 4.'):
        continue
    for keyword, img_file in CAPTION_IMG.items():
        if keyword in text:
            img_path = os.path.join(DIAGRAMS_DIR, img_file)
            if not os.path.exists(img_path):
                continue
            try:
                new_para = doc2.add_paragraph()
                new_para.alignment = 1
                new_para.add_run().add_picture(img_path, width=Inches(5.5))
                parent = para._element.getparent()
                pos = list(parent).index(para._element)
                new_elem = new_para._element
                parent.remove(new_elem)
                parent.insert(pos, new_elem)
                img_count += 1
                print(f"  ✅ {img_file}")
            except Exception as e:
                print(f"  ❌ {img_file}: {e}")
            break

doc2.save(OUTPUT)

# ===== Verify =====
doc3 = Document(OUTPUT)
total_imgs = sum(
    len(run._element.findall(qn('w:drawing')))
    for para in doc3.paragraphs
    for run in para.runs
)
old_words = ['Sistem Pakar Deteksi Hama', 'Dinas Pertanian', 'Brebes', 'SIPATAN', 'Certainty Factor', 'CF Kombinasi']
old_found = [w for w in old_words if any(w in p.text for p in doc3.paragraphs)]

print(f"\n📊 Final:")
print(f"  Paragraphs: {len(doc3.paragraphs)}")
print(f"  Images: {total_imgs}")
print(f"  Old content: {'⚠️ ' + str(old_found) if old_found else '✅ Clean'}")
print(f"  Size: {os.path.getsize(OUTPUT)/1024:.0f} KB")
