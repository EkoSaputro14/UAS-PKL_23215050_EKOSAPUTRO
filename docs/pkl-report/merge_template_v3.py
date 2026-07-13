#!/usr/bin/env python3
"""
Fix v3: Use para.style.name (display name) instead of XML style_val for heading detection.
"""
import os
from docx import Document
from docx.oxml.ns import qn

TEMPLATE = 'C:/Users/SMANSA/AppData/Local/hermes/cache/documents/doc_4ce6543a497c_LAPORAN PKL (1).docx'
SOURCE = 'LAPORAN_PKL_v14_UML.docx'
OUTPUT = 'LAPORAN_PKL_v15_Template.docx'

NEW_TITLE = "RANCANG BANGUN SISTEM CHATBOT AI BERBASIS PENGETAHUAN DENGAN RETRIEVAL-AUGMENTED GENERATION DAN PIPELINE CRM UNTUK OPTIMALISASI LAYANAN PELANGGAN"
NEW_STUDENT = "Eko Saputro"
NEW_NIM = "23215050"
NEW_SUPERVISOR = "Widianto Agung Nugroho"
OLD_STUDENT = "Moh. Arif Prasetyo"
OLD_NIM = "23215043"

print("📄 Loading...")
tpl_doc = Document(TEMPLATE)
src_doc = Document(SOURCE)

# ===== STEP 1: Text replacements =====
print("\n📝 Step 1: Text replacements...")
for para in tpl_doc.paragraphs:
    full_text = para.text
    if "SISTEM PAKAR DETEKSI HAMA" in full_text:
        for run in para.runs:
            if "SISTEM PAKAR" in run.text or "RANCANG BANGUN" in run.text:
                run.text = ""
        for run in para.runs:
            if not run.text.strip():
                run.text = NEW_TITLE
                break
    for run in para.runs:
        if OLD_STUDENT in run.text:
            run.text = run.text.replace(OLD_STUDENT, NEW_STUDENT)
        if OLD_NIM in run.text:
            run.text = run.text.replace(OLD_NIM, NEW_NIM)
        if "Nama Dosen Pembimbing" in run.text:
            run.text = run.text.replace("Nama Dosen Pembimbing", NEW_SUPERVISOR)
        if "NIPY." in run.text:
            run.text = run.text.replace("NIPY.", "")

print("  ✅ Done")

# ===== STEP 2: Detect headings using style.display_name =====
print("\n📝 Step 2: Detecting headings...")

def get_heading_level(para):
    """Return heading level (1,2,3) or None using display name."""
    if para.style and para.style.name:
        name = para.style.name
        if name == 'Heading 1' or name == 'Heading1':
            return 1
        if name == 'Heading 2' or name == 'Heading2':
            return 2
        if name == 'Heading 3' or name == 'Heading3':
            return 3
    return None

# Find BAB positions
tpl_babs = {}
for i, para in enumerate(tpl_doc.paragraphs):
    level = get_heading_level(para)
    if level != 1:
        continue
    text = para.text.strip()
    if 'BAB I' in text and 'PENDAHULUAN' not in text:
        if i + 1 < len(tpl_doc.paragraphs):
            next_text = tpl_doc.paragraphs[i + 1].text.strip()
            if 'PENDAHULUAN' in next_text:
                tpl_babs['BAB_I'] = i
    elif 'BAB II' in text:
        tpl_babs['BAB_II'] = i
    elif 'BAB III' in text:
        tpl_babs['BAB_III'] = i
    elif 'BAB IV' in text:
        tpl_babs['BAB_IV'] = i
    elif 'BAB V' in text:
        tpl_babs['BAB_V'] = i
    elif 'DAFTAR PUSTAKA' in text:
        tpl_babs['DAFTAR_PUSTAKA'] = i
    elif 'LAMPIRAN' in text:
        tpl_babs['LAMPIRAN'] = i

print(f"  BAB positions: {tpl_babs}")

# ===== STEP 3: Extract source content =====
print("\n📝 Step 3: Extracting source...")

src_content = {}
current_bab = None
current_content = []

for para in src_doc.paragraphs:
    text = para.text.strip()
    level = get_heading_level(para)
    
    if level == 1:
        if current_bab and current_content:
            src_content[current_bab] = current_content
        current_content = []
        
        if 'BAB I' in text and 'PENDAHULUAN' not in text:
            current_bab = 'BAB_I'
        elif 'BAB II' in text:
            current_bab = 'BAB_II'
        elif 'BAB III' in text:
            current_bab = 'BAB_III'
        elif 'BAB IV' in text:
            current_bab = 'BAB_IV'
        elif 'BAB V' in text:
            current_bab = 'BAB_V'
        elif 'DAFTAR PUSTAKA' in text:
            current_bab = 'DAFTAR_PUSTAKA'
        elif 'LAMPIRAN' in text:
            current_bab = 'LAMPIRAN'
        else:
            current_bab = None
        continue
    
    if current_bab:
        current_content.append({'text': text, 'style': para.style.name if para.style else ''})

if current_bab and current_content:
    src_content[current_bab] = current_content

for bab, content in src_content.items():
    print(f"  {bab}: {len(content)} paragraphs")

# ===== STEP 4: Replace content =====
print("\n📝 Step 4: Replacing BAB content...")

sorted_babs = sorted(tpl_babs.items(), key=lambda x: x[1])

for i, (tpl_key, start_idx) in enumerate(sorted_babs):
    end_idx = len(tpl_doc.paragraphs)
    for j in range(i + 1, len(sorted_babs)):
        _, next_idx = sorted_babs[j]
        end_idx = next_idx
        break
    
    if tpl_key not in src_content:
        print(f"  ⚠️ No source for {tpl_key}")
        continue
    
    section_paras = list(tpl_doc.paragraphs[start_idx:end_idx])
    src_paras = src_content[tpl_key]
    
    # Skip heading paragraph + optional subheading
    content_start = 1
    if tpl_key == 'BAB_I' and len(section_paras) > 1:
        if section_paras[1].text.strip() in ('PENDAHULUAN', 'PENDAHULUAN'):
            content_start = 2
    
    available = len(section_paras) - content_start
    
    for j, src_p in enumerate(src_paras):
        if j < available:
            target = section_paras[content_start + j]
            for run in target.runs:
                run.text = ''
            if target.runs:
                target.runs[0].text = src_p['text']
            elif src_p['text']:
                target.add_run(src_p['text'])
    
    replaced = min(len(src_paras), available)
    print(f"  ✅ {tpl_key}: {replaced}/{len(src_paras)}")

# ===== SAVE =====
print(f"\n💾 Saving...")
tpl_doc.save(OUTPUT)

# ===== VERIFY =====
doc2 = Document(OUTPUT)
print(f"\n📊 Verification:")

# Check title
for para in doc2.paragraphs[:30]:
    if "SISTEM CHATBOT AI" in para.text:
        if "RANCANG BANGUN RANCANG BANGUN" in para.text:
            print(f"  ❌ Title duplicated!")
        else:
            print(f"  ✅ Title OK: {para.text[:50]}")
        break

# Check all Babs have content
for key in ['BAB_I', 'BAB_II', 'BAB_III', 'BAB_IV', 'BAB_V']:
    found = False
    for para in doc2.paragraphs:
        text = para.text.strip()
        if key == 'BAB_I' and 'Bank Mandiri' in text and len(text) > 50:
            found = True
            break
        elif key == 'BAB_II' and 'Bank Mandiri KCP' in text:
            found = True
            break
        elif key == 'BAB_III' and 'Retrieval-Augmented' in text and len(text) > 50:
            found = True
            break
        elif key == 'BAB_IV' and 'Mimotes AI' in text and len(text) > 50:
            found = True
            break
        elif key == 'BAB_V' and 'Kesimpulan' in text:
            found = True
            break
    status = "✅" if found else "❌"
    print(f"  {status} {key} content: {'present' if found else 'MISSING'}")

print(f"  Size: {os.path.getsize(OUTPUT)/1024:.0f} KB")
