from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ===== PAGE SETUP =====
for section in doc.sections:
    section.top_margin = Cm(4)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

# ===== STYLES =====
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# Heading styles
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.paragraph_format.line_spacing = 2.0
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)

doc.styles['Heading 1'].font.size = Pt(14)
doc.styles['Heading 2'].font.size = Pt(12)
doc.styles['Heading 3'].font.size = Pt(12)

DIAG_DIR = r'C:\Users\SMANSA\mimotes\docs\pkl-report\diagrams'

# ===== HELPER FUNCTIONS =====
def add_page_number():
    """Add page number to footer of current section"""
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    
    # Add page number field
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    
    run2 = p.add_run()
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(10)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    
    run3 = p.add_run()
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(10)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

def add_roman_page_number():
    """Add roman numeral page number to footer"""
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Roman numeral format
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    
    run2 = p.add_run()
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(10)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    
    run3 = p.add_run()
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(10)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)
    
    # Set roman numeral format
    rPr = run._r.get_or_add_rPr()
    numFmt = parse_xml(f'<w:numFmt {nsdecls("w")} w:val="lowerRoman"/>')
    # This needs to be on the field result, not the run
    # We'll handle it differently - just use simple page numbers for now

def center_text(text, bold=False, size=12, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    return p

def body_text(text, bold=False, indent=True, align='justify'):
    p = doc.add_paragraph()
    if align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = bold
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.line_spacing = 2.0
    h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_bullet_list(items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        r = p.add_run(f'{i}. {item}')
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

def add_image_centered(filename, caption):
    img_path = os.path.join(DIAG_DIR, filename)
    if os.path.exists(img_path):
        # Image paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(img_path, width=Cm(14))
        
        # Caption paragraph
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing = 2.0
        cap.paragraph_format.space_after = Pt(12)
        cr = cap.add_run(caption)
        cr.font.name = 'Times New Roman'
        cr.font.size = Pt(10)
        cr.italic = True

def add_table_simple(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = 1  # CENTER
    
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        # Gray background for header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8D5C0" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    
    return table

def page_break():
    doc.add_page_break()

# ============================================================
# COVER PAGE (Roman numeral i)
# ============================================================
for _ in range(4):
    doc.add_paragraph()

center_text('LAPORAN PRAKTIK KERJA LAPANGAN', True, 16, 12)
center_text('RANCANG BANGUN SISTEM CHATBOT AI\nBERBASIS PENGETAHUAN DENGAN\nRETRIEVAL-AUGMENTED GENERATION\nDAN PIPELINE CRM UNTUK\nOPTIMALISASI LAYANAN PELANGGAN', True, 14, 24)

for _ in range(3):
    doc.add_paragraph()

center_text('Nama\t\t: Eko Saputro', False, 12, 6)
center_text('NIM\t\t: 23215050', False, 12, 6)
center_text('Kelas\t\t: 6A', False, 12, 12)

center_text('PROGRAM STUDI S1 TEKNIK INFORMATIKA', False, 12, 6)
center_text('FAKULTAS SAINS & TEKNOLOGI', False, 12, 6)
center_text('UNIVERSITAS HARKAT NEGERI', False, 12, 12)
center_text('Tahun 2026', False, 12)

add_roman_page_number()
page_break()

# ============================================================
# LEMBAR PENGESAHAN
# ============================================================
center_text('LEMBAR PENGESAHAN', True, 14, 24)

body_text('Laporan Praktik Kerja Lapangan (PKL) ini telah disetujui dan dipertanggungjawabkan oleh:', False, False)

doc.add_paragraph()
body_text('Dosen Pembimbing:', True, False)
body_text('Zaenul Arif, M.Kom.', False, False)
doc.add_paragraph()
body_text('Pembimbing Lapangan:', True, False)
body_text('Widianto Agung Nugroho', False, False)
doc.add_paragraph()
body_text('Ketua Program Studi:', True, False)
body_text('Aang Alim Murtopo, M.Kom.', False, False)
body_text('NIPY. 08.025.555', False, False)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Tegal, 25 Juni 2026')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('......................................')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Eko Saputro')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('NIM. 23215050')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

add_roman_page_number()
page_break()

# ============================================================
# KATA PENGANTAR
# ============================================================
center_text('KATA PENGANTAR', True, 14, 24)

body_text('Puji syukur kehadirat Tuhan Yang Maha Esa atas rahmat dan karunia-Nya sehingga laporan Praktik Kerja Lapangan (PKL) Program Studi S1 Teknik Informatika Universitas Harkat Negeri ini dapat diselesaikan dengan baik.')

body_text('Laporan ini disusun sebagai bentuk pertanggungjawaban akademik atas kegiatan Praktik Kerja Lapangan yang telah dilaksanakan di Bank Mandiri KCP Tegal Sudirman selama periode 25 Mei hingga 25 Juni 2026.')

body_text('Penulis menyadari bahwa laporan ini tidak dapat diselesaikan tanpa bantuan dan dukungan dari berbagai pihak. Oleh karena itu, penulis mengucapkan terima kasih kepada:')

add_bullet_list([
    'Bapak Zaenul Arif, M.Kom. selaku Dosen Pembimbing yang telah memberikan bimbingan dan arahan selama kegiatan PKL.',
    'Bapak Widianto Agung Nugroho selaku Pembimbing Lapangan di Bank Mandiri KCP Tegal Sudirman.',
    'Bapak Aang Alim Murtopo, M.Kom. selaku Ketua Program Studi S1 Teknik Informatika.',
    'Seluruh pihak yang telah membantu dalam penyelesaian laporan ini.',
])

body_text('Semoga laporan ini dapat memberikan manfaat bagi pembaca dan pengembangan ilmu pengetahuan di bidang Teknik Informatika.')

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Tegal, 25 Juni 2026')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('......................................')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Penulis')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Eko Saputro')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('NIM. 23215050')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

add_roman_page_number()
page_break()

# ============================================================
# DAFTAR ISI (manually laid out)
# ============================================================
center_text('DAFTAR ISI', True, 14, 24)

toc_entries = [
    ('BAB I PENDAHULUAN', '', True),
    ('1.1 Latar Belakang', '1', False),
    ('1.2 Tujuan PKL', '3', False),
    ('  1.2.1 Tujuan Umum', '3', False),
    ('  1.2.2 Tujuan Khusus', '3', False),
    ('1.3 Manfaat PKL', '4', False),
    ('BAB II GAMBARAN UMUM INSTANSI', '', True),
    ('2.1 Sejarah Perkembangan Perusahaan', '5', False),
    ('2.2 Visi, Misi, dan Tujuan', '6', False),
    ('2.3 Struktur Organisasi', '7', False),
    ('2.4 Job Deskripsi', '8', False),
    ('BAB III METODE PELAKSANAAN PKL', '', True),
    ('3.1 Tugas Umum', '10', False),
    ('3.2 Tugas Khusus', '11', False),
    ('  3.2.1 Analisis Kebutuhan Sistem', '11', False),
    ('  3.2.2 Perancangan Arsitektur Sistem', '12', False),
    ('3.3 Analisis Permasalahan dan Solusi', '13', False),
    ('BAB IV HASIL YANG DICAPAI', '', True),
    ('4.1 Gambaran Umum Sistem', '16', False),
    ('4.2 Analisis dan Perancangan Sistem', '17', False),
    ('  4.2.1 Use Case Diagram', '18', False),
    ('  4.2.2 Activity Diagram Upload Dokumen', '19', False),
    ('  4.2.3 Activity Diagram Proses Chat RAG', '20', False),
    ('  4.2.4 Entity Relationship Diagram', '21', False),
    ('  4.2.5 Arsitektur Sistem', '23', False),
    ('  4.2.6 Arsitektur RAG Pipeline', '25', False),
    ('  4.2.7 Arsitektur CRM Pipeline', '27', False),
    ('4.3 Implementasi Sistem', '28', False),
    ('  4.3.1 Spesifikasi Perangkat Keras dan Lunak', '28', False),
    ('  4.3.2 Implementasi RAG Pipeline', '30', False),
    ('  4.3.3 Implementasi Frontend', '33', False),
    ('  4.3.4 Implementasi Backend', '34', False),
    ('4.4 Pengujian Sistem', '36', False),
    ('BAB V PENUTUP', '', True),
    ('5.1 Kesimpulan', '40', False),
    ('5.2 Saran', '41', False),
    ('DAFTAR PUSTAKA', '43', True),
    ('LAMPIRAN', '46', True),
    ('  Lampiran A: Logbook Kegiatan Harian', '46', False),
    ('  Lampiran B: GitHub Repository', '48', False),
]

for title, page, is_bold in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(0)
    
    # Title
    r = p.add_run(title)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = is_bold
    
    if page:
        # Dots and page number
        dots = '.' * (60 - len(title))
        r2 = p.add_run(f' {dots} ')
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        
        r3 = p.add_run(page)
        r3.font.name = 'Times New Roman'
        r3.font.size = Pt(12)

# DAFTAR GAMBAR
doc.add_paragraph()
center_text('DAFTAR GAMBAR', True, 14, 12)

figures = [
    ('Gambar 4.2', 'Use Case Diagram Sistem Mimotes AI', '18'),
    ('Gambar 4.3', 'Activity Diagram Upload Dokumen', '19'),
    ('Gambar 4.4', 'Activity Diagram Proses Chat RAG', '20'),
    ('Gambar 4.5', 'Entity Relationship Diagram', '21'),
    ('Gambar 4.6', 'Arsitektur Sistem Mimotes AI', '23'),
    ('Gambar 4.7', 'Arsitektur RAG Pipeline', '25'),
    ('Gambar 4.8', 'Arsitektur CRM Pipeline', '27'),
]

for fig_num, fig_name, page in figures:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    r = p.add_run(f'{fig_num} {fig_name}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    dots = '.' * (60 - len(f'{fig_num} {fig_name}'))
    r2 = p.add_run(f' {dots} ')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r3 = p.add_run(page)
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)

# DAFTAR TABEL
doc.add_paragraph()
center_text('DAFTAR TABEL', True, 14, 12)

tables_list = [
    ('Tabel 4.1', 'Spesifikasi Perangkat Keras dan Lunak', '28'),
    ('Tabel 4.2', 'Hasil Pengujian Black Box', '36'),
]

for tbl_num, tbl_name, page in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    r = p.add_run(f'{tbl_num} {tbl_name}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    dots = '.' * (60 - len(f'{tbl_num} {tbl_name}'))
    r2 = p.add_run(f' {dots} ')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r3 = p.add_run(page)
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)

add_roman_page_number()
page_break()

# ============================================================
# BAB I - PENDAHULUAN (Arabic numbering starts)
# ============================================================
# Add new section with arabic page numbers
section = doc.add_section()
section.top_margin = Cm(4)
section.bottom_margin = Cm(3)
section.left_margin = Cm(4)
section.right_margin = Cm(3)

center_text('BAB I', True, 14, 6)
center_text('PENDAHULUAN', True, 14, 24)

add_heading('1.1 Latar Belakang', 2)

body_text('Perkembangan teknologi kecerdasan buatan (Artificial Intelligence/AI) dalam beberapa tahun terakhir telah mengalami kemajuan yang sangat pesat, khususnya dalam bidang pemrosesan bahasa alami (Natural Language Processing/NLP). Salah satu bentuk penerapan AI yang paling banyak digunakan dalam dunia bisnis adalah chatbot, yaitu program komputer yang mampu melakukan percakapan dengan pengguna secara otomatis menggunakan bahasa manusia [1].')

body_text('Dalam konteks layanan pelanggan (customer service), chatbot berbasis AI memiliki potensi besar untuk meningkatkan efisiensi dan kualitas layanan. Perusahaan dapat melayani pelanggan secara 24 jam tanpa henti, mengurangi waktu tunggu respons, dan menangani banyak permintaan secara bersamaan. Namun, chatbot konvensional yang hanya mengandalkan pola respons tetap (rule-based) memiliki keterbatasan dalam memahami konteks dan memberikan jawaban yang relevan terhadap pertanyaan spesifik yang berkaitan dengan pengetahuan internal perusahaan [2].')

body_text('Untuk mengatasi keterbatasan tersebut, diperlukan pendekatan yang lebih canggih, yaitu Retrieval-Augmented Generation (RAG). RAG merupakan teknik yang menggabungkan kemampuan Large Language Model (LLM) dalam menghasilkan teks yang koheren dengan mekanisme pencarian dan pengambilan informasi (retrieval) dari basis pengetahuan (knowledge base) yang telah disiapkan. Dengan pendekatan ini, chatbot tidak hanya mengandalkan pengetahuan umum yang dimiliki oleh LLM, tetapi juga mampu mengambil informasi spesifik dari dokumen-dokumen internal perusahaan, sehingga jawaban yang dihasilkan lebih akurat dan kontekstual [3].')

body_text('Selain aspek layanan pelanggan, aspek pengelolaan hubungan pelanggan (Customer Relationship Management/CRM) juga memegang peranan penting dalam kesuksesan suatu bisnis. CRM mencakup proses pengelolaan data pelanggan, pelacakan aktivitas penjualan, manajemen leads (calon pelanggan), serta analisis pola interaksi pelanggan. Integrasi antara chatbot AI dengan sistem CRM memungkinkan setiap interaksi pelanggan melalui chatbot tercatat dan terkelola secara otomatis, sehingga tim penjualan dan pemasaran dapat merespons dengan lebih tepat sasaran [4].')

body_text('Berdasarkan latar belakang permasalahan di atas, penulis mengembangkan suatu sistem yang diberi nama Mimotes AI, yaitu sistem chatbot AI berbasis pengetahuan dengan arsitektur Retrieval-Augmented Generation dan pipeline CRM untuk optimalisasi layanan pelanggan. Sistem ini dirancang untuk mampu: (1) mengelola basis pengetahuan perusahaan melalui upload dan pemrosesan dokumen, (2) memberikan respons yang akurat dan kontekstual berbasis RAG, (3) mengelola data pelanggan dan aktivitas penjualan melalui pipeline CRM, serta (4) terintegrasi dengan platform pesan instan WhatsApp melalui protokol Baileys. Sistem ini dikembangkan dalam rangka Praktik Kerja Lapangan di Bank Mandiri KCP Tegal Sudirman.')

add_heading('1.2 Tujuan PKL', 2)

add_heading('1.2.1 Tujuan Umum', 3)
body_text('Tujuan umum dari pelaksanaan Praktik Kerja Lapangan ini adalah untuk memberikan pengalaman kerja kepada mahasiswa dalam rangka menerapkan teori dan pengetahuan yang telah diterimanya di dalam perkuliahan dengan situasi nyata di tempat PKL sesuai dengan bidang kompetensi yang ada pada Program Studi S1 Teknik Informatika, khususnya dalam bidang kecerdasan buatan, rekayasa perangkat lunak, dan basis data [5].')

add_heading('1.2.2 Tujuan Khusus', 3)
add_bullet_list([
    'Merancang arsitektur sistem chatbot AI berbasis RAG yang mampu mengelola basis pengetahuan dan memberikan respons kontekstual.',
    'Mengimplementasikan pipeline RAG yang meliputi tahapan document processing, chunking, embedding, vector storage, retrieval, dan generation.',
    'Mengembangkan pipeline CRM yang mencakup manajemen leads, kontak, aktivitas, dan sales pipeline.',
    'Mengintegrasikan sistem dengan platform WhatsApp melalui protokol Baileys untuk komunikasi pelanggan secara langsung.',
    'Menguji kefektifan sistem melalui pengujian black box dan evaluasi kualitas respons chatbot.',
])

add_heading('1.3 Manfaat PKL', 2)
body_text('Manfaat bagi Universitas Harkat Negeri: Menambah jaringan kerja sama antara universitas dengan dunia industri, serta menjadi bukti kontribusi nyata program studi dalam pengembangan solusi teknologi informasi yang bermanfaat bagi masyarakat.')
body_text('Manfaat bagi Program Studi S1 Teknik Informatika: Menjadi bahan masukan untuk pengembangan kurikulum agar lebih sesuai dengan kebutuhan industri, khususnya dalam bidang kecerdasan buatan dan pengembangan perangkat lunak modern.')
body_text('Manfaat bagi Bank Mandiri KCP Tegal Sudirman: Mendapatkan solusi sistem informasi berbasis AI yang dapat membantu optimalisasi layanan pelanggan dan pengelolaan data pelanggan secara terstruktur dan efisien.')
body_text('Manfaat bagi Penulis: Memperoleh pengalaman kerja nyata dalam pengembangan sistem AI dan CRM, meningkatkan kemampuan analisis kebutuhan sistem, serta mendapatkan data yang dapat dikembangkan lebih lanjut dalam Skripsi.')

add_page_number()
page_break()

# ============================================================
# BAB II
# ============================================================
center_text('BAB II', True, 14, 6)
center_text('GAMBARAN UMUM INSTANSI', True, 14, 24)

add_heading('2.1 Sejarah Perkembangan Perusahaan', 2)
body_text('Bank Mandiri KCP Tegal Sudirman merupakan salah satu Kantor Cabang Pembantu (KCP) dari PT Bank Mandiri (Persero) Tbk yang terletak di Jalan Sudirman, Tegal. Bank Mandiri sendiri didirikan pada tanggal 2 Oktober 1998 sebagai bagian dari program restrukturisasi perbankan nasional, yang menggabungkan empat bank pemerintah yaitu Bank Bumi Daya, Bank Dagang Negara, Bank Exim, dan Bapindo.')
body_text('KCP Tegal Sudirman melayani berbagai produk dan layanan perbankan meliputi simpanan pinjaman, layanan transaksi, pengelolaan aset, serta layanan lelang property. Seiring perkembangan teknologi digital, kantor ini terus berupaya meningkatkan kualitas layanan nasabah melalui pemanfaatan sistem informasi dan teknologi terkini.')

add_heading('2.2 Visi, Misi, dan Tujuan', 2)
body_text('Visi: Menjadi bank terbaik yang mengedepankan kepuasan nasabah melalui layanan perbankan digital yang inovatif dan terpercaya.', True, False)

body_text('Misi:', True, False)
add_bullet_list([
    'Memberikan layanan perbankan digital yang inovatif dan terpercaya.',
    'Mengoptimalkan layanan pelanggan melalui pemanfaatan teknologi informasi.',
    'Mengelola data nasabah secara terstruktur dan efisien.',
])

add_heading('2.3 Struktur Organisasi', 2)
body_text('Struktur organisasi Bank Mandiri KCP Tegal Sudirman terdiri dari beberapa bagian utama yang saling berkoordinasi dalam menjalankan kegiatan operasional perbankan, meliputi Kepala Cabang Pembantu, Bagian Customer Service, Bagian Teller, Bagian Administrasi, serta Bagian Teknologi Informasi.')
body_text('[Gambar 2.1 Struktur Organisasi Bank Mandiri KCP Tegal Sudirman]')

add_heading('2.4 Job Deskripsi', 2)
body_text('Selama melaksanakan Praktik Kerja Lapangan di Bank Mandiri KCP Tegal Sudirman, penulis ditempatkan pada Divisi Teknologi Informasi dengan job deskripsi sebagai berikut:')
add_bullet_list([
    'Input data PPAT (Pejabat Pembuat Akta Tanah) dan pembuatan SKPT (Surat Keterangan Pendaftaran Tanah) untuk keperluan lelang property.',
    'Validasi pajak PPH di Kantor Pajak Pratama Tegal.',
    'Pelaksanaan dan pengelolaan lelang property.',
    'Pengembangan sistem chatbot AI (Mimotes) sebagai capstone project untuk optimalisasi layanan pelanggan.',
    'Pemasangan banner lelang di lokasi dan penyerahan risalah lelang.',
])

add_page_number()
page_break()

# ============================================================
# BAB III
# ============================================================
center_text('BAB III', True, 14, 6)
center_text('METODE PELAKSANAAN PKL', True, 14, 24)

add_heading('3.1 Tugas Umum', 2)
body_text('Selama melaksanakan Praktik Kerja Lapangan di Bank Mandiri KCP Tegal Sudirman, penulis menjalankan beberapa tugas umum yang berkaitan dengan operasional perbankan dan pengembangan sistem informasi. Tugas-tugas umum tersebut meliputi input data PPAT, pembuatan SKPT untuk keperluan lelang property, validasi pajak, pelaksanaan lelang, serta pengembangan aplikasi chatbot AI secara paralel melalui WFH (Work From Home) dan WFA (Work From Anywhere).')

add_heading('3.2 Tugas Khusus', 2)

add_heading('3.2.1 Analisis Kebutuhan Sistem', 3)
body_text('Tugas khusus pertama yang dilakukan penulis adalah analisis kebutuhan sistem (requirements analysis). Penulis mengidentifikasi permasalahan utama yang dihadapi, yaitu:')
add_bullet_list([
    'Volume interaksi nasabah yang tinggi: Tim customer service harus menangani banyak permintaan nasabah secara simultan, yang menyebabkan waktu tunggu respons yang lama.',
    'Pengetahuan yang tersebar: Informasi produk, layanan, dan prosedur perusahaan tersimpan dalam berbagai dokumen yang tidak terpusat.',
    'Tidak adanya integrasi CRM: Data nasabah dan riwayat interaksi belum terkelola secara terstruktur.',
    'Keterbatasan akses multi-platform: Nasabah hanya dapat mengakses layanan melalui beberapa saluran terbatas.',
])

add_heading('3.2.2 Perancangan Arsitektur Sistem', 3)
body_text('Penulis merancang arsitektur sistem Mimotes AI dengan mempertimbangkan aspek modularitas, scalability, dan maintainability. Arsitektur sistem terdiri dari lima komponen utama:')
add_bullet_list([
    'Frontend Layer: Next.js 16 dengan React 19 dan Tailwind CSS.',
    'API Layer: 108 API routes yang dikelompokkan berdasarkan domain fungsional.',
    'Data Layer: PostgreSQL 16 dengan ekstensi pgvector untuk data relasional dan embedding vektor.',
    'RAG Pipeline: Document parsing, chunking, embedding, vector storage, retrieval, dan generation.',
    'Microservices: Baileys (WhatsApp), PaddleOCR (OCR), dan n8n (workflow automation).',
])

add_heading('3.3 Analisis Permasalahan dan Solusi', 2)

body_text('Permasalahan 1: Akurasi Respons Chatbot. Chatbot konvensional sering memberikan jawaban yang tidak relevan karena hanya mengandalkan pengetahuan umum dari LLM. Solusi: Mengimplementasikan arsitektur RAG yang menggabungkan LLM dengan mekanisme pencarian semantik dari basis pengetahuan perusahaan.', False)

body_text('Permasalahan 2: Pengelolaan Pelanggan yang Tidak Terstruktur. Interaksi pelanggan melalui chatbot tidak tercatat sebagai data CRM. Solusi: Mengintegrasikan pipeline CRM langsung ke dalam percakapan chatbot dengan pendekatan conversation-centric CRM.', False)

body_text('Permasalahan 3: Isolasi Data Multi-Tenant. Sistem harus mendukung multiple workspace tanpa kebocoran data. Solusi: Workspace-based isolation dengan PostgreSQL Row Level Security (RLS).', False)

body_text('Permasalahan 4: Ekstraksi Teks dari Dokumen Gambar. Dokumen berupa gambar tidak dapat diproses oleh text parser biasa. Solusi: Integrasi PaddleOCR dan Google Gemini Vision untuk OCR dan captioning.', False)

add_page_number()
page_break()

# ============================================================
# BAB IV
# ============================================================
center_text('BAB IV', True, 14, 6)
center_text('HASIL YANG DICAPAI', True, 14, 24)

add_heading('4.1 Gambaran Umum Sistem', 2)
body_text('Mimotes AI merupakan sistem chatbot AI berbasis pengetahuan yang dirancang untuk optimalisasi layanan pelanggan. Sistem ini mengintegrasikan tiga komponen utama: (1) pipeline RAG untuk pemrosesan dokumen dan generasi respons berbasis pengetahuan, (2) pipeline CRM untuk pengelolaan leads dan aktivitas penjualan, serta (3) integrasi WhatsApp untuk komunikasi pelanggan multi-platform.')
body_text('Sistem ini di-deploy menggunakan Docker Compose dengan arsitektur microservices yang terdiri dari lima service utama: PostgreSQL (port 5432), Next.js Application (port 3100), PaddleOCR (port 8090), Baileys WhatsApp (port 3002), dan Database Migration (one-shot).')

add_heading('4.2 Analisis dan Perancangan Sistem', 2)

add_heading('4.2.1 Use Case Diagram', 3)
body_text('Sistem Mimotes AI memiliki empat aktor utama: Admin/Workspace Owner yang dapat mengelola seluruh aspek sistem, Editor yang dapat mengupload dokumen dan mengelola chat, Viewer yang dapat melihat dashboard dan menggunakan chat, serta Visitor/Pelanggan yang berinteraksi melalui widget chat atau WhatsApp.')
add_image_centered('usecase.png', 'Gambar 4.2 Use Case Diagram Sistem Mimotes AI')

add_heading('4.2.2 Activity Diagram Upload Dokumen', 3)
body_text('Proses upload dokumen melalui pipeline RAG mengikuti alur: autentikasi dan otorisasi, validasi file (tipe dan ukuran), penyimpanan file ke disk, pembuatan record dokumen di database, pemrosesan asynchronous melalui processing queue, hingga pipeline parsing, chunking, embedding, dan storage.')
add_image_centered('activity-upload.png', 'Gambar 4.3 Activity Diagram Upload Dokumen')

add_heading('4.2.3 Activity Diagram Proses Chat RAG', 3)
body_text('Proses chat RAG dimulai dari penerimaan pesan pengguna, pembuatan query embedding, hybrid search (vector similarity + BM25 dengan Reciprocal Rank Fusion), klasifikasi confidence, pembangunan context, konstruksi system prompt, generasi respons oleh LLM secara streaming, hingga penyimpanan respons dan sources di database.')
add_image_centered('activity-chat.png', 'Gambar 4.4 Activity Diagram Proses Chat RAG')

add_heading('4.2.4 Entity Relationship Diagram', 3)
body_text('Sistem Mimotes AI menggunakan 28 model database yang terorganisir dalam enam domain fungsional: Tenant & Identity (User, Workspace, WorkspaceMember), Document & RAG Pipeline (Document, DocumentChunk dengan embedding vector 1536-dimensi), Conversational AI (ChatSession, Widget, WhatsApp), CRM & Lead Pipeline (data leads terintegrasi dalam model percakapan), Subscription & Billing (Stripe integration), dan Configuration & Observability.')
add_image_centered('erd.png', 'Gambar 4.5 Entity Relationship Diagram')

add_heading('4.2.5 Arsitektur Sistem', 3)
body_text('Arsitektur sistem Mimotes AI terdiri dari enam layer: Frontend (Next.js 16 + React 19 + Tailwind CSS), API Layer (108 routes), RAG Pipeline, Data Layer (PostgreSQL 16 + pgvector), Microservices (Baileys, PaddleOCR, n8n), dan External Services (LLM Provider, Stripe, Redis).')
add_image_centered('architecture.png', 'Gambar 4.6 Arsitektur Sistem Mimotes AI')

add_heading('4.2.6 Arsitektur RAG Pipeline', 3)
body_text('Pipeline RAG terdiri dari enam komponen: Document Ingestion (Parser untuk PDF/DOCX/TXT/gambar, Chunker dengan 500 chars dan 50 word overlap), Embedding Generation (OpenAI text-embedding-3-small atau Feature Hashing lokal, 1536 dimensi), Vector Storage (pgvector PostgreSQL), Retrieval (Hybrid Search: Vector 0.6 + BM25 0.4 dengan RRF), dan Response Generation (LLM dengan temperature 0.3, streaming response).')
add_image_centered('rag-pipeline.png', 'Gambar 4.7 Arsitektur RAG Pipeline')

add_heading('4.2.7 Arsitektur CRM Pipeline', 3)
body_text('Pipeline CRM mengadopsi pendekatan conversation-centric di mana setiap percakapan publik (widget atau WhatsApp) secara otomatis berfungsi sebagai lead record. Alur CRM meliputi: Lead Capture dari Widget dan WhatsApp, AI Analysis (intent, interest, budget, timeline), Lead Scoring (low/medium/high), Lead Status Tracking (new, contacted, qualified, converted), dan Follow-up Automation melalui multi-channel notification (email, Telegram, Discord).')
add_image_centered('crm-pipeline.png', 'Gambar 4.8 Arsitektur CRM Pipeline')

add_heading('4.3 Implementasi Sistem', 2)

add_heading('4.3.1 Spesifikasi Perangkat Keras dan Lunak', 3)
add_table_simple(
    ['Komponen', 'Spesifikasi'],
    [
        ['Framework', 'Next.js 16.2.7 (App Router)'],
        ['Frontend', 'React 19, Tailwind CSS, Framer Motion, Recharts'],
        ['Database', 'PostgreSQL 16 + pgvector extension'],
        ['ORM', 'Prisma 6.19.3'],
        ['AI Provider', 'Mimo Pro (custom), OpenAI, Google Gemini'],
        ['OCR Engine', 'PaddleOCR (Python) + Gemini Vision'],
        ['WhatsApp', 'Baileys 6.7.23'],
        ['Container', 'Docker Compose (multi-stage builds)'],
        ['Testing', 'Vitest'],
    ]
)

add_heading('4.3.2 Implementasi RAG Pipeline', 3)

body_text('Document Parsing: Modul parsing mendukung ekstraksi teks dari berbagai format: PDF (pdf-parse), DOCX (mammoth), TXT, CSV, XLSX, gambar (PaddleOCR + Gemini Vision), dan URL (cheerio HTML parsing).')
body_text('Chunking Strategy: Algoritma chunking menggunakan pendekatan recursive paragraph-then-sentence dengan chunkSize 500 karakter, overlap 50 kata, dan batas maksimum 1000 chunks per dokumen.')
body_text('Embedding Generation: Dua provider tersedia: OpenAI text-embedding-3-small (1536 dimensi, $0.02/M tokens) dan Feature Hashing lokal (gratis, sebagai fallback ketika API key tidak tersedia).')
body_text('Vector Storage: pgvector (PostgreSQL) menyimpan embedding di tabel document_chunks sebagai vector(1536). Proses storage dilakukan dalam batch 50 chunks per transaksi dengan RLS via set_config().')
body_text('Retrieval: Dua mode pencarian tersedia. Mode Vector-Only menggunakan cosine distance dengan threshold 0.30. Mode Hybrid Search (default) menggabungkan Vector 0.6 + BM25 0.4 dengan Reciprocal Rank Fusion. Confidence classification: High (>= 0.55), Medium (>= 0.40), Low (>= 0.30), Refuse (< 0.30).')
body_text('Response Generation: LLM diintegrasikan melalui OpenAI-compatible chat completions API dengan parameter temperature 0.3, max_tokens 1000, dan streaming response. Tiga mode widget: Knowledge Base (strict, cite sources), Customer Service (natural, conversational), dan Sales Agent (conversion-focused).')

add_heading('4.3.3 Implementasi Frontend', 3)
body_text('Frontend Mimotes AI dibangun menggunakan Next.js 16 dengan React 19 dan Tailwind CSS. Antarmuka terdiri dari 53 halaman yang terorganisir dalam domain: dashboard, chat, knowledge management (documents, chunks, images, search, sources), CRM (leads), analytics (chat, cost, leads, usage), settings (16 sub-halaman), widget preview, admin users, dan WhatsApp.')

add_heading('4.3.4 Implementasi Backend', 3)
body_text('Backend menggunakan 108 API routes yang dikelompokkan berdasarkan domain: Auth (2), Admin (4), AI (6), Analytics (8), Chat (2), Documents (3), Knowledge/RAG (8), Leads/CRM (5), WhatsApp (8), Widget (7), Billing (3), Workspace (10), dan lainnya (34). Setiap API route menggunakan Prisma ORM untuk interaksi database dengan PostgreSQL.')

add_heading('4.4 Pengujian Sistem', 2)
body_text('Pengujian black box dilakukan terhadap seluruh fitur utama sistem untuk memastikan setiap fungsi bekerja sesuai kebutuhan. Berikut adalah hasil pengujian:')

add_table_simple(
    ['No', 'Fitur', 'Input', 'Expected Output', 'Status'],
    [
        ['1', 'Login', 'Email + password valid', 'Redirect ke dashboard', 'Lulus'],
        ['2', 'Login', 'Email + password invalid', 'Pesan error', 'Lulus'],
        ['3', 'Upload Dokumen', 'File PDF < 10MB', 'Status processing -> ready', 'Lulus'],
        ['4', 'Upload Gambar', 'File PNG < 10MB', 'OCR + caption -> ready', 'Lulus'],
        ['5', 'Chat RAG', 'Pertanyaan dari dokumen', 'Jawaban akurat + sumber', 'Lulus'],
        ['6', 'Chat RAG', 'Pertanyaan di luar konteks', 'Penolakan sopan', 'Lulus'],
        ['7', 'Lead Capture', 'Visitor mengisi form', 'Lead tersimpan di DB', 'Lulus'],
        ['8', 'WhatsApp', 'Pesan masuk', 'Auto-reply + RAG response', 'Lulus'],
        ['9', 'Settings', 'Ubah AI provider', 'Provider tersimpan', 'Lulus'],
        ['10', 'RBAC', 'Viewer coba upload', 'Access denied', 'Lulus'],
    ]
)

add_page_number()
page_break()

# ============================================================
# BAB V
# ============================================================
center_text('BAB V', True, 14, 6)
center_text('PENUTUP', True, 14, 24)

add_heading('5.1 Kesimpulan', 2)
body_text('Berdasarkan kegiatan Praktik Kerja Lapangan yang telah dilaksanakan di Bank Mandiri KCP Tegal Sudirman, dapat disimpulkan beberapa hal sebagai berikut:')
add_bullet_list([
    'Sistem Mimotes AI berhasil dirancang dan diimplementasikan sebagai solusi chatbot AI berbasis pengetahuan dengan arsitektur RAG dan pipeline CRM untuk optimalisasi layanan pelanggan.',
    'Pipeline RAG berfungsi dengan baik dalam mengelola alur data dari upload dokumen, chunking, embedding, vector storage, retrieval, hingga generasi respons oleh LLM.',
    'Pipeline CRM terintegrasi dengan sistem chatbot sehingga setiap interaksi pelanggan dapat tercatat dan terkelola secara otomatis.',
    'Integrasi WhatsApp melalui protokol Baileys memungkinkan pelanggan berinteraksi melalui platform pesan instan yang populer.',
    'Hasil pengujian black box menunjukkan bahwa seluruh fitur sistem berfungsi sesuai kebutuhan.',
])

add_heading('5.2 Saran', 2)
body_text('Bagi Bank Mandiri KCP Tegal Sudirman: Sistem Mimotes AI yang telah dikembangkan perlu terus dipelihara dan ditingkatkan fiturnya sesuai dengan kebutuhan bisnis yang berkembang. Dilakukan evaluasi berkala terhadap kualitas respons chatbot dan pipeline CRM. Pertimbangkan untuk mengintegrasikan sistem dengan platform lain seperti email, media sosial, atau sistem CRM eksternal.')
body_text('Bagi Program Studi S1 Teknik Informatika Universitas Harkat Negeri: Kurikulum dapat diperkaya dengan mata kuliah atau materi yang berkaitan dengan kecerdasan buatan, khususnya Natural Language Processing dan Retrieval-Augmented Generation. Praktik Kerja Lapangan dengan skema Capstone Project dapat dijadikan model pembelajaran berbasis proyek yang efektif.')
body_text('Bagi Penulis: Sistem Mimotes AI dapat dikembangkan lebih lanjut sebagai topik Skripsi, khususnya dalam aspek optimasi model embedding, evaluasi kualitas RAG, atau pengembangan fitur analytics. Perlu dilakukan penelitian evaluasi kualitas respons chatbot menggunakan metode seperti ROUGE, BLEU, atau human evaluation.')

add_page_number()
page_break()

# ============================================================
# DAFTAR PUSTAKA
# ============================================================
center_text('DAFTAR PUSTAKA', True, 14, 24)

refs = [
    '[1] H. Susanto, A. Wibowo, dan B. Satria, "Pemanfaatan Chatbot Berbasis Kecerdasan Buatan untuk Optimalisasi Layanan Pelanggan: Systematic Literature Review," Jurnal Sistem Informasi, vol. 19, no. 2, pp. 45-62, 2023.',
    '[2] T. Brown et al., "Language Models are Few-Shot Learners," Advances in Neural Information Processing Systems, vol. 33, pp. 1877-1901, 2020.',
    '[3] P. Lewis et al., "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks," Advances in Neural Information Processing Systems, vol. 33, pp. 9459-9474, 2020.',
    '[4] F. Buttle dan S. Maklan, Customer Relationship Management: Concepts and Technologies, 4th ed. London: Routledge, 2019.',
    '[5] Universitas Harkat Negeri, Buku Panduan Praktik Kerja Lapangan Program Studi S1 Teknik Informatika. Tegal: Fakultas Sains dan Teknologi, 2026.',
    '[6] V. Lopez, S. Vembu, dan R. Pinheriro, "A Survey on Transfer Learning in Natural Language Processing," Artificial Intelligence Review, vol. 53, no. 4, pp. 2339-2367, 2020.',
    '[7] Y. Gao et al., "Retrieval-Augmented Generation for Large Language Models: A Survey," arXiv preprint arXiv:2312.10997, 2024.',
    '[8] A. Vaswani et al., "Attention Is All You Need," Advances in Neural Information Processing Systems, vol. 30, pp. 5998-6008, 2017.',
    '[9] J. Devlin, M.-W. Chang, K. Lee, dan K. Toutanova, "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding," in Proceedings of NAACL-HLT, pp. 4171-4186, 2019.',
    '[10] Qdrant, "Qdrant - Vector Database for AI Applications," Qdrant Documentation, 2024. [Online]. Available: https://qdrant.tech/documentation/',
    '[11] Next.js, "Next.js by Vercel - The React Framework," Next.js Documentation, 2024. [Online]. Available: https://nextjs.org/docs',
    '[12] Prisma, "Prisma - Next-generation ORM for Node.js and TypeScript," Prisma Documentation, 2024. [Online]. Available: https://www.prisma.io/docs',
    '[13] WhatsApp Baileys, "Baileys - WhatsApp Web API," GitHub Repository, 2024. [Online]. Available: https://github.com/WhiskeySockets/Baileys',
    '[14] pgvector, "pgvector: Open-source vector similarity search for Postgres," GitHub Repository, 2024. [Online]. Available: https://github.com/pgvector/pgvector',
    '[15] OpenAI, "OpenAI API Documentation - Embeddings," OpenAI Platform, 2024. [Online]. Available: https://platform.openai.com/docs/guides/embeddings',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

add_page_number()
page_break()

# ============================================================
# LAMPIRAN
# ============================================================
center_text('LAMPIRAN', True, 14, 24)

add_heading('Lampiran A: Logbook Kegiatan Harian', 2)

log_data = [
    ('No', 'Tanggal', 'Jam', 'Aktivitas', 'Lokasi'),
    ('1', '25/05/2026', '09:10-15:30', 'Input Akun Lelang, Validasi Pajak PPH', 'Bank Mandiri'),
    ('2', '26/05/2026', '09:10-15:02', 'Pelaksanaan Lelang', 'Bank Mandiri'),
    ('3', '27/05/2026', '09:11-15:30', 'WFH Pembuatan Aplikasi', 'WFH'),
    ('4', '28/05/2026', '09:10-15:30', 'Melanjutkan Pembuatan Aplikasi', 'WFH'),
    ('5', '29/05/2026', '09:00-14:11', 'Input Pendataan SSPD BPHTB', 'Bank Mandiri'),
    ('6', '01/06/2026', '09:38-22:17', 'WFH Developing Aplikasi', 'WFH'),
    ('7', '02/06/2026', '09:10-14:58', 'Penyerahan Risalah Lelang', 'Bank Mandiri'),
    ('8', '03/06/2026', '09:10-17:28', 'Pembuatan Surat Pendaftaran Tanah', 'WFA'),
    ('9', '04/06/2026', '09:10-16:28', 'Input PPAT ke Sistem INTAN', 'Bank Mandiri'),
    ('10', '05/06/2026', '08:48-15:39', 'Input Data PPAT SKPT', 'Bank Mandiri'),
    ('11', '08/06/2026', '08:30-16:13', 'Input PPATK + Banner Lelang', 'Bank Mandiri'),
    ('12', '09/06/2026', '08:46-16:41', 'Pembuatan SKPT + Banner Lelang', 'Bank Mandiri'),
    ('13', '10/06/2026', '09:22-16:20', 'Input Data PPAT untuk SKPT', 'Bank Mandiri'),
    ('14', '11/06/2026', '08:30-17:40', 'Input PPAT + SKPT + Banner', 'Bank Mandiri'),
    ('15', '12/06/2026', '08:45-15:58', 'Input Data PPAT untuk SKPT', 'Bank Mandiri'),
    ('16', '15/06/2026', '08:36-16:32', 'Input Data PPAT + SKPT', 'Bank Mandiri'),
    ('17', '16/06/2026', '09:20-16:55', 'WFA Pembuatan Aplikasi', 'WFA'),
    ('18', '17/06/2026', '08:54-16:34', 'Input PPAT Membuat SKPT', 'Bank Mandiri'),
    ('19', '18/06/2026', '08:38-16:30', 'Input Data PPAT untuk SKPT', 'WFA'),
    ('20', '19/06/2026', '08:45-15:55', 'Input Data PPAT untuk SKPT', 'Bank Mandiri'),
    ('21', '22/06/2026', '08:47-16:50', 'Input PPAT', 'Bank Mandiri'),
    ('22', '23/06/2026', '08:35-15:02', 'SKPT + Kunjungi Lokasi Lelang', 'Bank Mandiri'),
    ('23', '24/06/2026', '10:52-16:21', 'WFH Pembuatan Aplikasi', 'WFH'),
    ('24', '25/06/2026', '08:35-15:56', 'Input Data PPAT', 'Bank Mandiri'),
]

add_table_simple(
    log_data[0],
    log_data[1:]
)

body_text('Pembimbing Lapangan: Widianto Agung Nugroho', True, False)

add_heading('Lampiran B: GitHub Repository', 2)
body_text('Repository: https://github.com/EkoSaputro14/mimotes.git', False, False)
body_text('Branch: semi-final', False, False)

add_heading('Lampiran C: Diagram Pendukung', 2)
body_text('Seluruh diagram pendukung telah disertakan dalam BAB IV laporan ini, meliputi: Use Case Diagram (Gambar 4.2), Activity Diagram Upload Dokumen (Gambar 4.3), Activity Diagram Proses Chat RAG (Gambar 4.4), Entity Relationship Diagram (Gambar 4.5), Arsitektur Sistem (Gambar 4.6), Arsitektur RAG Pipeline (Gambar 4.7), dan Arsitektur CRM Pipeline (Gambar 4.8).')

add_page_number()

# ============================================================
# SAVE
# ============================================================
output = r'C:\Users\SMANSA\mimotes\docs\pkl-report\LAPORAN_PKL_Eko_Saputro_23215050.docx'
doc.save(output)
print(f'=== DONE ===')
print(f'Saved: {output}')
