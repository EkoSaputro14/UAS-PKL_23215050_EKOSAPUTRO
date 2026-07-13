#!/usr/bin/env python3
"""
Rebuild v3: Fix multi-line heading detection, keep template formatting.
"""
import os
import copy
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

TEMPLATE = 'C:/Users/SMANSA/AppData/Local/hermes/cache/documents/doc_4ce6543a497c_LAPORAN PKL (1).docx'
SOURCE = 'LAPORAN_PKL_v14_UML.docx'
OUTPUT = 'LAPORAN_PKL_v15_Template.docx'
DIAGRAMS_DIR = 'diagrams'

NEW_TITLE = "RANCANG BANGUN SISTEM CHATBOT AI BERBASIS PENGETAHUAN DENGAN RETRIEVAL-AUGMENTED GENERATION DAN PIPELINE CRM UNTUK OPTIMALISASI LAYANAN PELANGGAN"
NEW_STUDENT = "Eko Saputro"
NEW_NIM = "23215050"
NEW_SUPERVISOR = "Widianto Agung Nugroho"
OLD_STUDENT = "Moh. Arif Prasetyo"
OLD_NIM = "23215043"

def get_heading_level(para):
    name = para.style.name if para.style else ''
    if name in ('Heading 1', 'Heading1'): return 1
    if name in ('Heading 2', 'Heading2'): return 2
    if name in ('Heading 3', 'Heading3'): return 3
    return None

def classify_bab(text):
    """Handle multi-line headings like 'BAB I\\nPENDAHULUAN'."""
    text = text.strip()
    first_line = text.split('\n')[0].strip()
    # Check first line only for BAB classification
    if 'BAB V' in first_line: return 'BAB_V'
    if 'BAB IV' in first_line: return 'BAB_IV'
    if 'BAB III' in first_line: return 'BAB_III'
    if 'BAB II' in first_line: return 'BAB_II'
    if 'BAB I' in first_line: return 'BAB_I'
    if 'DAFTAR PUSTAKA' in first_line: return 'DAFTAR_PUSTAKA'
    if 'LAMPIRAN' in first_line: return 'LAMPIRAN'
    return None

def get_bab_heading_text(text):
    """Get just the heading part (e.g. 'BAB I') from multi-line heading."""
    return text.strip().split('\n')[0].strip()

print("📄 Loading...")
tpl_doc = Document(TEMPLATE)
src_doc = Document(SOURCE)

# ===== STEP 1: Extract source sections =====
print("\n📝 Step 1: Extracting source...")
src_sections = {}
current_key = None
current_paras = []

for para in src_doc.paragraphs:
    text = para.text.strip()
    level = get_heading_level(para)
    
    if level == 1:
        if current_key:
            src_sections[current_key] = current_paras
        current_paras = []
        current_key = classify_bab(text)
        # Store the heading text for later
        if current_key:
            current_paras.append({'_heading': get_bab_heading_text(text)})
        continue
    
    if current_key:
        current_paras.append({
            'text': text,
            'style_name': para.style.name if para.style else 'Normal',
        })

if current_key:
    src_sections[current_key] = current_paras

for key, paras in src_sections.items():
    # Remove the _heading entry from count
    content_count = len([p for p in paras if '_heading' not in p])
    print(f"  {key}: {content_count} content paragraphs")

# ===== STEP 2: Find template sections =====
print("\n📝 Step 2: Finding template sections...")
tpl_sections = {}
for i, para in enumerate(tpl_doc.paragraphs):
    level = get_heading_level(para)
    if level == 1:
        key = classify_bab(para.text)
        if key:
            tpl_sections[key] = i

print(f"  {tpl_sections}")

# ===== STEP 3: Replace content =====
print("\n📝 Step 3: Replacing content...")

sorted_sections = sorted(tpl_sections.items(), key=lambda x: x[1])

for sec_i, (key, start_idx) in enumerate(sorted_sections):
    if key not in src_sections:
        print(f"  ⚠️ No source for {key}")
        continue
    
    end_idx = len(tpl_doc.paragraphs)
    for j in range(sec_i + 1, len(sorted_sections)):
        _, next_idx = sorted_sections[j]
        end_idx = next_idx
        break
    
    section_paras = list(tpl_doc.paragraphs[start_idx:end_idx])
    src_paras = src_sections[key]
    
    # Filter out _heading entries
    src_content = [p for p in src_paras if '_heading' not in p]
    
    # Skip heading paragraph + optional subheading
    content_start = 1
    if key == 'BAB_I' and len(section_paras) > 1:
        nxt = section_paras[1].text.strip()
        if 'PENDAHULUAN' in nxt or nxt == '':
            content_start = 2
    
    available = len(section_paras) - content_start
    
    # Replace existing paragraphs
    for j, src_p in enumerate(src_content):
        if j < available:
            target = section_paras[content_start + j]
            for run in target.runs:
                run.text = ''
            if target.runs:
                target.runs[0].text = src_p['text']
            elif src_p['text']:
                target.add_run(src_p['text'])
    
    # Insert excess paragraphs
    excess = len(src_content) - available
    if excess > 0 and available > 0:
        last_replaced = section_paras[content_start + available - 1]
        last_elem = last_replaced._element
        parent = last_elem.getparent()
        insert_pos = list(parent).index(last_elem) + 1
        
        # Use last paragraph as format template
        fmt_template = section_paras[-1]
        
        for k in range(excess):
            src_p = src_content[available + k]
            new_elem = copy.deepcopy(fmt_template._element)
            
            # Clear and set text
            for t_elem in new_elem.iter(qn('w:t')):
                t_elem.text = ''
            runs = new_elem.findall(f'.//{qn("w:r")}')
            if runs:
                t_elems = runs[0].findall(qn('w:t'))
                if t_elems:
                    t_elems[0].text = src_p['text']
                else:
                    from lxml import etree
                    t = etree.SubElement(runs[0], qn('w:t'))
                    t.text = src_p['text']
                    t.set(qn('xml:space'), 'preserve')
            
            parent.insert(insert_pos + k, new_elem)
        
        print(f"  ✅ {key}: {available} replaced + {excess} inserted")
    else:
        print(f"  ✅ {key}: {len(src_content)}/{available} replaced")

# ===== STEP 4: Fix front matter =====
print("\n📝 Step 4: Fixing front matter...")
for para in tpl_doc.paragraphs:
    if "SISTEM PAKAR" in para.text:
        for run in para.runs:
            if "SISTEM PAKAR" in run.text or "RANCANG BANGUN" in run.text:
                run.text = ""
        for run in para.runs:
            if not run.text.strip():
                run.text = NEW_TITLE
                break
    for run in para.runs:
        if OLD_STUDENT in run.text:
            run.text = run.text.replace(OLD_STUDENT, NEW_STUDENT)
        if OLD_NIM in run.text:
            run.text = run.text.replace(OLD_NIM, NEW_NIM)
        if "Nama Dosen Pembimbing" in run.text:
            run.text = run.text.replace("Nama Dosen Pembimbing", NEW_SUPERVISOR)
        if "NIPY." in run.text:
            run.text = run.text.replace("NIPY.", "")

# ===== STEP 5: Clean old captions =====
print("\n📝 Step 5: Cleaning old captions...")
old_captions = {
    "Pengumpulan Data": "Use Case Diagram Sistem Mimotes AI",
    "Pohon Keputusan": "Activity Diagram Upload Dokumen",
    "Activity Diagram Proses Diag": "Activity Diagram Proses Chat RAG",
    "Sequence diagram Proses": "Sequence Diagram Chat RAG",
    "Class Diagram": "ERD Domain Identity & Workspace",
    "Arsitektur Sistem SIPATAN": "ERD Domain RAG & Knowledge Base",
    "Flowchart Proses": "ERD Domain Chat & CRM",
    "Halaman Beranda": "ERD Domain Billing & Configuration",
    "Halaman Tentang": "ERD Ringkasan — Seluruh Relasi",
    "Halaman Cara Kerja": "Arsitektur Sistem Mimotes AI",
    "Halaman Deteksi": "Arsitektur RAG Pipeline",
    "Halaman Panel Admin": "Arsitektur CRM Pipeline",
    "Implementasi database": "Halaman Login",
    "Input Keyakinan": "Dashboard Admin",
    "Hasil Diag": "Upload Dokumen",
}
for para in tpl_doc.paragraphs:
    text = para.text.strip()
    if not text.startswith('Gambar 4.'):
        continue
    for old_kw, new_txt in old_captions.items():
        if old_kw in text:
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = f"Gambar {new_txt}"
            break

# ===== SAVE =====
print("\n💾 Saving...")
tpl_doc.save(OUTPUT)

# ===== STEP 6: Insert images =====
print("\n🖼️ Step 6: Inserting images...")
doc2 = Document(OUTPUT)

CAPTION_IMG = {
    "Use Case Diagram": "usecase.png",
    "Activity Diagram Upload": "activity-upload.png",
    "Activity Diagram Proses Chat": "activity-chat.png",
    "Sequence Diagram Chat RAG": "sequence-chat-rag.png",
    "ERD Domain Identity": "erd-a-identity.png",
    "ERD Domain RAG": "erd-b-rag.png",
    "ERD Domain Chat": "erd-c-chat-crm.png",
    "ERD Domain Billing": "erd-d-billing.png",
    "ERD Ringkasan": "erd-summary.png",
    "Arsitektur Sistem Mimotes": "architecture.png",
    "Arsitektur RAG": "rag-pipeline.png",
    "Arsitektur CRM": "crm-pipeline.png",
}

img_count = 0
for i, para in enumerate(doc2.paragraphs):
    text = para.text.strip()
    if not text.startswith('Gambar 4.'):
        continue
    for keyword, img_file in CAPTION_IMG.items():
        if keyword in text:
            img_path = os.path.join(DIAGRAMS_DIR, img_file)
            if not os.path.exists(img_path):
                continue
            try:
                new_para = doc2.add_paragraph()
                new_para.alignment = 1
                new_para.add_run().add_picture(img_path, width=Inches(5.5))
                parent = para._element.getparent()
                pos = list(parent).index(para._element)
                new_elem = new_para._element
                parent.remove(new_elem)
                parent.insert(pos, new_elem)
                img_count += 1
                print(f"  ✅ {img_file}")
            except Exception as e:
                print(f"  ❌ {img_file}: {e}")
            break

doc2.save(OUTPUT)

# ===== VERIFY =====
doc3 = Document(OUTPUT)
total_imgs = sum(
    len(run._element.findall(qn('w:drawing')))
    for para in doc3.paragraphs
    for run in para.runs
)
old_found = any('Sistem Pakar Deteksi Hama' in p.text for p in doc3.paragraphs)

print(f"\n📊 Final:")
print(f"  Paragraphs: {len(doc3.paragraphs)}")
print(f"  Images: {total_imgs}")
print(f"  Old content: {'⚠️ FOUND' if old_found else '✅ Clean'}")
print(f"  Size: {os.path.getsize(OUTPUT)/1024:.0f} KB")
