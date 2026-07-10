"""
Fix plain version TOC — split into steps with reload between each.
"""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

FILE = 'LAPORAN_PKL_v10_Plain_final.docx'

# ============================================
# STEP 1: Remove template sub-bab from DAFTAR ISI
# ============================================
print("=== Step 1: Remove template entries from DAFTAR ISI ===")
doc = Document(FILE)

REMOVE = [
    '5.1.1', '5.1.2', '5.2.1', '5.2.2', '5.2.3', '5.2.4', '5.2.5', '5.2.6',
    '5.3.1', '5.3.2', '5.3.3', '5.3.4', '5.3.5', '5.3.6',
    '5.4.1', '5.4.2', '5.4.3', '5.4.4', '5.4.5',
    '5.5.1', '5.5.2',
    '5.6.1', '5.6.2', '5.6.3', '5.6.4', '5.6.5', '5.6.6',
]

removed = 0
# Work backwards to avoid index shifting
to_remove = []
for i, p in enumerate(doc.paragraphs):
    if not p.style or not p.style.style_id or not p.style.style_id.startswith('TOC'):
        continue
    p_elem = p._p
    hyperlink = p_elem.find(qn('w:hyperlink'))
    if hyperlink is None:
        continue
    t_elems = hyperlink.findall('.//' + qn('w:t'))
    full = ''.join(t.text or '' for t in t_elems)
    for prefix in REMOVE:
        if full.startswith(prefix):
            to_remove.append(p_elem)
            break
    # Also check for duplicate BAB V with page 12
    if 'BAB V PENUTUP' in full and '\t12' in full:
        to_remove.append(p_elem)

for elem in to_remove:
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)
        removed += 1

doc.save(FILE)
print(f"  Removed {removed} entries")

# ============================================
# STEP 2: Fix DAFTAR GAMBAR — replace with complete list
# ============================================
print("\n=== Step 2: Fix DAFTAR GAMBAR ===")
doc = Document(FILE)

# Find DAFTAR GAMBAR heading
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.style_id == 'Heading1' and 'DAFTAR GAMBAR' in p.text:
        p_elem = p._p
        
        # Remove all non-heading paragraphs after it until next Heading 1
        next_elem = p_elem.getnext()
        while next_elem is not None:
            nxt = next_elem.getnext()
            tag = next_elem.tag.split('}')[-1]
            if tag == 'sectPr':
                break
            # Check if next element is a Heading 1
            pPr = next_elem.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None and pStyle.get(qn('w:val')) == 'Heading1':
                    break
            parent = next_elem.getparent()
            if parent is not None:
                parent.remove(next_elem)
            next_elem = nxt
        
        # Insert 15 figure entries
        figures = [
            ('Gambar 4.2 Use Case Diagram Sistem Mimotes AI', '13'),
            ('Gambar 4.3 Activity Diagram Upload Dokumen', '14'),
            ('Gambar 4.4 Activity Diagram Proses Chat RAG', '15'),
            ('Gambar 4.5 Entity Relationship Diagram (ERD)', '16'),
            ('Gambar 4.6 Arsitektur Sistem Mimotes AI', '16'),
            ('Gambar 4.7 Arsitektur RAG Pipeline', '16'),
            ('Gambar 4.8 Arsitektur CRM Pipeline', '16'),
            ('Gambar 4.9 Halaman Login', '16'),
            ('Gambar 4.10 Dashboard Admin', '16'),
            ('Gambar 4.11 Upload Dokumen', '16'),
            ('Gambar 4.12 Daftar Dokumen', '16'),
            ('Gambar 4.13 Chat AI', '16'),
            ('Gambar 4.14 Knowledge Search', '16'),
            ('Gambar 4.15 Analytics Chat', '16'),
            ('Gambar 4.16 Pengaturan AI Provider', '16'),
        ]
        
        insert_after = p_elem
        for title, page in figures:
            new_p = etree.Element(qn('w:p'))
            pPr = etree.SubElement(new_p, qn('w:pPr'))
            pStyle = etree.SubElement(pPr, qn('w:pStyle'))
            pStyle.set(qn('w:val'), 'TOC1')
            tabs = etree.SubElement(pPr, qn('w:tabs'))
            tab = etree.SubElement(tabs, qn('w:tab'))
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:leader'), 'dot')
            tab.set(qn('w:pos'), '8296')
            
            hl = etree.SubElement(new_p, qn('w:hyperlink'))
            r = etree.SubElement(hl, qn('w:r'))
            t = etree.SubElement(r, qn('w:t'))
            t.text = title
            t.set(qn('xml:space'), 'preserve')
            r_tab = etree.SubElement(hl, qn('w:r'))
            etree.SubElement(r_tab, qn('w:tab'))
            r_page = etree.SubElement(hl, qn('w:r'))
            t_page = etree.SubElement(r_page, qn('w:t'))
            t_page.text = page
            
            insert_after.addnext(new_p)
            insert_after = new_p
        
        print(f"  Inserted {len(figures)} figures")
        break

doc.save(FILE)

# ============================================
# STEP 3: Fix DAFTAR TABEL
# ============================================
print("\n=== Step 3: Fix DAFTAR TABEL ===")
doc = Document(FILE)

for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.style_id == 'Heading1' and 'DAFTAR TABEL' in p.text:
        p_elem = p._p
        
        # Remove old entries
        next_elem = p_elem.getnext()
        while next_elem is not None:
            nxt = next_elem.getnext()
            tag = next_elem.tag.split('}')[-1]
            if tag == 'sectPr':
                break
            pPr = next_elem.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None and pStyle.get(qn('w:val')) == 'Heading1':
                    break
            parent = next_elem.getparent()
            if parent is not None:
                parent.remove(next_elem)
            next_elem = nxt
        
        # Insert correct tables
        tables = [
            ('Tabel 4.1 Spesifikasi Perangkat Keras dan Lunak', '13'),
            ('Tabel 4.2 Hasil Pengujian Black Box', '14'),
        ]
        
        insert_after = p_elem
        for title, page in tables:
            new_p = etree.Element(qn('w:p'))
            pPr = etree.SubElement(new_p, qn('w:pPr'))
            pStyle = etree.SubElement(pPr, qn('w:pStyle'))
            pStyle.set(qn('w:val'), 'TOC1')
            tabs = etree.SubElement(pPr, qn('w:tabs'))
            tab = etree.SubElement(tabs, qn('w:tab'))
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:leader'), 'dot')
            tab.set(qn('w:pos'), '8296')
            
            hl = etree.SubElement(new_p, qn('w:hyperlink'))
            r = etree.SubElement(hl, qn('w:r'))
            t = etree.SubElement(r, qn('w:t'))
            t.text = title
            t.set(qn('xml:space'), 'preserve')
            r_tab = etree.SubElement(hl, qn('w:r'))
            etree.SubElement(r_tab, qn('w:tab'))
            r_page = etree.SubElement(hl, qn('w:r'))
            t_page = etree.SubElement(r_page, qn('w:t'))
            t_page.text = page
            
            insert_after.addnext(new_p)
            insert_after = new_p
        
        print(f"  Inserted {len(tables)} tables")
        break

doc.save(FILE)

# ============================================
# STEP 4: Fix DAFTAR LAMPIRAN
# ============================================
print("\n=== Step 4: Fix DAFTAR LAMPIRAN ===")
doc = Document(FILE)

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text == 'DAFTAR LAMPIRAN':
        p_elem = p._p
        
        # Remove old entries
        next_elem = p_elem.getnext()
        while next_elem is not None:
            nxt = next_elem.getnext()
            t_elems = next_elem.findall('.//' + qn('w:t'))
            txt = ''.join(t.text or '' for t in t_elems)
            if 'Lampiran' in txt and ('Surat' in txt or 'Absensi' in txt or 'Dokumentasi' in txt):
                parent = next_elem.getparent()
                if parent is not None:
                    parent.remove(next_elem)
                next_elem = nxt
            else:
                break
        
        # Insert correct entries
        insert_after = p_elem
        for entry in ['Lampiran A\tLogbook Kegiatan Harian', 'Lampiran B\tGitHub Repository', 'Lampiran C\tScreenshot Sistem']:
            new_p = etree.Element(qn('w:p'))
            r = etree.SubElement(new_p, qn('w:r'))
            t = etree.SubElement(r, qn('w:t'))
            t.text = entry
            t.set(qn('xml:space'), 'preserve')
            insert_after.addnext(new_p)
            insert_after = new_p
        
        print("  Replaced lampiran entries")
        break

doc.save(FILE)
print("\n=== ALL DONE ===")

# ============================================
# FINAL VERIFICATION
# ============================================
print("\n=== VERIFICATION ===")
doc = Document(FILE)
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() and i >= 96 and i < 220:
        style = p.style.style_id if p.style else ''
        text = p.text.strip()[:80]
        if style.startswith('TOC') or 'Heading' in (p.style.name if p.style else '') or 'DAFTAR' in text or 'LAMPIRAN' in text:
            print(f'[{i:3d}] {style:10s} | {text}')
