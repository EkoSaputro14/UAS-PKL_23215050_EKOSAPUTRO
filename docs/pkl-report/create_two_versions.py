"""
Create two versions of the PKL report:
1. Word Styles version — proper headings, captions, placeholder TOC
2. Plain Text version — current manual TOC preserved

Both versions fix duplicate captions.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import os, shutil

SRC = r'C:\Users\SMANSA\mimotes\docs\pkl-report\LAPORAN_PKL_Eko_Saputro_23215050.docx'
OUT_STYLES = r'C:\Users\SMANSA\mimotes\docs\pkl-report\LAPORAN_PKL_v10_Styles.docx'
OUT_PLAIN = r'C:\Users\SMANSA\mimotes\docs\pkl-report\LAPORAN_PKL_v10_Plain.docx'

# ============================================
# SHARED: Fix duplicate captions
# ============================================
def fix_captions(doc):
    """Fix duplicate text in caption paragraphs."""
    fixed = 0
    for p in doc.paragraphs:
        if p.style and 'Caption' in p.style.name:
            runs = p.runs
            if not runs:
                continue
            merged = ''.join(r.text for r in runs)
            # Fix known duplicates
            fixes = {
                'Gambar 4.2 Use Case Diagram Sistem Mimotes AI Sistem Mimotes AI': 'Gambar 4.2 Use Case Diagram Sistem Mimotes AI',
                'Gambar 4.3 Activity Diagram Upload Dokumen Dokumen': 'Gambar 4.3 Activity Diagram Upload Dokumen',
                'Gambar 4.6 Arsitektur Sistem Mimotes AI Mimotes AI': 'Gambar 4.6 Arsitektur Sistem Mimotes AI',
            }
            for old, new in fixes.items():
                if old in merged:
                    merged = merged.replace(old, new)
                    runs[0].text = merged
                    for r in runs[1:]:
                        r.text = ''
                    fixed += 1
                    break
    return fixed

# ============================================
# SHARED: Remove TOC entries between two headings
# ============================================
def find_para_index(doc, search_text, style_contains=None):
    """Find paragraph index matching text and optional style."""
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ''
        if search_text in text:
            if style_contains is None or style_contains in style:
                return i
    return None

def remove_toc_range(doc, start_idx, end_idx):
    """Remove paragraphs in range [start_idx, end_idx) from the document body."""
    body = doc.element.body
    all_p = list(body.findall(qn('w:p')))
    
    removed = 0
    # We need to remove from end to start to maintain indices
    for i in range(end_idx - 1, start_idx - 1, -1):
        if i < len(all_p):
            body.remove(all_p[i])
            removed += 1
    return removed

# ============================================
# VERSION 1: Word Styles (proper headings, placeholder TOC)
# ============================================
print("=== Building VERSION 1: Word Styles ===")
doc1 = Document(SRC)

# Fix captions
cap_fixes = fix_captions(doc1)
print(f"  Fixed {cap_fixes} duplicate captions")

# Find and remove manual TOC sections
# Section 1: DAFTAR ISI (para 96 heading → until para 203 which is "DAFTAR LAMBANG")
toc_start = find_para_index(doc1, 'DAFTAR ISI', 'Heading 1')
toc_end = find_para_index(doc1, 'BAB I', 'Heading 1')

if toc_start is not None and toc_end is not None:
    # Remove all paragraphs between DAFTAR ISI heading and BAB I heading
    # But keep the DAFTAR ISI heading itself — replace it with placeholder
    # First, remove the toc entries after DAFTAR ISI heading
    body = doc1.element.body
    all_p = list(body.findall(qn('w:p')))
    
    # Find the actual indices in XML
    removed = 0
    # Remove from end to start
    for i in range(toc_end - 1, toc_start, -1):
        if i < len(all_p):
            body.remove(all_p[i])
            removed += 1
    print(f"  Removed {removed} TOC entries (DAFTAR ISI)")
    
    # Add placeholder paragraph after DAFTAR ISI heading
    all_p2 = list(body.findall(qn('w:p')))
    if toc_start < len(all_p2):
        # Insert a new paragraph after the DAFTAR ISI heading
        placeholder = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'Normal')
        pPr.append(pStyle)
        placeholder.append(pPr)
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = '[DAFTAR ISI AKAN DIGENERATE OLEH MICROSOFT WORD — Klik kanan → Update Field → Update entire table]'
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        placeholder.append(r)
        
        # Insert after DAFTAR ISI heading
        toc_heading = all_p2[toc_start]
        toc_heading.addnext(placeholder)
        print("  Added TOC placeholder")

# Section 2: DAFTAR GAMBAR — find and replace
dg_heading = find_para_index(doc1, 'DAFTAR GAMBAR', 'Heading 1')
if dg_heading is not None:
    # Find next heading after DAFTAR GAMBAR
    next_h1 = None
    for i in range(dg_heading + 1, len(doc1.paragraphs)):
        if doc1.paragraphs[i].style and 'Heading 1' in doc1.paragraphs[i].style.name:
            next_h1 = i
            break
    
    if next_h1:
        body = doc1.element.body
        all_p = list(body.findall(qn('w:p')))
        removed = 0
        for i in range(next_h1 - 1, dg_heading, -1):
            if i < len(all_p):
                body.remove(all_p[i])
                removed += 1
        print(f"  Removed {removed} entries (DAFTAR GAMBAR)")
        
        # Add placeholder
        all_p2 = list(body.findall(qn('w:p')))
        if dg_heading < len(all_p2):
            placeholder = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), 'Normal')
            pPr.append(pStyle)
            placeholder.append(pPr)
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = '[DAFTAR GAMBAR AKAN DIGENERATE OLEH MICROSOFT WORD — Klik kanan → Insert Caption → Update Field]'
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            placeholder.append(r)
            all_p2[dg_heading].addnext(placeholder)
            print("  Added Daftar Gambar placeholder")

# Section 3: DAFTAR TABEL — find and replace
dt_heading = find_para_index(doc1, 'DAFTAR TABEL', 'Heading 1')
if dt_heading is not None:
    next_h1 = None
    for i in range(dt_heading + 1, len(doc1.paragraphs)):
        if doc1.paragraphs[i].style and 'Heading 1' in doc1.paragraphs[i].style.name:
            next_h1 = i
            break
    
    if next_h1:
        body = doc1.element.body
        all_p = list(body.findall(qn('w:p')))
        removed = 0
        for i in range(next_h1 - 1, dt_heading, -1):
            if i < len(all_p):
                body.remove(all_p[i])
                removed += 1
        print(f"  Removed {removed} entries (DAFTAR TABEL)")
        
        all_p2 = list(body.findall(qn('w:p')))
        if dt_heading < len(all_p2):
            placeholder = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), 'Normal')
            pPr.append(pStyle)
            placeholder.append(pPr)
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = '[DAFTAR TABEL AKAN DIGENERATE OLEH MICROSOFT WORD — Klik kanan → Insert Caption → Update Field]'
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            placeholder.append(r)
            all_p2[dt_heading].addnext(placeholder)
            print("  Added Daftar Tabel placeholder")

# Remove DAFTAR LAMPIRAN and DAFTAR LAMBANG (if they exist as plain text)
for search in ['DAFTAR LAMPIRAN', 'DAFTAR LAMBANG']:
    idx = find_para_index(doc1, search)
    if idx is not None:
        # Check if it's not a Heading 1 — if it's Normal style, it's manual
        p = doc1.paragraphs[idx]
        if p.style and 'Heading' not in p.style.name:
            body = doc1.element.body
            all_p = list(body.findall(qn('w:p')))
            if idx < len(all_p):
                body.remove(all_p[idx])
                # Also remove following entries
                for j in range(idx, idx + 5):
                    all_p2 = list(body.findall(qn('w:p')))
                    if j < len(all_p2):
                        t = all_p2[j].text if all_p2[j].text else ''
                        if 'Lampiran' in t or 'Lambang' in t or 'Singkatan' in t:
                            body.remove(all_p2[j])
                            print(f"  Removed: {t[:50]}")

doc1.save(OUT_STYLES)
print(f"\nVERSION 1 saved: {OUT_STYLES}")
print(f"Size: {os.path.getsize(OUT_STYLES) / 1024:.0f} KB")

# ============================================
# VERSION 2: Plain Text (current manual TOC preserved)
# ============================================
print("\n=== Building VERSION 2: Plain Text ===")
doc2 = Document(SRC)

# Fix captions only
cap_fixes2 = fix_captions(doc2)
print(f"  Fixed {cap_fixes2} duplicate captions")

doc2.save(OUT_PLAIN)
print(f"\nVERSION 2 saved: {OUT_PLAIN}")
print(f"Size: {os.path.getsize(OUT_PLAIN) / 1024:.0f} KB")

# ============================================
# VERIFICATION
# ============================================
print("\n=== VERIFICATION ===")

# Verify VERSION 1
doc_v1 = Document(OUT_STYLES)
h1_count = sum(1 for p in doc_v1.paragraphs if p.style and 'Heading 1' in p.style.name)
h2_count = sum(1 for p in doc_v1.paragraphs if p.style and 'Heading 2' in p.style.name)
h3_count = sum(1 for p in doc_v1.paragraphs if p.style and 'Heading 3' in p.style.name)
cap_count = sum(1 for p in doc_v1.paragraphs if p.style and 'Caption' in p.style.name)
toc_count = sum(1 for p in doc_v1.paragraphs if p.style and p.style.name.startswith('toc'))

# Check for placeholder text
has_toc_placeholder = any('[DAFTAR ISI AKAN DIGENERATE' in p.text for p in doc_v1.paragraphs)
has_dg_placeholder = any('[DAFTAR GAMBAR AKAN DIGENERATE' in p.text for p in doc_v1.paragraphs)
has_dt_placeholder = any('[DAFTAR TABEL AKAN DIGENERATE' in p.text for p in doc_v1.paragraphs)

print(f"\nVERSION 1 (Word Styles):")
print(f"  Heading 1: {h1_count}")
print(f"  Heading 2: {h2_count}")
print(f"  Heading 3: {h3_count}")
print(f"  Captions: {cap_count}")
print(f"  TOC entries: {toc_count} (should be 0)")
print(f"  TOC placeholder: {has_toc_placeholder}")
print(f"  DG placeholder: {has_dg_placeholder}")
print(f"  DT placeholder: {has_dt_placeholder}")

# Verify VERSION 2
doc_v2 = Document(OUT_PLAIN)
toc_v2 = sum(1 for p in doc_v2.paragraphs if p.style and p.style.name.startswith('toc'))
print(f"\nVERSION 2 (Plain Text):")
print(f"  TOC entries preserved: {toc_v2}")
