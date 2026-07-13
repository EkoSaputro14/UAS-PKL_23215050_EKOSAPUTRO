#!/usr/bin/env python3
"""
Merge v14 content into the new template.
Replace: title, student name, NIM, supervisor, and all BAB content.
"""
import copy
import os
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE = 'C:/Users/SMANSA/AppData/Local/hermes/cache/documents/doc_4ce6543a497c_LAPORAN PKL (1).docx'
SOURCE = 'LAPORAN_PKL_v14_UML.docx'
OUTPUT = 'LAPORAN_PKL_v15_Template.docx'
DIAGRAMS_DIR = 'diagrams'

# ===== NEW CONTENT =====
NEW_TITLE = "RANCANG BANGUN SISTEM CHATBOT AI BERBASIS PENGETAHUAN DENGAN RETRIEVAL-AUGMENTED GENERATION DAN PIPELINE CRM UNTUK OPTIMALISASI LAYANAN PELANGGAN"
NEW_STUDENT = "Eko Saputro"
NEW_NIM = "23215050"
NEW_SUPERVISOR = "Widianto Agung Nugroho"
NEW_INSTITUTION = "Bank Mandiri KCP Tegal Sudirman"

# Old content to replace
OLD_TITLE_KEYWORDS = ["SISTEM PAKAR DETEKSI HAMA", "SISTEM PAKAR"]
OLD_STUDENT = "Moh. Arif Prasetyo"
OLD_NIM = "23215043"
OLD_SUPERVISOR_KEYWORDS = ["Nama Dosen Pembimbing", "Dosen Pembimbing,"]

print("📄 Loading template...")
tpl_doc = Document(TEMPLATE)
src_doc = Document(SOURCE)

# ===== STEP 1: Replace text in template =====
print("\n📝 Step 1: Replacing text in template...")

replace_count = 0
for para in tpl_doc.paragraphs:
    for run in para.runs:
        original = run.text
        
        # Replace title
        for kw in OLD_TITLE_KEYWORDS:
            if kw in run.text:
                run.text = run.text.replace(kw, NEW_TITLE[:len(run.text)])
                replace_count += 1
        
        # Replace student name
        if OLD_STUDENT in run.text:
            run.text = run.text.replace(OLD_STUDENT, NEW_STUDENT)
            replace_count += 1
        
        # Replace NIM
        if OLD_NIM in run.text:
            run.text = run.text.replace(OLD_NIM, NEW_NIM)
            replace_count += 1
        
        # Replace supervisor name placeholder
        if "Nama Dosen Pembimbing" in run.text:
            run.text = run.text.replace("Nama Dosen Pembimbing", NEW_SUPERVISOR)
            replace_count += 1
        if "NIPY." in run.text:
            run.text = run.text.replace("NIPY.", "")
            replace_count += 1

# Also check paragraphs without runs (some paragraphs have text directly)
for para in tpl_doc.paragraphs:
    text = para.text
    if OLD_TITLE_KEYWORDS[0] in text:
        # This paragraph has the full title
        for run in para.runs:
            if OLD_TITLE_KEYWORDS[0] in run.text:
                run.text = NEW_TITLE
                replace_count += 1
                break

print(f"  ✅ {replace_count} text replacements made")

# ===== STEP 2: Find BAB positions in template =====
print("\n📝 Step 2: Finding BAB positions in template...")

tpl_babs = {}
for i, para in enumerate(tpl_doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    if style == 'Heading 1':
        if 'BAB I' in text and 'PENDAHULUAN' not in text:
            # Check next paragraph for PENDAHULUAN
            if i + 1 < len(tpl_doc.paragraphs):
                next_text = tpl_doc.paragraphs[i + 1].text.strip()
                if 'PENDAHULUAN' in next_text:
                    tpl_babs['BAB_I'] = i
                    print(f"  Found BAB I at paragraph {i}")
        elif 'BAB II' in text:
            tpl_babs['BAB_II'] = i
            print(f"  Found BAB II at paragraph {i}")
        elif 'BAB III' in text:
            tpl_babs['BAB_III'] = i
            print(f"  Found BAB III at paragraph {i}")
        elif 'BAB IV' in text:
            tpl_babs['BAB_IV'] = i
            print(f"  Found BAB IV at paragraph {i}")
        elif 'BAB V' in text:
            tpl_babs['BAB_V'] = i
            print(f"  Found BAB V at paragraph {i}")
        elif 'DAFTAR PUSTAKA' in text:
            tpl_babs['DAFTAR_PUSTAKA'] = i
            print(f"  Found DAFTAR PUSTAKA at paragraph {i}")
        elif 'LAMPIRAN' in text:
            tpl_babs['LAMPIRAN'] = i
            print(f"  Found LAMPIRAN at paragraph {i}")

# ===== STEP 3: Extract content from source =====
print("\n📝 Step 3: Extracting content from source...")

src_content = {}
current_bab = None
current_content = []

for para in src_doc.paragraphs:
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    # Detect BAB headings
    if style == 'Heading 1':
        if current_bab and current_content:
            src_content[current_bab] = current_content
        current_content = []
        
        if 'BAB I' in text:
            current_bab = 'BAB_I'
        elif 'BAB II' in text:
            current_bab = 'BAB_II'
        elif 'BAB III' in text:
            current_bab = 'BAB_III'
        elif 'BAB IV' in text:
            current_bab = 'BAB_IV'
        elif 'BAB V' in text:
            current_bab = 'BAB_V'
        elif 'DAFTAR PUSTAKA' in text:
            current_bab = 'DAFTAR_PUSTAKA'
        elif 'LAMPIRAN' in text:
            current_bab = 'LAMPIRAN'
        else:
            current_bab = None
        continue
    
    if current_bab:
        current_content.append({
            'text': text,
            'style': style,
            'runs': [(r.text, r.bold, r.italic) for r in para.runs],
        })

# Save last BAB
if current_bab and current_content:
    src_content[current_bab] = current_content

for bab, content in src_content.items():
    print(f"  {bab}: {len(content)} paragraphs")

# ===== STEP 4: Replace BAB content in template =====
# This is complex - we need to clear the old content and insert new content
# between the heading and the next heading

print("\n📝 Step 4: Replacing BAB content...")

# For simplicity, we'll replace the text of existing paragraphs
# and remove excess paragraphs if source has fewer

# Map source BAB keys to template BAB keys
bab_mapping = {
    'BAB_I': 'BAB_I',
    'BAB_II': 'BAB_II',
    'BAB_III': 'BAB_III',
    'BAB_IV': 'BAB_IV',
    'BAB_V': 'BAB_V',
}

for src_key, tpl_key in bab_mapping.items():
    if src_key not in src_content or tpl_key not in tpl_babs:
        print(f"  ⚠️ Skipping {src_key} (not found)")
        continue
    
    start_idx = tpl_babs[tpl_key]
    # Find the end (next heading or end of document)
    end_idx = len(tpl_doc.paragraphs)
    for next_key, next_idx in tpl_babs.items():
        if next_idx > start_idx:
            end_idx = min(end_idx, next_idx)
            break
    
    # Get paragraphs in this section
    section_paras = list(tpl_doc.paragraphs[start_idx:end_idx])
    src_paras = src_content[src_key]
    
    # Replace content (skip the heading itself)
    content_start = 1 if section_paras[0].text.strip().startswith('BAB') else 0
    # Also skip "PENDAHULUAN" after "BAB I"
    if tpl_key == 'BAB_I' and len(section_paras) > 1:
        if section_paras[1].text.strip() == 'PENDAHULUAN':
            content_start = 2
    
    available_slots = len(section_paras) - content_start
    
    for i, src_p in enumerate(src_paras):
        if i < available_slots:
            # Replace existing paragraph
            target_para = section_paras[content_start + i]
            # Clear existing runs
            for run in target_para.runs:
                run.text = ''
            # Set new text
            if target_para.runs:
                target_para.runs[0].text = src_p['text']
            elif src_p['text']:
                # No runs, add one
                run = target_para.add_run(src_p['text'])
    
    replaced = min(len(src_paras), available_slots)
    print(f"  ✅ {src_key}: replaced {replaced}/{len(src_paras)} paragraphs")

# ===== STEP 5: Save =====
print("\n💾 Step 5: Saving...")
tpl_doc.save(OUTPUT)
print(f"✅ Saved: {OUTPUT}")

# ===== VERIFY =====
doc2 = Document(OUTPUT)
print(f"\n📊 Verification:")
print(f"  Paragraphs: {len(doc2.paragraphs)}")
print(f"  Tables: {len(doc2.tables)}")

# Check key text replacements
for para in doc2.paragraphs[:100]:
    text = para.text.strip()
    if NEW_STUDENT in text:
        print(f"  ✅ Found student name: {text[:50]}")
    if NEW_NIM in text:
        print(f"  ✅ Found NIM: {text[:50]}")
    if NEW_TITLE[:30] in text:
        print(f"  ✅ Found title: {text[:60]}")

print(f"\n  File size: {os.path.getsize(OUTPUT)/1024:.0f} KB")
