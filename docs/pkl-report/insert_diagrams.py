#!/usr/bin/env python3
"""
Insert diagrams from v14 into v15 template.
Replace old figure captions and insert images.
"""
import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Cm, Pt
from lxml import etree

INPUT = 'LAPORAN_PKL_v15_Template.docx'
OUTPUT = 'LAPORAN_PKL_v15_Template.docx'
DIAGRAMS_DIR = 'diagrams'

# Mapping: old caption keyword → new caption + image file
# We'll find captions by their position in BAB IV content
DIAGRAM_MAP = [
    # (search_text_in_caption, new_caption, image_file)
    ("Use Case", "Gambar 4.2 Use Case Diagram Sistem Mimotes AI", "usecase.png"),
    ("Activity Diagram Proses Diag", "Gambar 4.3 Activity Diagram Upload Dokumen", "activity-upload.png"),
    ("Activity Diagram", "Gambar 4.4 Activity Diagram Proses Chat RAG", "activity-chat.png"),
    ("Sequence diagram Proses", "Gambar 4.5 Sequence Diagram Chat RAG", "sequence-chat-rag.png"),
    ("Class Diagram", "Gambar 4.6 ERD Domain Identity & Workspace", "erd-a-identity.png"),
    ("Pohon Keputusan", "Gambar 4.7 ERD Domain RAG & Knowledge Base", "erd-b-rag.png"),
    ("Arsitektur Sistem", "Gambar 4.8 ERD Domain Chat & CRM", "erd-c-chat-crm.png"),
    ("Flowchart Proses", "Gambar 4.9 ERD Domain Billing & Configuration", "erd-d-billing.png"),
    ("Halaman Beranda", "Gambar 4.10 ERD Ringkasan — Seluruh Relasi", "erd-summary.png"),
    ("Halaman Tentang", "Gambar 4.11 Arsitektur Sistem Mimotes AI", "architecture.png"),
    ("Halaman Cara Kerja", "Gambar 4.12 Arsitektur RAG Pipeline", "rag-pipeline.png"),
    ("Halaman Deteksi", "Gambar 4.13 Arsitektur CRM Pipeline", "crm-pipeline.png"),
    ("Halaman Panel Admin", "Gambar 4.14 Halaman Login", None),  # screenshot - skip
    ("Implementasi database", "Gambar 4.15 Dashboard Admin", None),
    ("Input Keyakinan", "Gambar 4.16 Upload Dokumen", None),
    ("Hasil Diag", "Gambar 4.17 Daftar Dokumen", None),
]

print("📄 Loading v15...")
doc = Document(INPUT)
body = doc.element.body

# ===== STEP 1: Find all figure captions in BAB IV =====
print("\n📝 Step 1: Finding figure captions...")

# Find where BAB IV content starts and ends
bab4_start = None
bab5_start = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    if style == 'Heading 1' and 'BAB IV' in text:
        bab4_start = i
    elif style == 'Heading 1' and 'BAB V' in text:
        bab5_start = i
        break

print(f"  BAB IV: paragraphs {bab4_start}-{bab5_start}")

# Find figure captions in BAB IV
figure_positions = []
if bab4_start is not None and bab5_start is not None:
    for i in range(bab4_start, bab5_start):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if text.startswith('Gambar 4.'):
            figure_positions.append((i, text, para))

print(f"  Found {len(figure_positions)} figure captions in BAB IV")
for idx, text, _ in figure_positions:
    print(f"    [{idx:3d}] {text[:60]}")

# ===== STEP 2: Match and replace captions + insert images =====
print("\n📝 Step 2: Replacing captions and inserting images...")

replaced = 0
for caption_idx, (para_idx, old_text, para) in enumerate(figure_positions):
    # Find matching diagram
    matched = False
    for search_text, new_caption, img_file in DIAGRAM_MAP:
        if search_text.lower() in old_text.lower():
            # Replace caption text
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = new_caption
            elif new_text:
                para.add_run(new_caption)
            
            # Insert image if available
            if img_file:
                img_path = os.path.join(DIAGRAMS_DIR, img_file)
                if os.path.exists(img_path):
                    # Create a new paragraph for the image (before the caption)
                    # We need to insert it into the XML body
                    para_elem = para._element
                    parent = para_elem.getparent()
                    
                    # Find position of this paragraph in parent
                    pos = list(parent).index(para_elem)
                    
                    # Create image paragraph
                    img_para = etree.SubElement(etree.Element('dummy'), qn('w:p'))
                    img_pPr = etree.SubElement(img_para, qn('w:pPr'))
                    img_jc = etree.SubElement(img_pPr, qn('w:jc'))
                    img_jc.set(qn('w:val'), 'center')
                    
                    # Add drawing element with image
                    img_r = etree.SubElement(img_para, qn('w:r'))
                    drawing = etree.SubElement(img_r, qn('w:drawing'))
                    
                    # Create inline image
                    inline = etree.SubElement(drawing, qn('wp:inline'))
                    inline.set(qn('distT'), '0')
                    inline.set(qn('distB'), '0')
                    inline.set(qn('distL'), '0')
                    inline.set(qn('distR'), '0')
                    
                    # extent (size)
                    extent = etree.SubElement(inline, qn('wp:extent'))
                    # Width: 14cm = ~5.5 inches = ~5.5*914400 EMUs
                    cx = int(14.0 / 2.54 * 914400)  # 14cm in EMUs
                    cy = int(cx * 0.6)  # rough aspect ratio
                    extent.set(qn('cx'), str(cx))
                    extent.set(qn('cy'), str(cy))
                    
                    # docPr
                    docPr = etree.SubElement(inline, qn('wp:docPr'))
                    docPr.set(qn('id'), str(replaced + 1))
                    docPr.set(qn('name'), f'Image {replaced + 1}')
                    
                    # Graphic
                    graphic = etree.SubElement(inline, qn('a:graphic'))
                    graphic.set(qn('xmlns:a'), 'http://schemas.openxmlformats.org/drawingml/2006/main')
                    graphicFrame = etree.SubElement(graphic, qn('a:graphicFrame'))
                    graphicFrame.set(qn('macro'), '')
                    
                    xfrm = etree.SubElement(graphicFrame, qn('a:xfrm'))
                    off = etree.SubElement(xfrm, qn('a:off'))
                    off.set(qn('x'), '0')
                    off.set(qn('y'), '0')
                    ext = etree.SubElement(xfrm, qn('a:ext'))
                    ext.set(qn('cx'), '0')
                    ext.set(qn('cy'), '0')
                    
                    prstGeom = etree.SubElement(graphicFrame, qn('a:prstGeom'))
                    prstGeom.set(qn('prst'), 'rect')
                    avLst = etree.SubElement(prstGeom, qn('a:avLst'))
                    
                    # pic:nvPicPr (picture non-visual properties)
                    # We need to use the pic namespace
                    pic_ns = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
                    
                    # For simplicity, let's use a different approach:
                    # Add the image as a relationship and reference it
                    # This is complex with raw XML, so let's use python-docx's add_picture method
                    # But that adds to the end of the document...
                    
                    # Actually, the simplest approach is to use python-docx's
                    # paragraph._element.addnext() or addprevious()
                    
                    # Let's just note the position and use python-docx later
                    pass
            
            replaced += 1
            matched = True
            print(f"  ✅ [{para_idx}] {old_text[:40]} → {new_caption[:40]}")
            break
    
    if not matched:
        print(f"  ⚠️ [{para_idx}] No match for: {old_text[:50]}")

print(f"\n  Replaced {replaced}/{len(figure_positions)} captions")

# ===== SAVE =====
print(f"\n💾 Saving...")
doc.save(OUTPUT)
print(f"✅ Saved: {OUTPUT}")

# ===== Now insert images using python-docx =====
print("\n🖼️ Step 3: Inserting images...")

doc2 = Document(OUTPUT)

# Find figure captions again and insert images before them
for i, para in enumerate(doc2.paragraphs):
    text = para.text.strip()
    if not text.startswith('Gambar 4.'):
        continue
    
    # Find matching image
    img_file = None
    for search_text, new_caption, img_f in DIAGRAM_MAP:
        if search_text.lower() in text.lower() and img_f:
            img_file = img_f
            break
    
    if not img_file:
        continue
    
    img_path = os.path.join(DIAGRAMS_DIR, img_file)
    if not os.path.exists(img_path):
        print(f"  ⚠️ Image not found: {img_path}")
        continue
    
    try:
        # Insert image before this paragraph
        # We use add_run on a previous paragraph, or add a new paragraph
        # The cleanest way is to insert a new paragraph before this one
        
        # Get the paragraph's XML element
        para_elem = para._element
        parent = para_elem.getparent()
        pos = list(parent).index(para_elem)
        
        # Create new paragraph with image
        new_para = doc2.add_paragraph()
        new_para.alignment = 1  # CENTER
        run = new_para.add_run()
        run.add_picture(img_path, width=Inches(5.5))
        
        # Move the new paragraph to before the caption
        new_para_elem = new_para._element
        parent.remove(new_para_elem)
        parent.insert(pos, new_para_elem)
        
        print(f"  ✅ Inserted {img_file} before '{text[:40]}'")
    except Exception as e:
        print(f"  ❌ Failed to insert {img_file}: {e}")

doc2.save(OUTPUT)
print(f"\n💾 Final save: {OUTPUT}")

# ===== VERIFY =====
doc3 = Document(OUTPUT)
img_count = 0
for para in doc3.paragraphs:
    for run in para.runs:
        if run._element.findall(qn('w:drawing')):
            img_count += 1

print(f"\n📊 Verification:")
print(f"  Images inserted: {img_count}")
print(f"  File size: {os.path.getsize(OUTPUT)/1024:.0f} KB")
