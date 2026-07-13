#!/usr/bin/env python3
"""
Final: Add figure captions and images to BAB IV content.
The v5 rebuild has clean content but no figure captions/images.
This script adds them.
"""
import os
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from lxml import etree

FILE = 'LAPORAN_PKL_v15_Template.docx'
DIAGRAMS_DIR = 'diagrams'

# Figure captions to add (positioned after relevant content sections)
FIGURES = [
    # (search_text_in_content, caption, image_file)
    # We'll insert captions at the end of relevant sub-sections
    ("4.2.3 Use Case Diagram", "Gambar 4.2 Use Case Diagram Sistem Mimotes AI", "usecase.png"),
    ("4.2.4 Activity Diagram Upload", "Gambar 4.3 Activity Diagram Upload Dokumen", "activity-upload.png"),
    ("4.2.5 Activity Diagram Proses Chat", "Gambar 4.4 Activity Diagram Proses Chat RAG", "activity-chat.png"),
    ("4.2.6 Sequence Diagram Chat RAG", "Gambar 4.5 Sequence Diagram Chat RAG", "sequence-chat-rag.png"),
    ("4.2.7 Entity Relationship Diagram", "Gambar 4.6 ERD Domain Identity & Workspace", "erd-a-identity.png"),
    # After ERD section, add remaining ERDs
    ("4.2.8 Arsitektur Sistem", "Gambar 4.11 Arsitektur Sistem Mimotes AI", "architecture.png"),
    ("4.2.9 Arsitektur RAG Pipeline", "Gambar 4.12 Arsitektur RAG Pipeline", "rag-pipeline.png"),
    ("4.2.10 Arsitektur CRM Pipeline", "Gambar 4.13 Arsitektur CRM Pipeline", "crm-pipeline.png"),
]

# Additional ERDs to insert after ERD section
ERD_FIGURES = [
    ("Gambar 4.7 ERD Domain RAG & Knowledge Base", "erd-b-rag.png"),
    ("Gambar 4.8 ERD Domain Chat & CRM", "erd-c-chat-crm.png"),
    ("Gambar 4.9 ERD Domain Billing & Configuration", "erd-d-billing.png"),
    ("Gambar 4.10 ERD Ringkasan — Seluruh Relasi", "erd-summary.png"),
]

# Screenshot captions (no images, just captions)
SCREENSHOTS = [
    "Gambar 4.14 Halaman Login",
    "Gambar 4.15 Dashboard Admin",
    "Gambar 4.16 Upload Dokumen",
    "Gambar 4.17 Daftar Dokumen",
    "Gambar 4.18 Chat AI",
    "Gambar 4.19 Knowledge Search",
    "Gambar 4.20 Analytics Chat",
    "Gambar 4.21 Pengaturan AI Provider",
]

print("📄 Loading...")
doc = Document(FILE)

# ===== Step 1: Find key positions in BAB IV =====
print("\n📝 Finding positions...")

positions = {}
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    # Match heading 3 sections
    for search, caption, img in FIGURES:
        if search in text:
            positions[search] = (i, caption, img)

for key, (idx, cap, img) in positions.items():
    print(f"  [{idx:3d}] {key}")

# ===== Step 2: Insert captions and images =====
print("\n📝 Inserting captions and images...")

# Process in reverse order to avoid index shifting
sorted_positions = sorted(positions.items(), key=lambda x: x[1][0], reverse=True)

for search, (para_idx, caption, img_file) in sorted_positions:
    para = doc.paragraphs[para_idx]
    
    # Create caption paragraph
    cap_para = doc.add_paragraph()
    cap_para.alignment = 1  # CENTER
    cap_para.text = caption
    # Make it smaller font
    for run in cap_para.runs:
        run.font.size = Inches(0.1)  # Will be set properly below
    
    # Create image paragraph
    img_para = doc.add_paragraph()
    img_para.alignment = 1  # CENTER
    
    img_path = os.path.join(DIAGRAMS_DIR, img_file)
    if os.path.exists(img_path):
        img_para.add_run().add_picture(img_path, width=Inches(5.5))
    
    # Insert image before caption, both after the section heading
    parent = para._element.getparent()
    pos = list(parent).index(para._element) + 1
    
    # Move image paragraph
    img_elem = img_para._element
    parent.remove(img_elem)
    parent.insert(pos, img_elem)
    
    # Move caption paragraph
    cap_elem = cap_para._element
    parent.remove(cap_elem)
    parent.insert(pos + 1, cap_elem)
    
    print(f"  ✅ {caption[:50]} + {img_file}")

# ===== Step 3: Add ERD figures after ERD section =====
print("\n📝 Adding ERD figures...")

# Find the ERD section heading
erd_section_idx = None
for i, para in enumerate(doc.paragraphs):
    if '4.2.7 Entity Relationship Diagram' in para.text:
        erd_section_idx = i
        break

if erd_section_idx:
    # Find next heading after ERD section
    next_heading_idx = len(doc.paragraphs)
    for i in range(erd_section_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name.startswith('Heading'):
            next_heading_idx = i
            break
    
    # Insert ERD figures before next heading
    insert_before = doc.paragraphs[next_heading_idx]._element if next_heading_idx < len(doc.paragraphs) else None
    parent = insert_before.getparent() if insert_before else doc.element.body
    insert_pos = list(parent).index(insert_before) if insert_before else len(list(parent))
    
    for j, (cap_text, img_file) in enumerate(ERD_FIGURES):
        img_path = os.path.join(DIAGRAMS_DIR, img_file)
        
        # Image paragraph
        img_para = doc.add_paragraph()
        img_para.alignment = 1
        if os.path.exists(img_path):
            img_para.add_run().add_picture(img_path, width=Inches(5.5))
        
        # Caption paragraph
        cap_para = doc.add_paragraph()
        cap_para.alignment = 1
        cap_para.text = cap_text
        
        # Insert
        img_elem = img_para._element
        cap_elem = cap_para._element
        parent.remove(img_elem)
        parent.remove(cap_elem)
        parent.insert(insert_pos + j * 2, img_elem)
        parent.insert(insert_pos + j * 2 + 1, cap_elem)
        
        print(f"  ✅ {cap_text[:50]} + {img_file}")

# ===== Step 4: Add screenshot captions =====
print("\n📝 Adding screenshot captions...")

# Find "4.6 Tampilan Antarmuka Sistem" section
tampilan_idx = None
for i, para in enumerate(doc.paragraphs):
    if '4.6 Tampilan' in para.text:
        tampilan_idx = i
        break

if tampilan_idx:
    # Find next heading
    next_h = len(doc.paragraphs)
    for i in range(tampilan_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name.startswith('Heading 1'):
            next_h = i
            break
    
    # Add screenshot captions
    insert_before = doc.paragraphs[next_h]._element if next_h < len(doc.paragraphs) else None
    parent = insert_before.getparent() if insert_before else doc.element.body
    insert_pos = list(parent).index(insert_before) if insert_before else len(list(parent))
    
    for j, cap_text in enumerate(SCREENSHOTS):
        cap_para = doc.add_paragraph()
        cap_para.alignment = 1
        cap_para.text = cap_text
        cap_elem = cap_para._element
        parent.remove(cap_elem)
        parent.insert(insert_pos + j, cap_elem)
        
        print(f"  ✅ {cap_text}")

# ===== Save =====
print("\n💾 Saving...")
doc.save(FILE)

# ===== Verify =====
doc2 = Document(FILE)
total_imgs = sum(
    len(run._element.findall(qn('w:drawing')))
    for para in doc2.paragraphs
    for run in para.runs
)
cap_count = sum(1 for p in doc2.paragraphs if p.text.strip().startswith('Gambar 4.'))

print(f"\n📊 Final:")
print(f"  Paragraphs: {len(doc2.paragraphs)}")
print(f"  Images: {total_imgs}")
print(f"  Figure captions: {cap_count}")
print(f"  Size: {os.path.getsize(FILE)/1024:.0f} KB")
