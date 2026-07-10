"""
Rebuild front matter completely — one clean pass.
Keep everything before DAFTAR ISI, replace everything between DAFTAR ISI and BAB I.
"""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

FILE = 'LAPORAN_PKL_v10_Plain_final.docx'
doc = Document(FILE)
body = doc.element.body
all_p = list(body.findall(qn('w:p')))

# Find key positions
positions = {}
for i, p_elem in enumerate(all_p):
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        continue
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        continue
    sv = pStyle.get(qn('w:val'))
    if sv != 'Heading1':
        continue
    t_elems = p_elem.findall('.//' + qn('w:t'))
    text = ''.join(t.text or '' for t in t_elems).strip()
    if 'DAFTAR ISI' in text and 'DAFTAR ISI' not in positions:
        positions['toc_start'] = i
    elif 'BAB I' in text:
        positions['body_start'] = i
        break

print(f"DAFTAR ISI at index: {positions.get('toc_start')}")
print(f"BAB I at index: {positions.get('body_start')}")

toc_start = positions['toc_start']
body_start = positions['body_start']

# Remove ALL paragraphs between DAFTAR ISI and BAB I (exclusive)
to_remove = all_p[toc_start + 1:body_start]
for elem in to_remove:
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)

print(f"Removed {len(to_remove)} paragraphs between DAFTAR ISI and BAB I")

# Now insert fresh content after DAFTAR ISI heading
toc_heading = all_p[toc_start]
insert_after = toc_heading

def make_toc_entry(text, page, style='TOC1'):
    """Create a TOC entry paragraph with proper tab structure."""
    new_p = etree.Element(qn('w:p'))
    pPr = etree.SubElement(new_p, qn('w:pPr'))
    pStyle = etree.SubElement(pPr, qn('w:pStyle'))
    pStyle.set(qn('w:val'), style)
    tabs = etree.SubElement(pPr, qn('w:tabs'))
    tab = etree.SubElement(tabs, qn('w:tab'))
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '8296')
    
    hl = etree.SubElement(new_p, qn('w:hyperlink'))
    r = etree.SubElement(hl, qn('w:r'))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r_tab = etree.SubElement(hl, qn('w:r'))
    etree.SubElement(r_tab, qn('w:tab'))
    r_page = etree.SubElement(hl, qn('w:r'))
    t_page = etree.SubElement(r_page, qn('w:t'))
    t_page.text = str(page)
    
    return new_p

def make_heading(text):
    """Create a Heading 1 paragraph."""
    new_p = etree.Element(qn('w:p'))
    pPr = etree.SubElement(new_p, qn('w:pPr'))
    pStyle = etree.SubElement(pPr, qn('w:pStyle'))
    pStyle.set(qn('w:val'), 'Heading1')
    r = etree.SubElement(new_p, qn('w:r'))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    return new_p

# Insert DAFTAR ISI content
toc_entries = [
    ('BAB I PENDAHULUAN', 1),
    ('  1.1 Latar Belakang', 1),
    ('  1.2 Rumusan Masalah', 1),
    ('  1.3 Tujuan PKL', 2),
    ('    1.3.1 Tujuan Umum', 2),
    ('    1.3.2 Tujuan Khusus', 2),
    ('  1.4 Manfaat PKL', 3),
    ('BAB II GAMBARAN UMUM INSTANSI', 4),
    ('  2.1 Sejarah Perkembangan Perusahaan', 4),
    ('  2.2 Visi, Misi, dan Tujuan', 4),
    ('    2.2.1 Visi', 4),
    ('    2.2.2 Misi', 5),
    ('  2.3 Struktur Organisasi', 5),
    ('  2.4 Job Deskripsi', 6),
    ('BAB III METODE PELAKSANAAN PKL', 7),
    ('  3.1 Landasan Teori', 7),
    ('    3.1.1 AI dan Chatbot', 7),
    ('    3.1.2 RAG', 7),
    ('    3.1.3 LLM', 8),
    ('    3.1.4 Embedding dan Vector DB', 8),
    ('    3.1.5 CRM', 8),
    ('    3.1.6 Next.js dan PostgreSQL', 9),
    ('    3.1.7 WhatsApp Baileys', 9),
    ('  3.2 Kerangka Berpikir', 9),
    ('  3.3 Metode Pengembangan', 9),
    ('BAB IV HASIL YANG DICAPAI', 10),
    ('  4.1 Gambaran Umum Sistem', 10),
    ('  4.2 Analisis dan Perancangan Sistem', 10),
    ('  4.3 Implementasi Sistem', 13),
    ('  4.4 Pengujian Sistem', 14),
    ('  4.5 Pembahasan Hasil', 15),
    ('  4.6 Tampilan Antarmuka Sistem', 16),
    ('BAB V PENUTUP', 18),
    ('  5.1 Kesimpulan', 18),
    ('  5.2 Saran', 18),
    ('DAFTAR PUSTAKA', 20),
    ('LAMPIRAN', 22),
    ('  Lampiran A: Logbook Kegiatan Harian', 22),
    ('  Lampiran B: GitHub Repository', 22),
    ('  Lampiran C: Screenshot Sistem', 22),
]

for text, page in toc_entries:
    # Determine style based on indentation
    if text.startswith('    '):
        style = 'TOC3'
        clean = text.strip()
    elif text.startswith('  '):
        style = 'TOC2'
        clean = text.strip()
    else:
        style = 'TOC1'
        clean = text
    
    entry = make_toc_entry(clean, page, style)
    insert_after.addnext(entry)
    insert_after = entry

print(f"Inserted {len(toc_entries)} DAFTAR ISI entries")

# Insert DAFTAR GAMBAR heading + entries
insert_after.addnext(make_heading('DAFTAR GAMBAR'))
insert_after = insert_after.getnext()

figures = [
    ('Gambar 4.2 Use Case Diagram Sistem Mimotes AI', 13),
    ('Gambar 4.3 Activity Diagram Upload Dokumen', 14),
    ('Gambar 4.4 Activity Diagram Proses Chat RAG', 15),
    ('Gambar 4.5 Entity Relationship Diagram (ERD)', 16),
    ('Gambar 4.6 Arsitektur Sistem Mimotes AI', 16),
    ('Gambar 4.7 Arsitektur RAG Pipeline', 16),
    ('Gambar 4.8 Arsitektur CRM Pipeline', 16),
    ('Gambar 4.9 Halaman Login', 16),
    ('Gambar 4.10 Dashboard Admin', 16),
    ('Gambar 4.11 Upload Dokumen', 16),
    ('Gambar 4.12 Daftar Dokumen', 16),
    ('Gambar 4.13 Chat AI', 16),
    ('Gambar 4.14 Knowledge Search', 16),
    ('Gambar 4.15 Analytics Chat', 16),
    ('Gambar 4.16 Pengaturan AI Provider', 16),
]

for title, page in figures:
    entry = make_toc_entry(title, page, 'TOC1')
    insert_after.addnext(entry)
    insert_after = entry

print(f"Inserted {len(figures)} DAFTAR GAMBAR entries")

# Insert DAFTAR TABEL heading + entries
insert_after.addnext(make_heading('DAFTAR TABEL'))
insert_after = insert_after.getnext()

tables = [
    ('Tabel 4.1 Spesifikasi Perangkat Keras dan Lunak', 13),
    ('Tabel 4.2 Hasil Pengujian Black Box', 14),
]

for title, page in tables:
    entry = make_toc_entry(title, page, 'TOC1')
    insert_after.addnext(entry)
    insert_after = entry

print(f"Inserted {len(tables)} DAFTAR TABEL entries")

doc.save(FILE)
print(f"\nSaved! {FILE}")

# Final verify
doc2 = Document(FILE)
print("\n=== FINAL VERIFICATION ===")
for i, p in enumerate(doc2.paragraphs):
    if i > 200:
        break
    text = p.text.strip()
    if not text:
        continue
    style = p.style.style_id if p.style else ''
    if style.startswith('TOC') or style == 'Heading1' or style == 'Heading2' or style == 'Heading3':
        print(f'[{i:3d}] {style:10s} | {text[:70]}')
