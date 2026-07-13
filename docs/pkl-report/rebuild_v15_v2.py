#!/usr/bin/env python3
"""
Rebuild v2: Keep template FORMATTING, only replace TEXT content.
Strategy:
1. Copy template completely
2. For each section (BAB I-V), find matching paragraphs in template
3. Replace text runs while KEEPING paragraph formatting
4. For extra paragraphs, copy format from nearest heading paragraph
"""
import os
import copy
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = 'C:/Users/SMANSA/AppData/Local/hermes/cache/documents/doc_4ce6543a497c_LAPORAN PKL (1).docx'
SOURCE = 'LAPORAN_PKL_v14_UML.docx'
OUTPUT = 'LAPORAN_PKL_v15_Template.docx'
DIAGRAMS_DIR = 'diagrams'

NEW_TITLE = "RANCANG BANGUN SISTEM CHATBOT AI BERBASIS PENGETAHUAN DENGAN RETRIEVAL-AUGMENTED GENERATION DAN PIPELINE CRM UNTUK OPTIMALISASI LAYANAN PELANGGAN"
NEW_STUDENT = "Eko Saputro"
NEW_NIM = "23215050"
NEW_SUPERVISOR = "Widianto Agung Nugroho"
OLD_STUDENT = "Moh. Arif Prasetyo"
OLD_NIM = "23215043"

print("📄 Loading...")
tpl_doc = Document(TEMPLATE)
src_doc = Document(SOURCE)

# ===== Helper: get heading level =====
def get_heading_level(para):
    name = para.style.name if para.style else ''
    if name in ('Heading 1', 'Heading1'): return 1
    if name in ('Heading 2', 'Heading2'): return 2
    if name in ('Heading 3', 'Heading3'): return 3
    return None

def classify_bab(text):
    text = text.strip()
    if 'BAB V' in text: return 'BAB_V'
    if 'BAB IV' in text: return 'BAB_IV'
    if 'BAB III' in text: return 'BAB_III'
    if 'BAB II' in text: return 'BAB_II'
    if 'BAB I' in text and 'PENDAHULUAN' not in text: return 'BAB_I'
    if 'DAFTAR PUSTAKA' in text: return 'DAFTAR_PUSTAKA'
    if 'LAMPIRAN' in text: return 'LAMPIRAN'
    return None

# ===== STEP 1: Extract source sections =====
print("\n📝 Step 1: Extracting source sections...")
src_sections = {}
current_key = None
current_paras = []

for para in src_doc.paragraphs:
    text = para.text.strip()
    level = get_heading_level(para)
    
    if level == 1:
        if current_key:
            src_sections[current_key] = current_paras
        current_paras = []
        current_key = classify_bab(text)
        continue
    
    if current_key:
        # Copy paragraph with all formatting
        current_paras.append({
            'text': text,
            'style_name': para.style.name if para.style else 'Normal',
            'bold': any(r.bold for r in para.runs if r.bold),
            'italic': any(r.italic for r in para.runs if r.italic),
            'element': copy.deepcopy(para._element),
        })

if current_key:
    src_sections[current_key] = current_paras

for key, paras in src_sections.items():
    print(f"  {key}: {len(paras)} paragraphs")

# ===== STEP 2: Find template section boundaries =====
print("\n📝 Step 2: Finding template sections...")

tpl_sections = {}
for i, para in enumerate(tpl_doc.paragraphs):
    level = get_heading_level(para)
    if level == 1:
        key = classify_bab(para.text)
        if key:
            tpl_sections[key] = i

print(f"  Template sections: {tpl_sections}")

# ===== STEP 3: Replace content paragraph by paragraph =====
print("\n📝 Step 3: Replacing content...")

# Sort sections by position
sorted_sections = sorted(tpl_sections.items(), key=lambda x: x[1])

for sec_i, (key, start_idx) in enumerate(sorted_sections):
    if key not in src_sections:
        print(f"  ⚠️ No source for {key}")
        continue
    
    # Find end of section
    end_idx = len(tpl_doc.paragraphs)
    for j in range(sec_i + 1, len(sorted_sections)):
        _, next_idx = sorted_sections[j]
        end_idx = next_idx
        break
    
    section_paras = list(tpl_doc.paragraphs[start_idx:end_idx])
    src_paras = src_sections[key]
    
    # Skip heading paragraph
    content_start = 1
    if key == 'BAB_I' and len(section_paras) > 1:
        if section_paras[1].text.strip() == 'PENDAHULUAN':
            content_start = 2
    
    available = len(section_paras) - content_start
    to_place = src_paras
    
    # Replace existing paragraphs
    for j, src_p in enumerate(to_place):
        if j < available:
            target = section_paras[content_start + j]
            # Clear all runs and set new text
            for run in target.runs:
                run.text = ''
            if target.runs:
                target.runs[0].text = src_p['text']
            elif src_p['text']:
                target.add_run(src_p['text'])
    
    # If source has MORE paragraphs than template slots, we need to INSERT
    excess = len(to_place) - available
    if excess > 0:
        # Insert excess paragraphs after the last replaced paragraph
        last_replaced = section_paras[content_start + available - 1] if available > 0 else section_paras[0]
        last_elem = last_replaced._element
        parent = last_elem.getparent()
        insert_pos = list(parent).index(last_elem) + 1
        
        for k in range(excess):
            src_p = to_place[available + k]
            # Clone the formatting from the last paragraph in section
            template_para = section_paras[-1]
            new_elem = copy.deepcopy(template_para._element)
            
            # Clear text and set new text
            for t_elem in new_elem.iter(qn('w:t')):
                t_elem.text = ''
            # Set text in first w:r/w:t
            runs = new_elem.findall(f'.//{qn("w:r")}')
            if runs:
                t_elems = runs[0].findall(qn('w:t'))
                if t_elems:
                    t_elems[0].text = src_p['text']
                else:
                    t = etree.SubElement(runs[0], qn('w:t'))
                    t.text = src_p['text']
                    t.set(qn('xml:space'), 'preserve')
            
            parent.insert(insert_pos + k, new_elem)
        
        print(f"  ✅ {key}: {available} replaced + {excess} inserted")
    else:
        print(f"  ✅ {key}: {len(to_place)}/{available} replaced")

# ===== STEP 4: Fix front matter text =====
print("\n📝 Step 4: Fixing front matter...")
for para in tpl_doc.paragraphs:
    if "SISTEM PAKAR" in para.text:
        for run in para.runs:
            if "SISTEM PAKAR" in run.text or "RANCANG BANGUN" in run.text:
                run.text = ""
        for run in para.runs:
            if not run.text.strip():
                run.text = NEW_TITLE
                break
    for run in para.runs:
        if OLD_STUDENT in run.text:
            run.text = run.text.replace(OLD_STUDENT, NEW_STUDENT)
        if OLD_NIM in run.text:
            run.text = run.text.replace(OLD_NIM, NEW_NIM)
        if "Nama Dosen Pembimbing" in run.text:
            run.text = run.text.replace("Nama Dosen Pembimbing", NEW_SUPERVISOR)
        if "NIPY." in run.text:
            run.text = run.text.replace("NIPY.", "")

print("  ✅ Done")

# ===== STEP 5: Clean old Daftar Gambar/Tabel =====
print("\n📝 Step 5: Cleaning Daftar Gambar/Tabel...")

for i, para in enumerate(tpl_doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    # Clean Daftar Tabel entries
    if style == 'table of figures' and text.startswith('Tabel 4.') and 'Penyakit' in text:
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = 'Tabel 4.1 Dokumen dalam Sistem Mimotes AI'
    
    if style == 'table of figures' and text.startswith('Tabel 4.') and 'Rule' in text:
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = 'Tabel 4.2 Chat Session dalam Sistem Mimotes AI'
    
    # Clean Daftar Gambar entries - replace old with new
    if style == 'table of figures' and text.startswith('Gambar 4.'):
        # Map old captions to new
        old_to_new = {
            "Pengumpulan Data": "Use Case Diagram Sistem Mimotes AI",
            "Pohon Keputusan": "Activity Diagram Upload Dokumen",
            "Use Case\t": "Activity Diagram Proses Chat RAG",
            "Activity Diagram Proses Diag": "Sequence Diagram Chat RAG",
            "Sequence diagram Proses": "ERD Domain Identity & Workspace",
            "Class Diagram": "ERD Domain RAG & Knowledge Base",
            "Arsitektur Sistem SIPATAN": "ERD Domain Chat & CRM",
            "Flowchart Proses": "ERD Domain Billing & Configuration",
            "Halaman Beranda": "ERD Ringkasan — Seluruh Relasi",
            "Halaman Tentang": "Arsitektur Sistem Mimotes AI",
            "Halaman Cara Kerja": "Arsitektur RAG Pipeline",
            "Halaman Deteksi": "Arsitektur CRM Pipeline",
            "Halaman Panel Admin": "Halaman Login",
            "Implementasi database": "Dashboard Admin",
            "Input Keyakinan": "Upload Dokumen",
            "Hasil Diag": "Daftar Dokumen",
        }
        for old_kw, new_txt in old_to_new.items():
            if old_kw in text:
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = f"Gambar {new_txt}"
                break

print("  ✅ Done")

# ===== SAVE =====
print("\n💾 Saving...")
tpl_doc.save(OUTPUT)

# ===== STEP 6: Insert images =====
print("\n🖼️ Step 6: Inserting images...")

doc2 = Document(OUTPUT)

CAPTION_IMG = {
    "Use Case Diagram": "usecase.png",
    "Activity Diagram Upload": "activity-upload.png",
    "Activity Diagram Proses Chat": "activity-chat.png",
    "Sequence Diagram Chat RAG": "sequence-chat-rag.png",
    "ERD Domain Identity": "erd-a-identity.png",
    "ERD Domain RAG": "erd-b-rag.png",
    "ERD Domain Chat": "erd-c-chat-crm.png",
    "ERD Domain Billing": "erd-d-billing.png",
    "ERD Ringkasan": "erd-summary.png",
    "Arsitektur Sistem Mimotes": "architecture.png",
    "Arsitektur RAG": "rag-pipeline.png",
    "Arsitektur CRM": "crm-pipeline.png",
}

img_inserted = 0
for i, para in enumerate(doc2.paragraphs):
    text = para.text.strip()
    if not text.startswith('Gambar 4.'):
        continue
    
    for keyword, img_file in CAPTION_IMG.items():
        if keyword in text:
            img_path = os.path.join(DIAGRAMS_DIR, img_file)
            if not os.path.exists(img_path):
                continue
            try:
                new_para = doc2.add_paragraph()
                new_para.alignment = 1  # CENTER
                new_para.add_run().add_picture(img_path, width=Inches(5.5))
                parent = para._element.getparent()
                pos = list(parent).index(para._element)
                new_elem = new_para._element
                parent.remove(new_elem)
                parent.insert(pos, new_elem)
                img_inserted += 1
                print(f"  ✅ {img_file}")
            except Exception as e:
                print(f"  ❌ {img_file}: {e}")
            break

doc2.save(OUTPUT)
print(f"\n  Images: {img_inserted}")

# ===== VERIFY =====
doc3 = Document(OUTPUT)
total_imgs = sum(
    len(run._element.findall(qn('w:drawing')))
    for para in doc3.paragraphs
    for run in para.runs
)

# Check for old content
old_found = False
for para in doc3.paragraphs:
    if 'Sistem Pakar Deteksi Hama' in para.text:
        old_found = True
        break

print(f"\n📊 Final:")
print(f"  Paragraphs: {len(doc3.paragraphs)}")
print(f"  Images: {total_imgs}")
print(f"  Old content: {'⚠️ FOUND' if old_found else '✅ Clean'}")
print(f"  Size: {os.path.getsize(OUTPUT)/1024:.0f} KB")
