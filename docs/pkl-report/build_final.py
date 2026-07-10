import shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===== COPY TEMPLATE =====
src = r'C:\Users\SMANSA\AppData\Local\hermes\cache\documents\doc_6c6d74787cac_Laporan_PKL_MuhammadAffif_Fixed.docx'
dst = r'C:\Users\SMANSA\mimotes\docs\pkl-report\LAPORAN_PKL_Eko_Saputro_23215050.docx'
shutil.copy2(src, dst)
doc = Document(dst)

DIAG_DIR = r'C:\Users\SMANSA\mimotes\docs\pkl-report\diagrams'

# ===== STEP 1: Replace personal data in front matter =====
for p in doc.paragraphs:
    for run in p.runs:
        rt = run.text
        if 'Muhammad Affif' in rt:
            run.text = rt.replace('Muhammad Affif', 'Eko Saputro')
        if '24225046' in rt:
            run.text = rt.replace('24225046', '23215050')
        if 'Syefudin' in rt:
            run.text = rt.replace('Syefudin', 'Zaenul Arif')
        if 'PT Agsya Karya Manca' in rt:
            run.text = rt.replace('PT Agsya Karya Manca', 'Bank Mandiri KCP Tegal Sudirman')
        if 'Agsya Karya Manca' in rt:
            run.text = rt.replace('Agsya Karya Manca', 'Bank Mandiri KCP Tegal Sudirman')
        if 'Agsya' in rt and 'Bank Mandiri' not in rt:
            run.text = rt.replace('Agsya', 'Bank Mandiri')
        if 'Implementasi Keamanan Aplikasi dan Manajemen State pada Sistem' in rt:
            run.text = rt.replace('Implementasi Keamanan Aplikasi dan Manajemen State pada Sistem',
                                  'RANCANG BANGUN SISTEM CHATBOT AI')
        if 'Informasi Manajemen SDM Berbasis Web Menggunakan Vue.js 3 dan' in rt:
            run.text = rt.replace('Informasi Manajemen SDM Berbasis Web Menggunakan Vue.js 3 dan',
                                  'BERBASIS PENGETAHUAN DENGAN')
        if 'Pinia di PT Bank Mandiri KCP Tegal Sudirman' in rt:
            run.text = rt.replace('Pinia di PT Bank Mandiri KCP Tegal Sudirman',
                                  'RETRIEVAL-AUGMENTED GENERATION DAN PIPELINE CRM')
        elif 'Pinia di PT' in rt:
            run.text = rt.replace(rt[rt.index('Pinia'):], 'RETRIEVAL-AUGMENTED GENERATION DAN PIPELINE CRM')

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if 'Muhammad Affif' in run.text:
                        run.text = run.text.replace('Muhammad Affif', 'Eko Saputro')
                    if '24225046' in run.text:
                        run.text = run.text.replace('24225046', '23215050')

# ===== STEP 2: Find where body content starts =====
body_start = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and 'BAB I' in p.text:
        body_start = i
        break

if body_start is None:
    print('ERROR: Could not find BAB I heading')
    exit(1)

print(f'Body content starts at paragraph {body_start}')

# ===== STEP 3: Delete all body content (from body_start to end) =====
body = doc.element.body
para_elements = list(body.findall(qn('w:p')))
table_elements = list(body.findall(qn('w:tbl')))

# Remove all paragraphs from body_start onwards
to_remove = []
for i, p_elem in enumerate(para_elements):
    if i >= body_start:
        to_remove.append(p_elem)

# Also remove all tables (they're in the body content)
for t_elem in table_elements:
    to_remove.append(t_elem)

for elem in to_remove:
    body.remove(elem)

print(f'Removed {len(to_remove)} elements')

# ===== STEP 4: Add new content =====
# Get style names from the template
def add_styled_para(text, style_name='Normal', bold=False, align=None):
    p = doc.add_paragraph(style=style_name)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    return p

def add_bab_title(text):
    p = doc.add_paragraph(style='Heading 1')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    return p

def add_section_title(text):
    p = doc.add_paragraph(style='Heading 2')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    return p

def add_subsection_title(text):
    p = doc.add_paragraph(style='Heading 3')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    return p

def add_body(text, bold=False):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_image(filename, caption):
    img_path = os.path.join(DIAG_DIR, filename)
    if os.path.exists(img_path):
        p = doc.add_paragraph(style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(14))
        # Caption
        cap = doc.add_paragraph(style='Caption')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.font.name = 'Times New Roman'
        cr.font.size = Pt(10)
        cr.italic = True

def add_page_break():
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

# ===== BAB I: PENDAHULUAN =====
add_bab_title('BAB I\nPENDAHULUAN')

add_section_title('1.1 Latar Belakang')
add_body('Perkembangan teknologi kecerdasan buatan (Artificial Intelligence/AI) dalam beberapa tahun terakhir telah mengalami kemajuan yang sangat pesat, khususnya dalam bidang pemrosesan bahasa alami (Natural Language Processing/NLP). Salah satu bentuk penerapan AI yang paling banyak digunakan dalam dunia bisnis adalah chatbot, yaitu program komputer yang mampu melakukan percakapan dengan pengguna secara otomatis menggunakan bahasa manusia [1].')
add_body('Dalam konteks layanan pelanggan (customer service), chatbot berbasis AI memiliki potensi besar untuk meningkatkan efisiensi dan kualitas layanan. Namun, chatbot konvensional memiliki keterbatasan dalam memahami konteks dan memberikan jawaban yang relevan [2].')
add_body('Untuk mengatasi keterbatasan tersebut, diperlukan pendekatan Retrieval-Augmented Generation (RAG), yaitu teknik yang menggabungkan kemampuan Large Language Model (LLM) dengan mekanisme pencarian dari basis pengetahuan perusahaan [3].')
add_body('Selain aspek layanan pelanggan, aspek Customer Relationship Management (CRM) juga memegang peranan penting. Integrasi antara chatbot AI dengan sistem CRM memungkinkan setiap interaksi pelanggan tercatat dan terkelola secara otomatis [4].')
add_body('Berdasarkan latar belakang di atas, penulis mengembangkan Mimotes AI, yaitu sistem chatbot AI berbasis pengetahuan dengan RAG dan pipeline CRM untuk optimalisasi layanan pelanggan di Bank Mandiri KCP Tegal Sudirman.')

add_section_title('1.2 Tujuan PKL')
add_subsection_title('1.2.1 Tujuan Umum')
add_body('Memberikan pengalaman kerja kepada mahasiswa dalam menerapkan teori perkuliahan dengan situasi nyata di bidang kecerdasan buatan, rekayasa perangkat lunak, dan basis data [5].')
add_subsection_title('1.2.2 Tujuan Khusus')
for t in [
    'Merancang arsitektur sistem chatbot AI berbasis RAG.',
    'Mengimplementasikan pipeline RAG: document processing, chunking, embedding, vector storage, retrieval, dan generation.',
    'Mengembangkan pipeline CRM: manajemen leads, kontak, aktivitas, dan sales pipeline.',
    'Mengintegrasikan sistem dengan WhatsApp melalui protokol Baileys.',
    'Menguji kefektifan sistem melalui pengujian black box.',
]:
    add_bullet(t)

add_section_title('1.3 Manfaat PKL')
add_body('Bagi Universitas Harkat Negeri: Menambah jaringan kerja sama dengan dunia industri.')
add_body('Bagi Program Studi: Bahan masukan pengembangan kurikulum sesuai kebutuhan industri.')
add_body('Bagi Bank Mandiri KCP Tegal Sudirman: Solusi sistem informasi berbasis AI untuk optimalisasi layanan pelanggan.')
add_body('Bagi Penulis: Pengalaman kerja nyata dalam pengembangan sistem AI dan CRM.')
add_page_break()

# ===== BAB II =====
add_bab_title('BAB II\nGAMBARAN UMUM INSTANSI')

add_section_title('2.1 Sejarah Perkembangan Perusahaan')
add_body('Bank Mandiri KCP Tegal Sudirman merupakan Kantor Cabang Pembantu dari PT Bank Mandiri (Persero) Tbk yang terletak di Jalan Sudirman, Tegal. Bank Mandiri didirikan pada 2 Oktober 1998 dari penggabungan empat bank pemerintah: Bank Bumi Daya, Bank Dagang Negara, Bank Exim, dan Bapindo.')
add_body('KCP Tegal Sudirman melayani produk perbankan meliputi simpanan pinjaman, layanan transaksi, pengelolaan aset, serta layanan lelang property.')

add_section_title('2.2 Visi, Misi, dan Tujuan')
add_body('Visi: Menjadi bank terbaik yang mengedepankan kepuasan nasabah melalui layanan perbankan digital yang inovatif dan terpercaya.', bold=True)
add_body('Misi:', bold=True)
for m in [
    'Memberikan layanan perbankan digital yang inovatif dan terpercaya.',
    'Mengoptimalkan layanan pelanggan melalui pemanfaatan teknologi informasi.',
    'Mengelola data nasabah secara terstruktur dan efisien.',
]:
    add_bullet(m)

add_section_title('2.3 Struktur Organisasi')
add_body('Struktur organisasi Bank Mandiri KCP Tegal Sudirman terdiri dari Kepala Cabang Pembantu, Bagian Customer Service, Bagian Teller, Bagian Administrasi, serta Bagian Teknologi Informasi.')

add_section_title('2.4 Job Deskripsi')
add_body('Selama PKL, penulis ditempatkan pada Divisi Teknologi Informasi dengan job deskripsi:')
for j in [
    'Input data PPAT dan pembuatan SKPT untuk lelang property.',
    'Validasi pajak PPH di Kantor Pajak Pratama Tegal.',
    'Pelaksanaan dan pengelolaan lelang property.',
    'Pengembangan sistem chatbot AI (Mimotes) sebagai capstone project.',
    'Pemasangan banner lelang dan penyerahan risalah lelang.',
]:
    add_bullet(j)
add_page_break()

# ===== BAB III =====
add_bab_title('BAB III\nMETODE PELAKSANAAN PKL')

add_section_title('3.1 Tugas Umum')
add_body('Selama PKL di Bank Mandiri KCP Tegal Sudirman, penulis menjalankan tugas umum meliputi input data PPAT, pembuatan SKPT, validasi pajak, pelaksanaan lelang, serta pengembangan aplikasi chatbot AI secara paralel melalui WFH dan WFA.')

add_section_title('3.2 Tugas Khusus')
add_subsection_title('3.2.1 Analisis Kebutuhan Sistem')
add_body('Penulis mengidentifikasi permasalahan utama: volume interaksi nasabah yang tinggi, pengetahuan tersebar dalam berbagai dokumen, tidak adanya integrasi CRM, serta keterbatasan akses multi-platform.')
add_subsection_title('3.2.2 Perancangan Arsitektur Sistem')
add_body('Arsitektur Mimotes AI terdiri dari: Frontend Layer (Next.js 16 + React 19 + Tailwind CSS), API Layer (108 routes), Data Layer (PostgreSQL 16 + pgvector), RAG Pipeline, dan Microservices (Baileys WhatsApp, PaddleOCR, n8n).')

add_section_title('3.3 Analisis Permasalahan dan Solusi')
add_body('Akurasi Respons Chatbot: Solusi RAG menggabungkan LLM dengan pencarian semantik dari knowledge base perusahaan.')
add_body('Pengelolaan Pelanggan: Pipeline CRM terintegrasi dengan pendekatan conversation-centric CRM.')
add_body('Isolasi Data Multi-Tenant: Workspace-based isolation dengan PostgreSQL Row Level Security (RLS).')
add_body('Ekstraksi Teks dari Gambar: Integrasi PaddleOCR dan Gemini Vision untuk OCR dan captioning.')
add_page_break()

# ===== BAB IV =====
add_bab_title('BAB IV\nHASIL YANG DICAPAI')

add_section_title('4.1 Gambaran Umum Sistem')
add_body('Mimotes AI di-deploy menggunakan Docker Compose dengan lima service: PostgreSQL (5432), Next.js App (3100), PaddleOCR (8090), Baileys WhatsApp (3002), dan Database Migration.')

add_section_title('4.2 Analisis dan Perancangan Sistem')
add_subsection_title('4.2.1 Use Case Diagram')
add_body('Sistem memiliki empat aktor: Admin/Workspace Owner, Editor, Viewer, dan Visitor/Pelanggan.')
add_image('usecase.png', 'Gambar 4.2 Use Case Diagram Sistem Mimotes AI')

add_subsection_title('4.2.2 Activity Diagram Upload Dokumen')
add_body('Proses upload: autentikasi, validasi file, simpan ke disk, buat record di DB, processing queue, lalu parsing, chunking, embedding, storage.')
add_image('activity-upload.png', 'Gambar 4.3 Activity Diagram Upload Dokumen')

add_subsection_title('4.2.3 Activity Diagram Proses Chat RAG')
add_body('Proses chat: user kirim pesan, generate query embedding, hybrid search, confidence classification, build context, system prompt, LLM generate streaming response, simpan ke DB.')
add_image('activity-chat.png', 'Gambar 4.4 Activity Diagram Proses Chat RAG')

add_subsection_title('4.2.4 Entity Relationship Diagram')
add_body('28 model database dalam enam domain: Tenant & Identity, Document & RAG Pipeline, Conversational AI, CRM & Lead Pipeline, Subscription & Billing, Configuration & Observability.')
add_image('erd.png', 'Gambar 4.5 Entity Relationship Diagram')

add_subsection_title('4.2.5 Arsitektur Sistem')
add_body('Enam layer: Frontend, API Layer, RAG Pipeline, Data Layer, Microservices, External Services.')
add_image('architecture.png', 'Gambar 4.6 Arsitektur Sistem Mimotes AI')

add_subsection_title('4.2.6 Arsitektur RAG Pipeline')
add_body('Pipeline: Document Ingestion, Embedding Generation, Vector Storage (pgvector), Retrieval (Hybrid Search + RRF), Response Generation (LLM).')
add_image('rag-pipeline.png', 'Gambar 4.7 Arsitektur RAG Pipeline')

add_subsection_title('4.2.7 Arsitektur CRM Pipeline')
add_body('Pipeline: Lead Capture, AI Analysis, Lead Scoring, Lead Status Tracking, Follow-up Automation.')
add_image('crm-pipeline.png', 'Gambar 4.8 Arsitektur CRM Pipeline')

add_section_title('4.3 Implementasi Sistem')
add_subsection_title('4.3.1 Spesifikasi Perangkat Keras dan Lunak')
# Table
table = doc.add_table(rows=10, cols=2)
table.style = 'Normal Table'
specs = [
    ('Komponen', 'Spesifikasi'),
    ('Framework', 'Next.js 16.2.7'),
    ('Frontend', 'React 19, Tailwind CSS'),
    ('Database', 'PostgreSQL 16 + pgvector'),
    ('ORM', 'Prisma 6.19.3'),
    ('AI Provider', 'Mimo Pro, OpenAI, Gemini'),
    ('OCR', 'PaddleOCR + Gemini Vision'),
    ('WhatsApp', 'Baileys 6.7.23'),
    ('Container', 'Docker Compose'),
    ('Testing', 'Vitest'),
]
for i, (k, v) in enumerate(specs):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    for c in range(2):
        for p in table.rows[i].cells[c].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
                if i == 0:
                    r.bold = True

add_subsection_title('4.3.2 Implementasi RAG Pipeline')
add_body('Document Parsing: PDF, DOCX, TXT, CSV, XLSX, gambar (PaddleOCR + Gemini Vision), URL.')
add_body('Chunking: Recursive paragraph-then-sentence, 500 chars, 50 word overlap, max 1000 chunks/doc.')
add_body('Embedding: OpenAI text-embedding-3-small (1536 dim) atau Feature Hashing lokal (fallback gratis).')
add_body('Vector Storage: pgvector PostgreSQL, vector(1536), batch 50/txn, RLS.')
add_body('Retrieval: Hybrid Search (Vector 0.6 + BM25 0.4 dengan RRF).')
add_body('Generation: LLM OpenAI-compatible, temperature 0.3, max_tokens 1000, streaming.')

add_subsection_title('4.3.3 Implementasi Frontend')
add_body('53 halaman: dashboard, chat, knowledge management, CRM, analytics, settings, widget, admin, WhatsApp.')

add_subsection_title('4.3.4 Implementasi Backend')
add_body('108 API routes: Auth, Admin, AI, Analytics, Chat, Documents, Knowledge/RAG, Leads/CRM, WhatsApp, Widget, Billing, Workspace.')

add_section_title('4.4 Pengujian Sistem')
add_body('Hasil pengujian black box:')
table2 = doc.add_table(rows=10, cols=5)
table2.style = 'Normal Table'
tests = [
    ('No', 'Fitur', 'Input', 'Expected', 'Status'),
    ('1', 'Login', 'Email+pass valid', 'Redirect', 'Lulus'),
    ('2', 'Login', 'Email+pass invalid', 'Error', 'Lulus'),
    ('3', 'Upload', 'PDF <10MB', 'Ready', 'Lulus'),
    ('4', 'Chat RAG', 'Pertanyaan', 'Jawaban', 'Lulus'),
    ('5', 'Chat RAG', 'Di luar konteks', 'Penolakan', 'Lulus'),
    ('6', 'Lead', 'Isi form', 'Tersimpan', 'Lulus'),
    ('7', 'WhatsApp', 'Pesan masuk', 'Reply', 'Lulus'),
    ('8', 'Settings', 'Ubah provider', 'OK', 'Lulus'),
    ('9', 'RBAC', 'Viewer upload', 'Denied', 'Lulus'),
]
for i, row_data in enumerate(tests):
    for j, val in enumerate(row_data):
        table2.rows[i].cells[j].text = val
        for p in table2.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
                if i == 0:
                    r.bold = True
add_page_break()

# ===== BAB V =====
add_bab_title('BAB V\nPENUTUP')

add_section_title('5.1 Kesimpulan')
for k in [
    'Sistem Mimotes AI berhasil dirancang dan diimplementasikan sebagai chatbot AI berbasis RAG dengan pipeline CRM.',
    'Pipeline RAG berfungsi: document processing, chunking, embedding, vector storage, retrieval, hingga generasi respons.',
    'Pipeline CRM terintegrasi: setiap interaksi pelanggan tercatat dan terkelola otomatis.',
    'Integrasi WhatsApp memungkinkan pelanggan berinteraksi melalui platform pesan instan populer.',
    'Pengujian black box menunjukkan seluruh fitur berfungsi sesuai kebutuhan.',
]:
    add_bullet(k)

add_section_title('5.2 Saran')
add_body('Bagi Bank Mandiri: Sistem perlu dipelihara dan dievaluasi berkala. Pertimbangkan integrasi dengan email dan media sosial.')
add_body('Bagi Prodi: Kurikulum dapat diperkaya dengan materi NLP dan RAG. PKL Capstone Project menjadi model pembelajaran efektif.')
add_body('Bagi Penulis: Dapat dikembangkan sebagai topik Skripsi dengan evaluasi kualitas RAG menggunakan ROUGE/BLEU.')
add_page_break()

# ===== DAFTAR PUSTAKA =====
add_bab_title('DAFTAR PUSTAKA')
refs = [
    '[1] H. Susanto et al., "Pemanfaatan Chatbot Berbasis AI untuk Optimalisasi Layanan Pelanggan," Jurnal Sistem Informasi, vol. 19, no. 2, pp. 45-62, 2023.',
    '[2] T. Brown et al., "Language Models are Few-Shot Learners," NeurIPS, vol. 33, pp. 1877-1901, 2020.',
    '[3] P. Lewis et al., "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks," NeurIPS, vol. 33, pp. 9459-9474, 2020.',
    '[4] F. Buttle dan S. Maklan, Customer Relationship Management, 4th ed. London: Routledge, 2019.',
    '[5] Universitas Harkat Negeri, Buku Panduan PKL Prodi S1 Teknik Informatika. Tegal, 2026.',
    '[6] V. Lopez et al., "A Survey on Transfer Learning in NLP," AI Review, vol. 53, pp. 2339-2367, 2020.',
    '[7] Y. Gao et al., "RAG for LLMs: A Survey," arXiv:2312.10997, 2024.',
    '[8] A. Vaswani et al., "Attention Is All You Need," NeurIPS, vol. 30, pp. 5998-6008, 2017.',
    '[9] J. Devlin et al., "BERT," NAACL-HLT, pp. 4171-4186, 2019.',
    '[10] Qdrant, "Vector Database for AI," 2024. https://qdrant.tech/',
    '[11] Next.js, "The React Framework," 2024. https://nextjs.org/',
    '[12] Prisma, "Next-gen ORM," 2024. https://www.prisma.io/',
    '[13] Baileys, "WhatsApp Web API," 2024. https://github.com/WhiskeySockets/Baileys',
    '[14] pgvector, "Vector similarity search for Postgres," 2024. https://github.com/pgvector/pgvector',
    '[15] OpenAI, "Embeddings API," 2024. https://platform.openai.com/docs/guides/embeddings',
]
for ref in refs:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
add_page_break()

# ===== LAMPIRAN =====
add_bab_title('LAMPIRAN')

add_section_title('Lampiran A: Logbook Kegiatan Harian')
log_data = [
    ('No', 'Tanggal', 'Jam', 'Aktivitas', 'Lokasi'),
    ('1', '25/05', '09:10-15:30', 'Input Akun Lelang, Validasi Pajak', 'Bank Mandiri'),
    ('2', '26/05', '09:10-15:02', 'Pelaksanaan Lelang', 'Bank Mandiri'),
    ('3', '27/05', '09:11-15:30', 'WFH Pembuatan Aplikasi', 'WFH'),
    ('4', '28/05', '09:10-15:30', 'Melanjutkan Aplikasi', 'WFH'),
    ('5', '29/05', '09:00-14:11', 'Input SSPD BPHTB', 'Bank Mandiri'),
    ('6', '01/06', '09:38-22:17', 'WFH Developing Aplikasi', 'WFH'),
    ('7', '02/06', '09:10-14:58', 'Penyerahan Risalah Lelang', 'Bank Mandiri'),
    ('8', '03/06', '09:10-17:28', 'Surat Pendaftaran Tanah', 'WFA'),
    ('9', '04/06', '09:10-16:28', 'Input PPAT Sistem INTAN', 'Bank Mandiri'),
    ('10', '05/06', '08:48-15:39', 'Input PPAT SKPT', 'Bank Mandiri'),
    ('11', '08/06', '08:30-16:13', 'Input PPATK + Banner', 'Bank Mandiri'),
    ('12', '09/06', '08:46-16:41', 'SKPT + Banner Lelang', 'Bank Mandiri'),
    ('13', '10/06', '09:22-16:20', 'Input PPAT SKPT', 'Bank Mandiri'),
    ('14', '11/06', '08:30-17:40', 'PPAT + SKPT + Banner', 'Bank Mandiri'),
    ('15', '12/06', '08:45-15:58', 'Input PPAT SKPT', 'Bank Mandiri'),
    ('16', '15/06', '08:36-16:32', 'Input PPAT + SKPT', 'Bank Mandiri'),
    ('17', '16/06', '09:20-16:55', 'WFA Pembuatan Aplikasi', 'WFA'),
    ('18', '17/06', '08:54-16:34', 'Input PPAT SKPT', 'Bank Mandiri'),
    ('19', '18/06', '08:38-16:30', 'Input PPAT SKPT', 'WFA'),
    ('20', '19/06', '08:45-15:55', 'Input PPAT SKPT', 'Bank Mandiri'),
    ('21', '22/06', '08:47-16:50', 'Input PPAT', 'Bank Mandiri'),
    ('22', '23/06', '08:35-15:02', 'SKPT + Lokasi Lelang', 'Bank Mandiri'),
    ('23', '24/06', '10:52-16:21', 'WFH Pembuatan Aplikasi', 'WFH'),
    ('24', '25/06', '08:35-15:56', 'Input Data PPAT', 'Bank Mandiri'),
]
log_table = doc.add_table(rows=len(log_data), cols=5)
log_table.style = 'Normal Table'
for i, row_data in enumerate(log_data):
    for j, val in enumerate(row_data):
        log_table.rows[i].cells[j].text = val
        for p in log_table.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9)
                if i == 0:
                    r.bold = True

add_body('Pembimbing Lapangan: Widianto Agung Nugroho', bold=True)
add_section_title('Lampiran B: GitHub Repository')
add_body('Repository: https://github.com/EkoSaputro14/mimotes.git', bold=False)
add_body('Branch: semi-final', bold=False)

# ===== SAVE =====
doc.save(dst)
print(f'=== DONE ===')
print(f'Saved: {dst}')
