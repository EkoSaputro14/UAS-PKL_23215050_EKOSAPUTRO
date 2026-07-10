"""
Fix academic tone — avoid implying Bank Mandiri officially uses the system.
Based on dosen pembimbing advice.
"""
from docx import Document
import os

SRC = r'C:\Users\SMANSA\mimotes\docs\pkl-report\LAPORAN_PKL_Eko_Saputro_23215050.docx'
doc = Document(SRC)

changes = []

def fix(p, old, new, label):
    runs = p.runs
    if not runs:
        return
    merged = ''.join(r.text for r in runs)
    if old in merged:
        merged = merged.replace(old, new)
        runs[0].text = merged
        for r in runs[1:]:
            r.text = ''
        changes.append(f'  ✅ {label}')

# ============================================
# FIX 1: BAB I Latar Belakang (para 210)
# ============================================
# BEFORE: "penulis mengembangkan Mimotes AI, yaitu sistem chatbot AI berbasis 
#          pengetahuan dengan arsitektur RAG dan pipeline CRM untuk optimalisasi 
#          layanan pelanggan di Bank Mandiri KCP Tegal Sudirman."
# AFTER:  "penulis mengembangkan Mimotes AI, yaitu prototipe sistem chatbot AI 
#          berbasis pengetahuan dengan arsitektur RAG dan pipeline CRM yang 
#          dikembangkan selama kegiatan Praktik Kerja Lapangan untuk mendukung 
#          optimalisasi layanan pelanggan."
fix(doc.paragraphs[210],
    'penulis mengembangkan Mimotes AI, yaitu sistem chatbot AI berbasis pengetahuan dengan arsitektur RAG dan pipeline CRM untuk optimalisasi layanan pelanggan di Bank Mandiri KCP Tegal Sudirman.',
    'penulis mengembangkan Mimotes AI, yaitu prototipe sistem chatbot AI berbasis pengetahuan dengan arsitektur RAG dan pipeline CRM yang dikembangkan selama kegiatan Praktik Kerja Lapangan untuk mendukung optimalisasi layanan pelanggan.',
    'BAB I: "sistem untuk Bank Mandiri" → "prototipe selama PKL"')

# ============================================
# FIX 2: Manfaat PKL Bagi Bank Mandiri (para 230)
# ============================================
# BEFORE: "Mendapatkan solusi sistem informasi berbasis AI yang dapat membantu 
#          optimalisasi layanan pelanggan..."
# AFTER:  "Menjadi objek studi kelayakan penerapan sistem informasi berbasis AI 
#          untuk optimalisasi layanan pelanggan..."
fix(doc.paragraphs[230],
    'Mendapatkan solusi sistem informasi berbasis AI yang dapat membantu optimalisasi layanan pelanggan dan pengelolaan data pelanggan secara terstruktur dan efisien.',
    'Menjadi objek studi kelayakan penerapan sistem informasi berbasis AI untuk optimalisasi layanan pelanggan dan pengelolaan data pelanggan secara terstruktur dan efisien.',
    'Manfaat PKL: "Mendapatkan solusi" → "Menjadi objek studi kelayakan"')

# ============================================
# FIX 3: Saran Bagi Bank Mandiri (para 386)
# ============================================
# BEFORE: "Sistem Mimotes AI yang telah dikembangkan perlu terus dipelihara 
#          dan ditingkatkan fiturnya sesuai dengan kebutuhan bisnis yang berkembang.
#          Dilakukan evaluasi berkala..."
# AFTER:  "Apabila sistem ini hendak diadopsi secara produksi, diperlukan 
#          evaluasi kelayakan lebih lanjut, termasuk uji coba terbatas pada 
#          lingkungan staging sebelum deployment penuh. Evaluasi berkala..."
fix(doc.paragraphs[386],
    'Sistem Mimotes AI yang telah dikembangkan perlu terus dipelihara dan ditingkatkan fiturnya sesuai dengan kebutuhan bisnis yang berkembang. Dilakukan evaluasi berkala terhadap kualitas respons chatbot dan pipeline CRM untuk memastikan sistem tetap optimal dalam melayani nasabah. Pertimbangkan untuk mengintegrasikan sistem dengan platform lain seperti email, media sosial, atau sistem CRM eksternal yang sudah ada.',
    'Apabila sistem ini hendak diadopsi secara produksi, diperlukan evaluasi kelayakan lebih lanjut, termasuk uji coba terbatas pada lingkungan staging sebelum deployment penuh. Evaluasi berkala terhadap kualitas respons chatbot dan pipeline CRM perlu dilakukan untuk memastikan sistem memenuhi standar operasional bank. Selain itu, pertimbangkan integrasi dengan platform lain seperti email, media sosial, atau sistem CRM eksternal yang sudah ada.',
    'Saran: "perlu dipelihara" → "apabila hendak diadopsi, perlu evaluasi kelayakan"')

# ============================================
# SAVE
# ============================================
doc.save(SRC)

print(f'Applied {len(changes)} academic tone fixes:')
for c in changes:
    print(c)
print(f'\nFile: {SRC}')
print(f'Size: {os.path.getsize(SRC) / 1024:.0f} KB')
