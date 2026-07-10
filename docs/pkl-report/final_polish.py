"""
FINAL DOCUMENT POLISHING — comprehensive audit and fix.
No content changes, only formatting, consistency, and visual fixes.
"""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import os

FILE = 'LAPORAN_PKL_v12_Final.docx'
doc = Document(FILE)
changes = []

# ============================================
# 1. FIX DUPLICATE CAPTIONS
# ============================================
print("=== 1. Fix Duplicate Captions ===")
cap_fixes = {
    'Gambar 4.2 Use Case Diagram Sistem Mimotes AI Sistem Mimotes AI': 'Gambar 4.2 Use Case Diagram Sistem Mimotes AI',
    'Gambar 4.3 Activity Diagram Upload Dokumen Dokumen': 'Gambar 4.3 Activity Diagram Upload Dokumen',
    'Gambar 4.6 Arsitektur Sistem Mimotes AI Mimotes AI': 'Gambar 4.6 Arsitektur Sistem Mimotes AI',
}

for p in doc.paragraphs:
    style = p.style.name if p.style else ''
    if 'Caption' not in style:
        continue
    runs = p.runs
    if not runs:
        continue
    merged = ''.join(r.text for r in runs)
    for old, new in cap_fixes.items():
        if old in merged:
            runs[0].text = new
            for r in runs[1:]:
                r.text = ''
            changes.append(f'Caption fix: "{old[:40]}" -> "{new[:40]}"')
            print(f'  Fixed: {new}')

# Also fix via XML for captions without runs
body = doc.element.body
for p_elem in body.findall('.//' + qn('w:p')):
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        continue
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        continue
    style_val = pStyle.get(qn('w:val'))
    if style_val != 'Caption':
        continue
    t_elems = p_elem.findall('.//' + qn('w:t'))
    text = ''.join(t.text or '' for t in t_elems)
    for old, new in cap_fixes.items():
        if old in text:
            for t in t_elems:
                if t.text and old in t.text:
                    t.text = t.text.replace(old, new)
            changes.append(f'XML Caption fix: {new[:40]}')

# ============================================
# 2. FIX COMMA IN LEMBAR PERSETUJUAN
# ============================================
print("\n=== 2. Fix Lembar Persetujuan ===")
for i, p in enumerate(doc.paragraphs):
    if i < 36 or i > 58:
        continue
    text = p.text.strip()
    if text == ',':
        # Remove the comma paragraph
        runs = p.runs
        if runs:
            runs[0].text = ''
            for r in runs[1:]:
                r.text = ''
        changes.append(f'Removed standalone comma at para[{i}]')
        print(f'  Removed comma at para[{i}]')

# ============================================
# 3. FIX CAPSTONE PROJECT SENTENCE
# ============================================
print("\n=== 3. Fix Capstone Project Sentence ===")
for p in doc.paragraphs:
    runs = p.runs
    if not runs:
        continue
    merged = ''.join(r.text for r in runs)
    old = 'Pengembangan sistem chatbot AI (Mimotes) sebagai capstone project untuk optimalisasi layanan pelanggan.'
    new = 'Pengembangan prototipe sistem chatbot AI (Mimotes) sebagai capstone project untuk mendukung optimalisasi layanan pelanggan.'
    if old in merged:
        merged = merged.replace(old, new)
        runs[0].text = merged
        for r in runs[1:]:
            r.text = ''
        changes.append('Fixed capstone project sentence')
        print(f'  Fixed: "sistem" -> "prototipe sistem"')

# Also fix via XML
for p_elem in body.findall('.//' + qn('w:p')):
    t_elems = p_elem.findall('.//' + qn('w:t'))
    text = ''.join(t.text or '' for t in t_elems)
    if 'Pengembangan sistem chatbot AI (Mimotes) sebagai capstone project untuk optimalisasi' in text:
        for t in t_elems:
            if t.text and 'Pengembangan sistem chatbot AI' in t.text:
                t.text = t.text.replace(
                    'Pengembangan sistem chatbot AI (Mimotes) sebagai capstone project untuk optimalisasi layanan pelanggan.',
                    'Pengembangan prototipe sistem chatbot AI (Mimotes) sebagai capstone project untuk mendukung optimalisasi layanan pelanggan.'
                )

# ============================================
# 4. AUDIT DAFTAR ISI vs ACTUAL CONTENT
# ============================================
print("\n=== 4. Audit Daftar Isi ===")

# Collect actual headings
actual_headings = []
for i, p in enumerate(doc.paragraphs):
    style = p.style.style_id if p.style else ''
    text = p.text.strip()
    if style == 'Heading1' and text:
        actual_headings.append(('H1', text.replace('\n', ' '), i))
    elif style == 'Heading2' and text:
        actual_headings.append(('H2', text, i))
    elif style == 'Heading3' and text:
        actual_headings.append(('H3', text, i))

print("Actual headings in document:")
for level, text, idx in actual_headings:
    print(f"  {level} [{idx:3d}]: {text[:60]}")

# Check TOC entries
print("\nTOC entries:")
toc_entries = []
for i, p in enumerate(doc.paragraphs):
    style = p.style.style_id if p.style else ''
    if style.startswith('TOC'):
        text = p.text.strip()
        if text:
            toc_entries.append((style, text, i))
            print(f"  {style} [{i:3d}]: {text[:60]}")

# ============================================
# 5. AUDIT DAFTAR GAMBAR vs ACTUAL CAPTIONS
# ============================================
print("\n=== 5. Audit Daftar Gambar ===")
actual_captions = []
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ''
    text = p.text.strip()
    if 'Caption' in style and text.startswith('Gambar'):
        actual_captions.append((text, i))

print("Actual captions:")
for text, idx in actual_captions:
    print(f"  [{idx:3d}] {text[:70]}")

# Check if all captions are in Daftar Gambar
dg_entries = []
in_dg = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'DAFTAR GAMBAR' in text and p.style and p.style.style_id == 'Heading1':
        in_dg = True
        continue
    if in_dg:
        if p.style and p.style.style_id == 'Heading1':
            break
        style = p.style.style_id if p.style else ''
        if style.startswith('TOC'):
            dg_entries.append(text)

print("\nDaftar Gambar entries:")
for text in dg_entries:
    print(f"  {text[:70]}")

# Check missing
caption_names = set()
for text, _ in actual_captions:
    # Extract figure number
    parts = text.split('\t')
    name = parts[0].strip() if parts else text
    caption_names.add(name)

dg_names = set()
for text in dg_entries:
    parts = text.split('\t')
    name = parts[0].strip() if parts else text
    dg_names.add(name)

missing = caption_names - dg_names
extra = dg_names - caption_names
if missing:
    print(f"\n  MISSING from Daftar Gambar: {missing}")
if extra:
    print(f"\n  EXTRA in Daftar Gambar: {extra}")
if not missing and not extra:
    print("\n  All captions match Daftar Gambar ✅")

# ============================================
# 6. AUDIT DAFTAR TABEL vs ACTUAL TABLES
# ============================================
print("\n=== 6. Audit Daftar Tabel ===")
actual_table_captions = []
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ''
    text = p.text.strip()
    if 'Caption' in style and text.startswith('Tabel'):
        actual_table_captions.append((text, i))

print("Actual table captions:")
for text, idx in actual_table_captions:
    print(f"  [{idx:3d}] {text[:70]}")

# Count tables
print(f"\nTotal tables in doc: {len(doc.tables)}")

# ============================================
# SAVE
# ============================================
doc.save(FILE)
print(f"\n=== CHANGES APPLIED: {len(changes)} ===")
for c in changes:
    print(f"  - {c}")
print(f"\nFile: {FILE}")
print(f"Size: {os.path.getsize(FILE) / 1024:.0f} KB")
