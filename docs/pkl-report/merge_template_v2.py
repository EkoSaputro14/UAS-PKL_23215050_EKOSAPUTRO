#!/usr/bin/env python3
"""
Fix merge: handle heading style differences, fix title duplication,
properly replace BAB II-IV content.
"""
import copy
import os
from docx import Document
from docx.oxml.ns import qn

TEMPLATE = 'C:/Users/SMANSA/AppData/Local/hermes/cache/documents/doc_4ce6543a497c_LAPORAN PKL (1).docx'
SOURCE = 'LAPORAN_PKL_v14_UML.docx'
OUTPUT = 'LAPORAN_PKL_v15_Template.docx'

NEW_TITLE = "RANCANG BANGUN SISTEM CHATBOT AI BERBASIS PENGETAHUAN DENGAN RETRIEVAL-AUGMENTED GENERATION DAN PIPELINE CRM UNTUK OPTIMALISASI LAYANAN PELANGGAN"
NEW_STUDENT = "Eko Saputro"
NEW_NIM = "23215050"
NEW_SUPERVISOR = "Widianto Agung Nugroho"

OLD_STUDENT = "Moh. Arif Prasetyo"
OLD_NIM = "23215043"

print("📄 Loading...")
tpl_doc = Document(TEMPLATE)
src_doc = Document(SOURCE)

# ===== STEP 1: Fix all text replacements =====
print("\n📝 Step 1: Text replacements...")

# First, fix the title by clearing ALL runs and rewriting
for para in tpl_doc.paragraphs:
    full_text = para.text
    # Fix title duplication - replace entire paragraph text
    if "SISTEM PAKAR DETEKSI HAMA" in full_text:
        for run in para.runs:
            if "SISTEM PAKAR" in run.text or "RANCANG BANGUN" in run.text:
                run.text = ""
        # Set first meaningful run to new title
        for run in para.runs:
            if not run.text.strip():
                run.text = NEW_TITLE
                break
    
    # Replace student name
    for run in para.runs:
        if OLD_STUDENT in run.text:
            run.text = run.text.replace(OLD_STUDENT, NEW_STUDENT)
        if OLD_NIM in run.text:
            run.text = run.text.replace(OLD_NIM, NEW_NIM)
        if "Nama Dosen Pembimbing" in run.text:
            run.text = run.text.replace("Nama Dosen Pembimbing", NEW_SUPERVISOR)
        if "NIPY." in run.text:
            run.text = run.text.replace("NIPY.", "")

print("  ✅ Text replacements done")

# ===== STEP 2: Detect headings in template (with space) =====
print("\n📝 Step 2: Detecting headings...")

def get_style_name(para):
    """Get style name, handling both 'Heading 1' and 'Heading1' formats."""
    if para.style and para.style.name:
        return para.style.name
    # Check XML directly
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            return pStyle.get(qn('w:val'), '')
    return ''

def is_heading(para, level=None):
    style = get_style_name(para)
    if level is None:
        return style.startswith('Heading')
    return style == f'Heading {level}' or style == f'Heading{level}'

# Find all heading positions
tpl_headings = []
for i, para in enumerate(tpl_doc.paragraphs):
    if is_heading(para):
        text = para.text.strip()
        style = get_style_name(para)
        tpl_headings.append((i, text, style))

print(f"  Found {len(tpl_headings)} headings in template:")
for idx, text, style in tpl_headings:
    print(f"    [{idx:3d}] {style:15s} | {text[:60]}")

# ===== STEP 3: Find BAB positions in template =====
print("\n📝 Step 3: Finding BAB positions...")

tpl_babs = {}
for idx, text, style in tpl_headings:
    if 'BAB I' in text and 'PENDAHULUAN' not in text:
        # Check if next paragraph is PENDAHULUAN
        if idx + 1 < len(tpl_doc.paragraphs):
            next_text = tpl_doc.paragraphs[idx + 1].text.strip()
            if 'PENDAHULUAN' in next_text:
                tpl_babs['BAB_I'] = idx
    elif 'BAB II' in text:
        tpl_babs['BAB_II'] = idx
    elif 'BAB III' in text:
        tpl_babs['BAB_III'] = idx
    elif 'BAB IV' in text:
        tpl_babs['BAB_IV'] = idx
    elif 'BAB V' in text:
        tpl_babs['BAB_V'] = idx
    elif 'DAFTAR PUSTAKA' in text:
        tpl_babs['DAFTAR_PUSTAKA'] = idx
    elif 'LAMPIRAN' in text:
        tpl_babs['LAMPIRAN'] = idx

print(f"  BAB positions: {tpl_babs}")

# ===== STEP 4: Extract source content =====
print("\n📝 Step 4: Extracting source content...")

src_content = {}
current_bab = None
current_content = []

for para in src_doc.paragraphs:
    text = para.text.strip()
    style = get_style_name(para)
    
    # Detect BAB headings (source uses Heading1 without space)
    if style in ('Heading1', 'Heading 1'):
        if current_bab and current_content:
            src_content[current_bab] = current_content
        current_content = []
        
        if 'BAB I' in text and 'PENDAHULUAN' not in text:
            current_bab = 'BAB_I'
        elif 'BAB II' in text:
            current_bab = 'BAB_II'
        elif 'BAB III' in text:
            current_bab = 'BAB_III'
        elif 'BAB IV' in text:
            current_bab = 'BAB_IV'
        elif 'BAB V' in text:
            current_bab = 'BAB_V'
        elif 'DAFTAR PUSTAKA' in text:
            current_bab = 'DAFTAR_PUSTAKA'
        elif 'LAMPIRAN' in text:
            current_bab = 'LAMPIRAN'
        else:
            current_bab = None
        continue
    
    if current_bab:
        current_content.append({
            'text': text,
            'style': style,
        })

if current_bab and current_content:
    src_content[current_bab] = current_content

for bab, content in src_content.items():
    print(f"  {bab}: {len(content)} paragraphs")

# ===== STEP 5: Replace content section by section =====
print("\n📝 Step 5: Replacing BAB content...")

# Sort bab positions
sorted_babs = sorted(tpl_babs.items(), key=lambda x: x[1])

for i, (tpl_key, start_idx) in enumerate(sorted_babs):
    # Find end of section
    end_idx = len(tpl_doc.paragraphs)
    for j in range(i + 1, len(sorted_babs)):
        next_key, next_idx = sorted_babs[j]
        end_idx = next_idx
        break
    
    src_key = tpl_key  # Same key mapping
    if src_key not in src_content:
        print(f"  ⚠️ No source content for {src_key}")
        continue
    
    section_paras = list(tpl_doc.paragraphs[start_idx:end_idx])
    src_paras = src_content[src_key]
    
    # Find where content starts (after heading + subheading like PENDAHULUAN)
    content_start = 1
    if tpl_key == 'BAB_I' and len(section_paras) > 1:
        if section_paras[1].text.strip() == 'PENDAHULUAN':
            content_start = 2
    
    available = len(section_paras) - content_start
    
    # Replace content
    for j, src_p in enumerate(src_paras):
        if j < available:
            target = section_paras[content_start + j]
            # Clear all runs
            for run in target.runs:
                run.text = ''
            # Set text in first run
            if target.runs:
                target.runs[0].text = src_p['text']
            elif src_p['text']:
                target.add_run(src_p['text'])
    
    replaced = min(len(src_paras), available)
    print(f"  ✅ {tpl_key}: {replaced}/{len(src_paras)} paragraphs replaced")

# ===== SAVE =====
print(f"\n💾 Saving...")
tpl_doc.save(OUTPUT)
print(f"✅ Saved: {OUTPUT}")

# ===== VERIFY =====
doc2 = Document(OUTPUT)
print(f"\n📊 Verification:")

# Check title
title_ok = False
for para in doc2.paragraphs[:30]:
    if "SISTEM CHATBOT AI" in para.text:
        title_ok = True
        # Check no duplication
        if "RANCANG BANGUN RANCANG BANGUN" in para.text:
            print(f"  ❌ Title still duplicated!")
        else:
            print(f"  ✅ Title: {para.text[:60]}")
        break

# Check student
for para in doc2.paragraphs[:30]:
    if NEW_STUDENT in para.text:
        print(f"  ✅ Student: {para.text.strip()}")
        break

# Check BAB content
for para in doc2.paragraphs:
    if 'Bank Mandiri KCP Tegal Sudirman' in para.text:
        print(f"  ✅ BAB II content present")
        break

for para in doc2.paragraphs:
    if 'Retrieval-Augmented Generation' in para.text and len(para.text) > 100:
        print(f"  ✅ BAB III/IV content present")
        break

print(f"  File size: {os.path.getsize(OUTPUT)/1024:.0f} KB")
