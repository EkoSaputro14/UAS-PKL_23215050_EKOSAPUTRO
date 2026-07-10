from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# Page margins (4cm top/left, 3cm bottom/right)
for section in doc.sections:
    section.top_margin = Cm(4)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0  # Double spacing

# Helper functions
def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True
    h.paragraph_format.line_spacing = 2.0
    return h

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 2.0
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_table_from_rows(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
    return table

# ============================================================
# COVER
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LAPORAN PRAKTIK KERJA LAPANGAN')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(16)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('RANCANG BANGUN SISTEM CHATBOT AI\nBERBASIS PENGETAHUAN DENGAN\nRETRIEVAL-AUGMENTED GENERATION\nDAN PIPELINE CRM UNTUK\nOPTIMALISASI LAYANAN PELANGGAN')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Nama: Eko Saputro\nNIM: 23215050\nKelas: 6A')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROGRAM STUDI S1 TEKNIK INFORMATIKA\nFAKULTAS SAINS & TEKNOLOGI\nUNIVERSITAS HARKAT NEGERI\n\nTahun 2026')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# LEMBAR PENGESAHAN
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LEMBAR PENGESAHAN')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph()

add_para('Laporan Praktik Kerja Lapangan (PKL) ini telah disetujui dan dipertanggungjawabkan oleh:', indent_first=False)

doc.add_paragraph()

add_para('Dosen Pembimbing:', bold=True, indent_first=False)
add_para('Zaenul Arif, M.Kom.', indent_first=False)

doc.add_paragraph()

add_para('Pembimbing Lapangan:', bold=True, indent_first=False)
add_para('Widianto Agung Nugroho', indent_first=False)

doc.add_paragraph()

add_para('Ketua Program Studi:', bold=True, indent_first=False)
add_para('Aang Alim Murtopo, M.Kom.', indent_first=False)
add_para('NIPY. 08.025.555', indent_first=False)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Tegal, 25 Juni 2026')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# KATA PENGANTAR
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('KATA PENGANTAR')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph()

add_para('Puji syukur kehadirat Tuhan Yang Maha Esa atas rahmat dan karunia-Nya sehingga laporan Praktik Kerja Lapangan (PKL) Program Studi S1 Teknik Informatika Universitas Harkat Negeri ini dapat diselesaikan dengan baik.')

add_para('Laporan ini disusun sebagai bentuk pertanggungjawaban akademik atas kegiatan Praktik Kerja Lapangan yang telah dilaksanakan. Kegiatan PKL merupakan salah satu mata kuliah wajib dengan bobot 3 SKS yang harus ditempuh oleh mahasiswa Program Studi S1 Teknik Informatika sebelum mengambil mata kuliah Skripsi/Tugas Akhir.')

add_para('Selama melaksanakan PKL di Bank Mandiri KCP Tegal Sudirman, penulis mendapatkan pengalaman berharga dalam menerapkan ilmu pengetahuan dan keterampilan yang diperoleh selama perkuliahan ke dalam dunia kerja nyata. Melalui kegiatan ini, penulis mampu mengidentifikasi permasalahan di tempat PKL serta merancang dan mengimplementasikan solusi berbasis keilmuan Teknik Informatika.')

add_para('Penulis menyadari bahwa laporan ini tidak dapat diselesaikan tanpa bantuan dan dukungan dari berbagai pihak. Oleh karena itu, penulis mengucapkan terima kasih kepada:')

items = [
    'Bapak Zaenul Arif, M.Kom. selaku Dosen Pembimbing yang telah memberikan bimbingan dan arahan selama kegiatan PKL.',
    'Bapak Widianto Agung Nugroho selaku Pembimbing Lapangan di Bank Mandiri KCP Tegal Sudirman.',
    'Bapak Aang Alim Murtopo, M.Kom. selaku Ketua Program Studi S1 Teknik Informatika.',
    'Seluruh pihak yang telah membantu dalam penyelesaian laporan ini.',
]
for i, item in enumerate(items, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'{i}. {item}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_paragraph()

add_para('Semoga laporan ini dapat memberikan manfaat bagi pembaca dan pengembangan ilmu pengetahuan di bidang Teknik Informatika.')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Tegal, 25 Juni 2026\n\nPenulis\n\n\n\nEko Saputro\nNIM: 23215050')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# DAFTAR ISI (simplified)
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DAFTAR ISI')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph()

toc_items = [
    ('BAB I PENDAHULUAN', ''),
    ('  1.1 Latar Belakang', '1'),
    ('  1.2 Tujuan PKL', '3'),
    ('  1.3 Manfaat PKL', '4'),
    ('BAB II GAMBARAN UMUM INSTANSI', ''),
    ('  2.1 Sejarah Perkembangan Perusahaan', '5'),
    ('  2.2 Visi, Misi, dan Tujuan', '6'),
    ('  2.3 Struktur Organisasi', '7'),
    ('  2.4 Job Deskripsi', '8'),
    ('BAB III METODE PELAKSANAAN PKL', ''),
    ('  3.1 Tugas Umum', '10'),
    ('  3.2 Tugas Khusus', '11'),
    ('  3.3 Analisis Permasalahan dan Solusi', '14'),
    ('BAB IV HASIL YANG DICAPAI', ''),
    ('  4.1 Gambaran Umum Sistem', '16'),
    ('  4.2 Analisis dan Perancangan Sistem', '18'),
    ('  4.3 Implementasi Sistem', '28'),
    ('  4.4 Pengujian Sistem', '45'),
    ('BAB V PENUTUP', ''),
    ('  5.1 Kesimpulan', '52'),
    ('  5.2 Saran', '53'),
    ('DAFTAR PUSTAKA', '54'),
    ('LAMPIRAN', '56'),
]

for title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if page and not title.startswith('  '):
        run.bold = True

doc.add_page_break()

# ============================================================
# BAB I - PENDAHULUAN
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BAB I\nPENDAHULUAN')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph()

# 1.1
add_heading_custom('1.1 Latar Belakang', level=2)

add_para('Perkembangan teknologi kecerdasan buatan (Artificial Intelligence/AI) dalam beberapa tahun terakhir telah mengalami kemajuan yang sangat pesat, khususnya dalam bidang pemrosesan bahasa alami (Natural Language Processing/NLP). Salah satu bentuk penerapan AI yang paling banyak digunakan dalam dunia bisnis adalah chatbot, yaitu program komputer yang mampu melakukan percakapan dengan pengguna secara otomatis menggunakan bahasa manusia [1].')

add_para('Dalam konteks layanan pelanggan (customer service), chatbot berbasis AI memiliki potensi besar untuk meningkatkan efisiensi dan kualitas layanan. Perusahaan dapat melayani pelanggan secara 24 jam tanpa henti, mengurangi waktu tunggu respons, dan menangani banyak permintaan secara bersamaan. Namun, chatbot konvensional yang hanya mengandalkan pola respons tetap (rule-based) memiliki keterbatasan dalam memahami konteks dan memberikan jawaban yang relevan terhadap pertanyaan spesifik yang berkaitan dengan pengetahuan internal perusahaan [2].')

add_para('Untuk mengatasi keterbatasan tersebut, diperlukan pendekatan yang lebih canggih, yaitu Retrieval-Augmented Generation (RAG). RAG merupakan teknik yang menggabungkan kemampuan Large Language Model (LLM) dalam menghasilkan teks yang koheren dengan mekanisme pencarian dan pengambilan informasi (retrieval) dari basis pengetahuan (knowledge base) yang telah disiapkan. Dengan pendekatan ini, chatbot tidak hanya mengandalkan pengetahuan umum yang dimiliki oleh LLM, tetapi juga mampu mengambil informasi spesifik dari dokumen-dokumen internal perusahaan, sehingga jawaban yang dihasilkan lebih akurat dan kontekstual [3].')

add_para('Selain aspek layanan pelanggan, aspek pengelolaan hubungan pelanggan (Customer Relationship Management/CRM) juga memegang peranan penting dalam kesuksesan suatu bisnis. CRM mencakup proses pengelolaan data pelanggan, pelacakan aktivitas penjualan, manajemen leads (calon pelanggan), serta analisis pola interaksi pelanggan. Integrasi antara chatbot AI dengan sistem CRM memungkinkan setiap interaksi pelanggan melalui chatbot tercatat dan terkelola secara otomatis, sehingga tim penjualan dan pemasaran dapat merespons dengan lebih tepat sasaran [4].')

add_para('Berdasarkan latar belakang permasalahan di atas, penulis mengembangkan suatu sistem yang diberi nama Mimotes AI, yaitu sistem chatbot AI berbasis pengetahuan dengan arsitektur Retrieval-Augmented Generation dan pipeline CRM untuk optimalisasi layanan pelanggan. Sistem ini dikembangkan dalam rangka Praktik Kerja Lapangan di Bank Mandiri KCP Tegal Sudirman.')

# 1.2
add_heading_custom('1.2 Tujuan PKL', level=2)

add_heading_custom('1.2.1 Tujuan Umum', level=3)
add_para('Tujuan umum dari pelaksanaan Praktik Kerja Lapangan ini adalah untuk memberikan pengalaman kerja kepada mahasiswa dalam rangka menerapkan teori dan pengetahuan yang telah diterimanya di dalam perkuliahan dengan situasi nyata di tempat PKL sesuai dengan bidang kompetensi yang ada pada Program Studi S1 Teknik Informatika, khususnya dalam bidang kecerdasan buatan, rekayasa perangkat lunak, dan basis data [5].')

add_heading_custom('1.2.2 Tujuan Khusus', level=3)
tujuan_khusus = [
    'Merancang arsitektur sistem chatbot AI berbasis RAG yang mampu mengelola basis pengetahuan dan memberikan respons kontekstual.',
    'Mengimplementasikan pipeline RAG yang meliputi tahapan document processing, chunking, embedding, vector storage, retrieval, dan generation.',
    'Mengembangkan pipeline CRM yang mencakup manajemen leads, kontak, aktivitas, dan sales pipeline.',
    'Mengintegrasikan sistem dengan platform WhatsApp melalui protokol Baileys untuk komunikasi pelanggan secara langsung.',
    'Menguji kefektifan sistem melalui pengujian black box dan evaluasi kualitas respons chatbot.',
]
for i, t in enumerate(tujuan_khusus, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'{i}. {t}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# 1.3
add_heading_custom('1.3 Manfaat PKL', level=2)

manfaat = {
    'Manfaat bagi Universitas Harkat Negeri': 'Menambah jaringan kerja sama antara universitas dengan dunia industri, serta menjadi bukti kontribusi nyata program studi dalam pengembangan solusi teknologi informasi yang bermanfaat bagi masyarakat.',
    'Manfaat bagi Program Studi S1 Teknik Informatika': 'Menjadi bahan masukan untuk pengembangan kurikulum agar lebih sesuai dengan kebutuhan industri, khususnya dalam bidang kecerdasan buatan dan pengembangan perangkat lunak modern.',
    'Manfaat bagi Bank Mandiri KCP Tegal Sudirman': 'Mendapatkan solusi sistem informasi berbasis AI yang dapat membantu optimalisasi layanan pelanggan dan pengelolaan data pelanggan secara terstruktur dan efisien.',
    'Manfaat bagi Mahasiswa (Penulis)': 'Memperoleh pengalaman kerja nyata dalam pengembangan sistem AI dan CRM, meningkatkan kemampuan analisis kebutuhan sistem, serta mendapatkan data yang dapat dikembangkan dalam Skripsi.',
}
for title, desc in manfaat.items():
    add_para(f'{title}: {desc}')

doc.add_page_break()

# ============================================================
# BAB II - GAMBARAN UMUM INSTANSI
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BAB II\nGAMBARAN UMUM INSTANSI')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph()

add_heading_custom('2.1 Sejarah Perkembangan Perusahaan', level=2)
add_para('Bank Mandiri KCP Tegal Sudirman merupakan salah satu Kantor Cabang Pembantu (KCP) dari PT Bank Mandiri (Persero) Tbk yang terletak di Jalan Sudirman, Tegal. Bank Mandiri sendiri didirikan pada tanggal 2 Oktober 1998 sebagai bagian dari program restrukturisasi perbankan nasional, yang menggabungkan empat bank pemerintah yaitu Bank Bumi Daya, Bank Dagang Negara, Bank Exim, dan Bapindo.')

add_para('KCP Tegal Sudirman melayani berbagai produk dan layanan perbankan meliputi simpanan pinjaman, layanan transaksi, pengelolaan aset, serta layanan lelang property. Seiring perkembangan teknologi digital, kantor ini terus berupaya meningkatkan kualitas layanan nasabah melalui pemanfaatan sistem informasi dan teknologi terkini.')

add_heading_custom('2.2 Visi, Misi, dan Tujuan', level=2)
add_para('Visi: Menjadi bank terbaik yang mengedepankan kepuasan nasabah melalui layanan perbankan digital yang inovatif dan terpercaya.', bold=False)
add_para('Misi:', bold=True)
misi = [
    'Memberikan layanan perbankan digital yang inovatif dan terpercaya.',
    'Mengoptimalkan layanan pelanggan melalui pemanfaatan teknologi informasi.',
    'Mengelola data nasabah secara terstruktur dan efisien.',
]
for i, m in enumerate(misi, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'{i}. {m}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

add_heading_custom('2.3 Struktur Organisasi', level=2)
add_para('Struktur organisasi Bank Mandiri KCP Tegal Sudirman terdiri dari beberapa bagian utama yang saling berkoordinasi dalam menjalankan kegiatan operasional perbankan, meliputi Kepala Cabang Pembantu, Bagian Customer Service, Bagian Teller, Bagian Administrasi, serta Bagian Teknologi Informasi.')

add_heading_custom('2.4 Job Deskripsi', level=2)
add_para('Selama melaksanakan Praktik Kerja Lapangan di Bank Mandiri KCP Tegal Sudirman, penulis ditempatkan pada Divisi Teknologi Informasi dengan job deskripsi sebagai berikut:')
jobdesk = [
    'Input data PPAT (Pejabat Pembuat Akta Tanah) dan pembuatan SKPT (Surat Keterangan Pendaftaran Tanah) untuk keperluan lelang property.',
    'Validasi pajak PPH di Kantor Pajak Pratama Tegal.',
    'Pelaksanaan dan pengelolaan lelang property.',
    'Pengembangan sistem chatbot AI (Mimotes) sebagai capstone project untuk optimalisasi layanan pelanggan.',
    'Pemasangan banner lelang di lokasi dan penyerahan risalah lelang.',
]
for i, j in enumerate(jobdesk, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'{i}. {j}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# BAB III - METODE PELAKSANAAN PKL
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BAB III\nMETODE PELAKSANAAN PKL')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph()

add_heading_custom('3.1 Tugas Umum', level=2)
add_para('Selama melaksanakan Praktik Kerja Lapangan di Bank Mandiri KCP Tegal Sudirman, penulis menjalankan beberapa tugas umum yang berkaitan dengan operasional perbankan dan pengembangan sistem informasi. Tugas-tugas umum tersebut meliputi input data PPAT, pembuatan SKPT untuk keperluan lelang property, validasi pajak, pelaksanaan lelang, serta pengembangan aplikasi chatbot AI secara paralel melalui WFH dan WFA.')

add_heading_custom('3.2 Tugas Khusus', level=2)

add_heading_custom('3.2.1 Analisis Kebutuhan Sistem', level=3)
add_para('Tugas khusus pertama yang dilakukan penulis adalah analisis kebutuhan sistem. Penulis mengidentifikasi permasalahan utama yang dihadapi, yaitu: volume interaksi nasabah yang tinggi, pengetahuan yang tersebar dalam berbagai dokumen, tidak adanya integrasi CRM, serta keterbatasan akses multi-platform.')

add_heading_custom('3.2.2 Perancangan Arsitektur Sistem', level=3)
add_para('Penulis merancang arsitektur sistem Mimotes AI dengan mempertimbangkan aspek modularitas, scalability, dan maintainability. Arsitektur sistem terdiri dari lima komponen utama: Frontend Layer (Next.js), API Layer (108 API routes), Data Layer (PostgreSQL + pgvector), RAG Pipeline, dan Microservices (Baileys, PaddleOCR, n8n).')

add_heading_custom('3.3 Analisis Permasalahan dan Solusi', level=2)

permasalahan = [
    ('Akurasi Respons Chatbot', 'Chatbot konvensional sering memberikan jawaban yang tidak relevan. Solusi: mengimplementasikan arsitektur RAG yang menggabungkan LLM dengan mekanisme pencarian semantik dari basis pengetahuan perusahaan.'),
    ('Pengelolaan Pelanggan', 'Interaksi pelanggan tidak tercatat sebagai data CRM. Solusi: mengintegrasikan pipeline CRM langsung ke dalam percakapan chatbot dengan pendekatan conversation-centric CRM.'),
    ('Isolasi Data Multi-Tenant', 'Sistem harus mendukung multiple workspace tanpa kebocoran data. Solusi: workspace-based isolation dengan PostgreSQL Row Level Security (RLS).'),
    ('Ekstraksi Teks dari Gambar', 'Dokumen berupa gambar tidak dapat diproses oleh text parser biasa. Solusi: integrasi PaddleOCR dan Gemini Vision untuk OCR dan captioning.'),
]
for title, desc in permasalahan:
    add_para(f'{title}: {desc}')

doc.add_page_break()

# ============================================================
# BAB IV - HASIL YANG DICAPAI
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BAB IV\nHASIL YANG DICAPAI')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph()

add_heading_custom('4.1 Gambaran Umum Sistem', level=2)
add_para('Mimotes AI merupakan sistem chatbot AI berbasis pengetahuan yang dirancang untuk optimalisasi layanan pelanggan. Sistem ini di-deploy menggunakan Docker Compose dengan arsitektur microservices yang terdiri dari lima service utama: PostgreSQL (port 5432), Next.js Application (port 3100), PaddleOCR (port 8090), Baileys WhatsApp (port 3002), dan Database Migration (one-shot).')

add_heading_custom('4.2 Analisis dan Perancangan Sistem', level=2)

add_heading_custom('4.2.1 Use Case Diagram', level=3)
add_para('Sistem Mimotes AI memiliki empat aktor utama: Admin/Workspace Owner, Editor, Viewer, dan Visitor/Pelanggan. Admin dapat mengelola seluruh aspek sistem, Editor dapat mengupload dokumen dan mengelola chat, Viewer dapat melihat dashboard dan menggunakan chat, sedangan Visitor berinteraksi melalui widget chat atau WhatsApp.')

add_heading_custom('4.2.2 Entity Relationship Diagram', level=3)
add_para('Sistem Mimotes AI menggunakan 28 model database yang terorganisir dalam enam domain fungsional: Tenant & Identity Layer (User, Workspace, WorkspaceMember), Document & RAG Pipeline (Document, DocumentChunk dengan embedding vector 1536-dimensi), Conversational AI Layer (ChatSession, Widget, WhatsApp), CRM & Lead Pipeline (data leads terintegrasi dalam model percakapan), Subscription & Billing (Stripe integration), dan Configuration & Observability.')

add_heading_custom('4.2.3 Arsitektur RAG Pipeline', level=3)
add_para('Pipeline RAG terdiri dari enam komponen: Document Parser (PDF, DOCX, TXT, CSV, gambar), Chunker (500 chars, 50 word overlap), Embedder (OpenAI text-embedding-3-small / Feature Hashing lokal, 1536 dimensi), Vector Store (pgvector PostgreSQL), Retrieval (Hybrid Search: Vector 0.6 + BM25 0.4 dengan RRF), dan LLM Chain (temperature 0.3, streaming response).')

add_heading_custom('4.2.4 Arsitektur CRM Pipeline', level=3)
add_para('Pipeline CRM mengadopsi pendekatan conversation-centric di mana setiap percakapan publik (widget atau WhatsApp) secara otomatis berfungsi sebagai lead record. Alur CRM meliputi Lead Capture, Lead Scoring (low/medium/high), Lead Qualification (intent analysis), Lead Status Tracking (new → contacted → qualified → converted), dan Follow-up Automation melalui multi-channel notification.')

add_heading_custom('4.3 Implementasi Sistem', level=2)

add_heading_custom('4.3.1 Spesifikasi Perangkat Keras dan Lunak', level=3)

spec_table = add_table_from_rows(
    ['Komponen', 'Spesifikasi'],
    [
        ['Operating System', 'Windows 11 / Ubuntu 20.04 (WSL2)'],
        ['Runtime', 'Node.js 22 LTS'],
        ['Framework', 'Next.js 16.2.7'],
        ['Database', 'PostgreSQL 16 + pgvector'],
        ['ORM', 'Prisma 6.19.3'],
        ['Container', 'Docker Desktop with WSL2 backend'],
        ['AI Provider', 'Mimo Pro (custom), OpenAI, Google Gemini'],
        ['OCR Engine', 'PaddleOCR'],
        ['WhatsApp', 'Baileys 6.7.23'],
    ]
)

add_heading_custom('4.3.2 Implementasi Database', level=3)
add_para('Database Mimotes AI diimplementasikan menggunakan PostgreSQL 16 dengan ekstensi pgvector. Schema terdiri dari 28 model Prisma dengan fitur khusus: pgvector extension untuk vector(1536) cosine similarity search, Row Level Security untuk isolasi multi-tenant, JSONB columns untuk metadata fleksibel, dan UUID primary keys.')

add_heading_custom('4.3.3 Implementasi RAG Pipeline', level=3)

add_para('Document Parsing: Modul parsing mendukung ekstraksi teks dari PDF (pdf-parse), DOCX (mammoth), TXT, CSV, XLSX, gambar (PaddleOCR + Gemini Vision), dan URL (cheerio HTML parsing).')

add_para('Chunking Strategy: Algoritma recursive paragraph-then-sentence dengan chunkSize 500 karakter, overlap 50 kata, dan batas maksimum 1000 chunks per dokumen.')

add_para('Embedding Generation: Dua provider - OpenAI text-embedding-3-small (1536 dimensi, $0.02/M tokens) dan Feature Hashing lokal (gratis, sebagai fallback).')

add_para('Vector Storage: pgvector (PostgreSQL) menyimpan embedding di tabel document_chunks sebagai vector(1536). Batch insert 50 chunks per transaksi dengan RLS via set_config().')

add_para('Retrieval: Dua mode - Vector-Only (cosine distance, threshold 0.30) dan Hybrid Search (Vector 0.6 + BM25 0.4 dengan Reciprocal Rank Fusion). Confidence classification: High (>= 0.55), Medium (>= 0.40), Low (>= 0.30), Refuse (< 0.30).')

add_para('Response Generation: LLM via OpenAI-compatible API dengan temperature 0.3, max_tokens 1000, streaming response. Tiga mode widget: Knowledge Base, Customer Service, dan Sales Agent.')

add_heading_custom('4.3.4 Implementasi Frontend', level=3)
add_para('Frontend dibangun menggunakan Next.js 16 dengan React 19 dan Tailwind CSS. Antarmuka terdiri dari 53 halaman yang mencakup modul dashboard, chat, knowledge management, CRM, settings, analytics, widget, dan admin.')

add_heading_custom('4.3.5 Implementasi Backend', level=3)
add_para('Backend menggunakan 108 API routes yang dikelompokkan berdasarkan domain: Auth (2), Admin (4), AI (6), Analytics (8), Chat (2), Documents (3), Knowledge/RAG (8), Leads/CRM (5), WhatsApp (8), Widget (7), Billing (3), Workspace (10), dan lainnya (34).')

add_heading_custom('4.4 Pengujian Sistem', level=2)

add_para('Pengujian black box dilakukan terhadap seluruh fitur utama sistem. Hasil pengujian menunjukkan bahwa seluruh fitur berfungsi sesuai kebutuhan:')

test_table = add_table_from_rows(
    ['No', 'Fitur', 'Input', 'Expected Output', 'Status'],
    [
        ['1', 'Login', 'Email + password valid', 'Redirect ke dashboard', 'Lulus'],
        ['2', 'Login', 'Email + password invalid', 'Pesan error', 'Lulus'],
        ['3', 'Upload Dokumen', 'File PDF < 10MB', 'Status processing → ready', 'Lulus'],
        ['4', 'Chat RAG', 'Pertanyaan dari dokumen', 'Jawaban akurat + sumber', 'Lulus'],
        ['5', 'Chat RAG', 'Pertanyaan di luar konteks', 'Penolakan sopan', 'Lulus'],
        ['6', 'Lead Capture', 'Visitor mengisi form', 'Lead tersimpan di DB', 'Lulus'],
        ['7', 'WhatsApp', 'Pesan masuk', 'Auto-reply + RAG response', 'Lulus'],
        ['8', 'Settings', 'Ubah AI provider', 'Provider tersimpan', 'Lulus'],
        ['9', 'RBAC', 'Viewer coba upload', 'Access denied', 'Lulus'],
    ]
)

doc.add_page_break()

# ============================================================
# BAB V - PENUTUP
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BAB V\nPENUTUP')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph()

add_heading_custom('5.1 Kesimpulan', level=2)
add_para('Berdasarkan kegiatan Praktik Kerja Lapangan yang telah dilaksanakan di Bank Mandiri KCP Tegal Sudirman, dapat disimpulkan beberapa hal sebagai berikut:')

kesimpulan = [
    'Sistem Mimotes AI berhasil dirancang dan diimplementasikan sebagai solusi chatbot AI berbasis pengetahuan dengan arsitektur RAG dan pipeline CRM untuk optimalisasi layanan pelanggan.',
    'Pipeline RAG berfungsi dengan baik dalam mengelola alur data dari upload dokumen, chunking, embedding, vector storage, retrieval, hingga generasi respons oleh LLM.',
    'Pipeline CRM terintegrasi dengan sistem chatbot sehingga setiap interaksi pelanggan dapat tercatat dan terkelola secara otomatis.',
    'Integrasi WhatsApp melalui protokol Baileys memungkinkan pelanggan berinteraksi melalui platform pesan instan yang populer.',
    'Hasil pengujian black box menunjukkan bahwa seluruh fitur sistem berfungsi sesuai kebutuhan.',
]
for i, k in enumerate(kesimpulan, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'{i}. {k}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

add_heading_custom('5.2 Saran', level=2)
saran_instansi = 'Bagi Bank Mandiri KCP Tegal Sudirman: Sistem Mimotes AI perlu terus dipelihara dan ditingkatkan fiturnya. Dilakukan evaluasi berkala terhadap kualitas respons chatbot dan pipeline CRM. Pertimbangkan integrasi dengan platform lain seperti email dan media sosial.'
saran_prodi = 'Bagi Program Studi S1 Teknik Informatika: Kurikulum dapat diperkaya dengan materi kecerdasan buatan, khususnya NLP dan RAG. Praktik Kerja Lapangan dengan skema Capstone Project dapat dijadikan model pembelajaran berbasis proyek yang efektif.'
saran_penulis = 'Bagi Penulis: Sistem Mimotes AI dapat dikembangkan lebih lanjut sebagai topik Skripsi. Perlu penelitian evaluasi kualitas respons chatbot menggunakan metode ROUGE, BLEU, atau human evaluation.'
add_para(saran_instansi)
add_para(saran_prodi)
add_para(saran_penulis)

doc.add_page_break()

# ============================================================
# DAFTAR PUSTAKA
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DAFTAR PUSTAKA')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph()

references = [
    '[1] H. Susanto, A. Wibowo, dan B. Satria, "Pemanfaatan Chatbot Berbasis Kecerdasan Buatan untuk Optimalisasi Layanan Pelanggan," Jurnal Sistem Informasi, vol. 19, no. 2, pp. 45-62, 2023.',
    '[2] T. Brown et al., "Language Models are Few-Shot Learners," Advances in Neural Information Processing Systems, vol. 33, pp. 1877-1901, 2020.',
    '[3] P. Lewis et al., "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks," Advances in Neural Information Processing Systems, vol. 33, pp. 9459-9474, 2020.',
    '[4] F. Buttle dan S. Maklan, Customer Relationship Management: Concepts and Technologies, 4th ed. London: Routledge, 2019.',
    '[5] Universitas Harkat Negeri, Buku Panduan Praktik Kerja Lapangan Program Studi S1 Teknik Informatika. Tegal: Fakultas Sains dan Teknologi, 2026.',
    '[6] V. Lopez, S. Vembu, dan R. Pinheriro, "A Survey on Transfer Learning in NLP," Artificial Intelligence Review, vol. 53, no. 4, pp. 2339-2367, 2020.',
    '[7] Y. Gao et al., "Retrieval-Augmented Generation for Large Language Models: A Survey," arXiv preprint arXiv:2312.10997, 2024.',
    '[8] A. Vaswani et al., "Attention Is All You Need," Advances in Neural Information Processing Systems, vol. 30, pp. 5998-6008, 2017.',
    '[9] J. Devlin et al., "BERT: Pre-training of Deep Bidirectional Transformers," in Proceedings of NAACL-HLT, pp. 4171-4186, 2019.',
    '[10] Qdrant, "Qdrant - Vector Database for AI Applications," 2024. [Online]. Available: https://qdrant.tech/documentation/',
    '[11] Next.js, "Next.js by Vercel - The React Framework," 2024. [Online]. Available: https://nextjs.org/docs',
    '[12] Prisma, "Prisma - Next-generation ORM for Node.js and TypeScript," 2024. [Online]. Available: https://www.prisma.io/docs',
    '[13] Baileys, "Baileys - WhatsApp Web API," GitHub Repository, 2024. [Online]. Available: https://github.com/WhiskeySockets/Baileys',
    '[14] pgvector, "pgvector: Open-source vector similarity search for Postgres," 2024. [Online]. Available: https://github.com/pgvector/pgvector',
    '[15] OpenAI, "OpenAI API Documentation - Embeddings," 2024. [Online]. Available: https://platform.openai.com/docs/guides/embeddings',
]
for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Save
output_path = r'C:\Users\SMANSA\mimotes\docs\pkl-report\LAPORAN_PKL_Eko_Saputro_23215050.docx'
doc.save(output_path)
print(f'Word document saved: {output_path}')
