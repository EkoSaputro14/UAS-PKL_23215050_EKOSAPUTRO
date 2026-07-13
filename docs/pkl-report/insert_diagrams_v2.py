#!/usr/bin/env python3
"""
Insert diagrams using python-docx's add_picture method.
Simpler approach: find caption → insert image before it.
"""
import os
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

INPUT = 'LAPORAN_PKL_v15_Template.docx'
OUTPUT = 'LAPORAN_PKL_v15_Template.docx'
DIAGRAMS_DIR = 'diagrams'

# Mapping: search keyword in caption → (new caption text, image file)
CAPTION_MAP = {
    "Halaman Deteksi": ("Gambar 4.2 Use Case Diagram Sistem Mimotes AI", "usecase.png"),
    "Halaman Panel Admin": ("Gambar 4.3 Activity Diagram Upload Dokumen", "activity-upload.png"),
    "Implementasi database": ("Gambar 4.4 Activity Diagram Proses Chat RAG", "activity-chat.png"),
    "Input Keyakinan": ("Gambar 4.5 Sequence Diagram Chat RAG", "sequence-chat-rag.png"),
    "Hasil Diag": ("Gambar 4.6 ERD Domain Identity & Workspace", "erd-a-identity.png"),
}

# Additional images to insert (with new captions, matched to content positions)
EXTRA_DIAGRAMS = [
    ("4.7 ERD Domain RAG & Knowledge Base", "erd-b-rag.png"),
    ("4.8 ERD Domain Chat & CRM", "erd-c-chat-crm.png"),
    ("4.9 ERD Domain Billing & Configuration", "erd-d-billing.png"),
    ("4.10 ERD Ringkasan — Seluruh Relasi", "erd-summary.png"),
    ("4.11 Arsitektur Sistem Mimotes AI", "architecture.png"),
    ("4.12 Arsitektur RAG Pipeline", "rag-pipeline.png"),
    ("4.13 Arsitektur CRM Pipeline", "crm-pipeline.png"),
]

print("📄 Loading...")
doc = Document(INPUT)

# ===== STEP 1: Find and replace existing captions in BAB IV =====
print("\n📝 Step 1: Replacing existing captions...")

# Find BAB IV range
bab4_start = None
bab5_start = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    if style == 'Heading 1' and 'BAB IV' in text:
        bab4_start = i
    elif style == 'Heading 1' and 'BAB V' in text:
        bab5_start = i

print(f"  BAB IV: {bab4_start}-{bab5_start}")

# Find captions and replace
replaced_count = 0
for i in range(bab4_start or 0, bab5_start or len(doc.paragraphs)):
    para = doc.paragraphs[i]
    text = para.text.strip()
    
    for keyword, (new_caption, img_file) in CAPTION_MAP.items():
        if keyword in text:
            # Replace caption
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = new_caption
            
            # Insert image before caption
            img_path = os.path.join(DIAGRAMS_DIR, img_file)
            if os.path.exists(img_path):
                try:
                    new_para = doc.add_paragraph()
                    new_para.alignment = 1  # CENTER
                    new_para.add_run().add_picture(img_path, width=Inches(5.5))
                    
                    # Move before caption
                    parent = para._element.getparent()
                    pos = list(parent).index(para._element)
                    new_elem = new_para._element
                    parent.remove(new_elem)
                    parent.insert(pos, new_elem)
                    
                    replaced_count += 1
                    print(f"  ✅ {keyword} → {new_caption[:40]} + {img_file}")
                except Exception as e:
                    print(f"  ❌ {keyword}: {e}")
            break

print(f"\n  Replaced {replaced_count} captions with images")

# ===== STEP 2: Insert extra diagrams after the last replaced caption =====
print("\n📝 Step 2: Inserting extra diagrams...")

# Find the position after the last caption in BAB IV
last_caption_idx = None
for i in range(bab5_start or len(doc.paragraphs) - 1, (bab4_start or 0) - 1, -1):
    text = doc.paragraphs[i].text.strip()
    if text.startswith('Gambar 4.'):
        last_caption_idx = i
        break

if last_caption_idx:
    # Insert extra diagrams after the last caption
    insert_after = doc.paragraphs[last_caption_idx]._element
    
    for caption_text, img_file in EXTRA_DIAGRAMS:
        img_path = os.path.join(DIAGRAMS_DIR, img_file)
        if not os.path.exists(img_path):
            print(f"  ⚠️ Not found: {img_file}")
            continue
        
        try:
            # Add image
            img_para = doc.add_paragraph()
            img_para.alignment = 1  # CENTER
            img_para.add_run().add_picture(img_path, width=Inches(5.5))
            
            # Move after last_caption
            parent = insert_after.getparent()
            pos = list(parent).index(insert_after) + 1
            img_elem = img_para._element
            parent.remove(img_elem)
            parent.insert(pos, img_elem)
            
            # Add caption after image
            cap_para = doc.add_paragraph()
            cap_para.alignment = 1  # CENTER
            cap_para.add_run(f"Gambar {caption_text}")
            cap_elem = cap_para._element
            parent.remove(cap_elem)
            parent.insert(pos + 1, cap_elem)
            
            # Update insert_after for next diagram
            insert_after = cap_elem
            
            print(f"  ✅ {caption_text[:40]} + {img_file}")
        except Exception as e:
            print(f"  ❌ {caption_text[:30]}: {e}")

# ===== SAVE =====
print(f"\n💾 Saving...")
doc.save(OUTPUT)

# ===== VERIFY =====
doc2 = Document(OUTPUT)
img_count = 0
for para in doc2.paragraphs:
    for run in para.runs:
        drawings = run._element.findall(qn('w:drawing'))
        img_count += len(drawings)

print(f"\n📊 Verification:")
print(f"  Images: {img_count}")
print(f"  Size: {os.path.getsize(OUTPUT)/1024:.0f} KB")
