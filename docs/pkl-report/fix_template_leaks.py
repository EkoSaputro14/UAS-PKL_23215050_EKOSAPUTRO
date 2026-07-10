"""
Fix template leaks v3 — merge all runs per paragraph, replace, write back.
Buka dokumen yang sudah ada, fix yang salah, simpan. TIDAK rebuild dari awal.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SRC = r'C:\Users\SMANSA\mimotes\docs\pkl-report\LAPORAN_PKL_Eko_Saputro_23215050.docx'

CORRECT_TITLE = "RANCANG BANGUN SISTEM CHATBOT AI BERBASIS PENGETAHUAN DENGAN RETRIEVAL-AUGMENTED GENERATION DAN PIPELINE CRM UNTUK OPTIMALISASI LAYANAN PELANGGAN"

# Paragraph-level: merge all run text, then do replace on merged string
REPLACEMENTS = [
    # Cover title — old full title (may span multiple runs)
    ("Implementasi Keamanan Aplikasi dan Manajemen State pada Sistem Informasi Manajemen SDM Berbasis Web Menggunakan Vue.js 3 dan Pinia di PT Agsya Karya Manca", CORRECT_TITLE),
    ("Implementasi Keamanan Aplikasi dan Manajemen State pada Sistem Informasi Manajemen SDM Berbasis Web Menggunakan Vue.js 3 dan Pinia di PT", CORRECT_TITLE),
    ("Implementasi Keamanan Aplikasi dan Manajemen State pada Sistem RETRIEVAL-AUGMENTED", CORRECT_TITLE[:60]),
    # Name/ID/NIM
    ("Muhammad Affif", "Eko Saputro"),
    ("Affif", "Eko Saputro"),
    ("24225046", "23215050"),
    ("Syefudin", "Zaenul Arif"),
    # Company
    ("PT Agsya Karya Manca", "Bank Mandiri KCP Tegal Sudirman"),
    ("PT Agsya", "Bank Mandiri"),
    ("Agsya Karya Manca", "Bank Mandiri KCP Tegal Sudirman"),
    ("Agsya", "Bank Mandiri"),
    # TOC / body tech replacements
    ("Vue.js 3", "Next.js"),
    ("Vue.js", "Next.js"),
    ("Pinia", "State Management"),
    ("Go Fiber", "Next.js API Routes"),
    ("JWT", "NextAuth"),
    ("OAuth2", "Credentials"),
    ("Clean Architecture", "App Router"),
]

doc = Document(SRC)
changes = 0

for i, p in enumerate(doc.paragraphs):
    # Merge all run text into one string
    runs = p.runs
    if not runs:
        continue
    merged = ''.join(r.text for r in runs)
    if not merged.strip():
        continue

    original = merged
    for old, new in REPLACEMENTS:
        if old in merged:
            merged = merged.replace(old, new)

    if merged != original:
        # Write merged text to first run, clear rest
        runs[0].text = merged
        for r in runs[1:]:
            r.text = ''
        changes += 1
        print(f"  Fixed para[{i}]: {p.text[:80]}")

# Also fix tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                runs = p.runs
                if not runs:
                    continue
                merged = ''.join(r.text for r in runs)
                original = merged
                for old, new in REPLACEMENTS:
                    if old in merged:
                        merged = merged.replace(old, new)
                if merged != original:
                    runs[0].text = merged
                    for r in runs[1:]:
                        r.text = ''

doc.save(SRC)
print(f"\nFixed {changes} paragraphs. Saved: {SRC}")
print(f"Size: {os.path.getsize(SRC) / 1024:.0f} KB")
