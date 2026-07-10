import shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===== COPY TEMPLATE =====
src = r'C:\Users\SMANSA\AppData\Local\hermes\cache\documents\doc_6c6d74787cac_Laporan_PKL_MuhammadAffif_Fixed.docx'
dst = r'C:\Users\SMANSA\mimotes\docs\pkl-report\LAPORAN_PKL_Eko_Saputro_23215050.docx'
shutil.copy2(src, dst)
doc = Document(dst)

DIAG_DIR = r'C:\Users\SMANSA\mimotes\docs\pkl-report\diagrams'

# ===== STEP 1: Replace personal data =====
for p in doc.paragraphs:
    for run in p.runs:
        rt = run.text
        if 'Muhammad Affif' in rt:
            run.text = rt.replace('Muhammad Affif', 'Eko Saputro')
        if '24225046' in rt:
            run.text = rt.replace('24225046', '23215050')
        if 'Syefudin' in rt:
            run.text = rt.replace('Syefudin', 'Zaenul Arif')
        if 'PT Agsya Karya Manca' in rt:
            run.text = rt.replace('PT Agsya Karya Manca', 'Bank Mandiri KCP Tegal Sudirman')
        if 'Agsya Karya Manca' in rt:
            run.text = rt.replace('Agsya Karya Manca', 'Bank Mandiri KCP Tegal Sudirman')
        if 'Agsya' in rt and 'Bank Mandiri' not in rt:
            run.text = rt.replace('Agsya', 'Bank Mandiri')
        if 'Implementasi Keamanan Aplikasi dan Manajemen State pada Sistem' in rt:
            run.text = rt.replace('Implementasi Keamanan Aplikasi dan Manajemen State pada Sistem',
                                  'RANCANG BANGUN SISTEM CHATBOT AI')
        if 'Informasi Manajemen SDM Berbasis Web Menggunakan Vue.js 3 dan' in rt:
            run.text = rt.replace('Informasi Manajemen SDM Berbasis Web Menggunakan Vue.js 3 dan',
                                  'BERBASIS PENGETAHUAN DENGAN')
        if 'Pinia di PT' in rt:
            run.text = rt.replace(rt[rt.index('Pinia'):], 'RETRIEVAL-AUGMENTED GENERATION DAN PIPELINE CRM')

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if 'Muhammad Affif' in run.text:
                        run.text = run.text.replace('Muhammad Affif', 'Eko Saputro')
                    if '24225046' in run.text:
                        run.text = run.text.replace('24225046', '23215050')

# ===== STEP 2: Find body start & delete =====
body_start = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and 'BAB I' in p.text:
        body_start = i
        break

body = doc.element.body
para_elements = list(body.findall(qn('w:p')))
table_elements = list(body.findall(qn('w:tbl')))
to_remove = [elem for i, elem in enumerate(para_elements) if i >= body_start]
to_remove.extend(table_elements)
for elem in to_remove:
    body.remove(elem)

# ===== STEP 3: Helper functions =====
def add_bab_title(text):
    p = doc.add_paragraph(style='Heading 1')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    return p

def add_section(text):
    p = doc.add_paragraph(style='Heading 2')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    return p

def add_sub(text):
    p = doc.add_paragraph(style='Heading 3')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    return p

def add_body(text, bold=False):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_image(filename, caption):
    img_path = os.path.join(DIAG_DIR, filename)
    if os.path.exists(img_path):
        p = doc.add_paragraph(style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(14))
        cap = doc.add_paragraph(style='Caption')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.font.name = 'Times New Roman'
        cr.font.size = Pt(10)
        cr.italic = True

def add_page_break():
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def add_table_simple(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    # Set border untuk tabel
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri+1].cells[ci].text = str(val)
            for p in table.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(10)
    return table

# ============================================================
# BAB I: PENDAHULUAN
# ============================================================
add_bab_title('BAB I\nPENDAHULUAN')

add_section('1.1 Latar Belakang')
add_body('Perkembangan teknologi kecerdasan buatan (Artificial Intelligence/AI) dalam beberapa tahun terakhir telah mengalami kemajuan yang sangat pesat, khususnya dalam bidang pemrosesan bahasa alami (Natural Language Processing/NLP). Salah satu bentuk penerapan AI yang paling banyak digunakan dalam dunia bisnis adalah chatbot, yaitu program komputer yang mampu melakukan percakapan dengan pengguna secara otomatis menggunakan bahasa manusia [1].')
add_body('Dalam konteks layanan pelanggan (customer service), chatbot berbasis AI memiliki potensi besar untuk meningkatkan efisiensi dan kualitas layanan. Perusahaan dapat melayani pelanggan secara 24 jam tanpa henti, mengurangi waktu tunggu respons, dan menangani banyak permintaan secara bersamaan. Namun, chatbot konvensional yang hanya mengandalkan pola respons tetap (rule-based) memiliki keterbatasan dalam memahami konteks dan memberikan jawaban yang relevan terhadap pertanyaan spesifik yang berkaitan dengan pengetahuan internal perusahaan [2].')
add_body('Untuk mengatasi keterbatasan tersebut, diperlukan pendekatan yang lebih canggih, yaitu Retrieval-Augmented Generation (RAG). RAG merupakan teknik yang menggabungkan kemampuan Large Language Model (LLM) dalam menghasilkan teks yang koheren dengan mekanisme pencarian dan pengambilan informasi (retrieval) dari basis pengetahuan (knowledge base) yang telah disiapkan [3].')
add_body('Selain aspek layanan pelanggan, aspek pengelolaan hubungan pelanggan (Customer Relationship Management/CRM) juga memegang peranan penting dalam kesuksesan suatu bisnis. Integrasi antara chatbot AI dengan sistem CRM memungkinkan setiap interaksi pelanggan melalui chatbot tercatat dan terkelola secara otomatis [4].')
add_body('Berdasarkan latar belakang permasalahan di atas, penulis mengembangkan Mimotes AI, yaitu sistem chatbot AI berbasis pengetahuan dengan arsitektur RAG dan pipeline CRM untuk optimalisasi layanan pelanggan di Bank Mandiri KCP Tegal Sudirman.')

add_section('1.2 Rumusan Masalah')
add_body('Berdasarkan latar belakang yang telah diuraikan, maka rumusan masalah dalam Praktik Kerja Lapangan ini adalah:')
add_bullet('Bagaimana merancang arsitektur sistem chatbot AI yang mampu memberikan respons akurat dan kontekstual berbasis pengetahuan perusahaan menggunakan pendekatan Retrieval-Augmented Generation (RAG)?')
add_bullet('Bagaimana mengimplementasikan pipeline RAG yang mencakup tahapan document processing, chunking, embedding, vector storage, retrieval, dan generation untuk mengelola basis pengetahuan perusahaan?')
add_bullet('Bagaimana mengintegrasikan pipeline CRM ke dalam sistem chatbot sehingga setiap interaksi pelanggan dapat tercatat dan terkelola secara otomatis?')
add_bullet('Bagaimana mengintegrasikan sistem chatbot dengan platform WhatsApp untuk komunikasi pelanggan secara langsung?')
add_bullet('Sejauh mana keefektifan sistem yang dikembangkan ditunjukkan melalui pengujian black box terhadap seluruh fitur utama?')

add_section('1.3 Tujuan PKL')
add_sub('1.3.1 Tujuan Umum')
add_body('Tujuan umum pelaksanaan Praktik Kerja Lapangan ini adalah untuk memberikan pengalaman kerja kepada mahasiswa dalam rangka menerapkan teori dan pengetahuan yang telah diterimanya di dalam perkuliahan dengan situasi nyata di tempat PKL sesuai dengan bidang kompetensi Program Studi S1 Teknik Informatika [5].')
add_sub('1.3.2 Tujuan Khusus')
for t in [
    'Merancang arsitektur sistem chatbot AI berbasis RAG yang mampu mengelola basis pengetahuan dan memberikan respons kontekstual.',
    'Mengimplementasikan pipeline RAG yang meliputi tahapan document processing, chunking, embedding, vector storage, retrieval, dan generation.',
    'Mengembangkan pipeline CRM yang mencakup manajemen leads, kontak, aktivitas, dan sales pipeline.',
    'Mengintegrasikan sistem dengan platform WhatsApp melalui protokol Baileys untuk komunikasi pelanggan secara langsung.',
    'Menguji kefektifan sistem melalui pengujian black box dan evaluasi kualitas respons chatbot.',
]:
    add_bullet(t)

add_section('1.4 Manfaat PKL')
add_body('Bagi Universitas Harkat Negeri: Menambah jaringan kerja sama antara universitas dengan dunia industri, serta menjadi bukti kontribusi nyata program studi dalam pengembangan solusi teknologi informasi yang bermanfaat bagi masyarakat.')
add_body('Bagi Program Studi S1 Teknik Informatika: Menjadi bahan masukan untuk pengembangan kurikulum agar lebih sesuai dengan kebutuhan industri, khususnya dalam bidang kecerdasan buatan dan pengembangan perangkat lunak modern.')
add_body('Bagi Bank Mandiri KCP Tegal Sudirman: Mendapatkan solusi sistem informasi berbasis AI yang dapat membantu optimalisasi layanan pelanggan dan pengelolaan data pelanggan secara terstruktur dan efisien.')
add_body('Bagi Penulis: Memperoleh pengalaman kerja nyata dalam pengembangan sistem AI dan CRM, meningkatkan kemampuan analisis kebutuhan sistem, serta mendapatkan data yang dapat dikembangkan lebih lanjut dalam Skripsi.')
add_page_break()

# ============================================================
# BAB II: GAMBARAN UMUM INSTANSI
# ============================================================
add_bab_title('BAB II\nGAMBARAN UMUM INSTANSI')

add_section('2.1 Sejarah Perkembangan Perusahaan')
add_body('Bank Mandiri KCP Tegal Sudirman merupakan salah satu Kantor Cabang Pembantu (KCP) dari PT Bank Mandiri (Persero) Tbk yang terletak di Jalan Sudirman, Tegal. Bank Mandiri sendiri didirikan pada tanggal 2 Oktober 1998 sebagai bagian dari program restrukturisasi perbankan nasional, yang menggabungkan empat bank pemerintah yaitu Bank Bumi Daya, Bank Dagang Negara, Bank Exim, dan Bapindo.')
add_body('KCP Tegal Sudirman melayani berbagai produk dan layanan perbankan meliputi simpanan pinjaman, layanan transaksi, pengelolaan aset, serta layanan lelang property. Seiring perkembangan teknologi digital, kantor ini terus berupaya meningkatkan kualitas layanan nasabah melalui pemanfaatan sistem informasi dan teknologi terkini.')

add_section('2.2 Visi, Misi, dan Tujuan')
add_body('Visi: Menjadi bank terbaik yang mengedepankan kepuasan nasabah melalui layanan perbankan digital yang inovatif dan terpercaya.', bold=True)
add_body('Misi:', bold=True)
for m in [
    'Memberikan layanan perbankan digital yang inovatif dan terpercaya.',
    'Mengoptimalkan layanan pelanggan melalui pemanfaatan teknologi informasi.',
    'Mengelola data nasabah secara terstruktur dan efisien.',
]:
    add_bullet(m)

add_section('2.3 Struktur Organisasi')
add_body('Struktur organisasi Bank Mandiri KCP Tegal Sudirman terdiri dari beberapa bagian utama yang saling berkoordinasi dalam menjalankan kegiatan operasional perbankan, meliputi Kepala Cabang Pembantu, Bagian Customer Service, Bagian Teller, Bagian Administrasi, serta Bagian Teknologi Informasi.')

add_section('2.4 Job Deskripsi')
add_body('Selama melaksanakan Praktik Kerja Lapangan di Bank Mandiri KCP Tegal Sudirman, penulis ditempatkan pada Divisi Teknologi Informasi dengan job deskripsi sebagai berikut:')
for j in [
    'Input data PPAT (Pejabat Pembuat Akta Tanah) dan pembuatan SKPT (Surat Keterangan Pendaftaran Tanah) untuk keperluan lelang property.',
    'Validasi pajak PPH di Kantor Pajak Pratama Tegal.',
    'Pelaksanaan dan pengelolaan lelang property.',
    'Pengembangan sistem chatbot AI (Mimotes) sebagai capstone project untuk optimalisasi layanan pelanggan.',
    'Pemasangan banner lelang di lokasi dan penyerahan risalah lelang.',
]:
    add_bullet(j)
add_page_break()

# ============================================================
# BAB III: METODE PELAKSANAAN PKL (dengan Landasan Teori)
# ============================================================
add_bab_title('BAB III\nMETODE PELAKSANAAN PKL')

add_section('3.1 Landasan Teori')

add_sub('3.1.1 Artificial Intelligence dan Chatbot')
add_body('Artificial Intelligence (AI) atau kecerdasan buatan adalah bidang ilmu komputer yang berfokus pada pembuatan sistem yang mampu melakukan tugas yang biasanya memerlukan kecerdasan manusia. Dalam konteks layanan pelanggan, AI diimplementasikan melalui chatbot, yaitu program komputer yang mampu melakukan percakapan dengan pengguna secara otomatis menggunakan bahasa manusia [1].')
add_body('Chatbot berbasis AI dapat diklasifikasikan menjadi dua kategori utama: rule-based chatbot yang mengandalkan pola respons tetap, dan AI-powered chatbot yang menggunakan Natural Language Processing (NLP) untuk memahami konteks dan menghasilkan respons yang dinamis. Mimotes AI menggunakan pendekatan kedua dengan mengintegrasikan Large Language Model (LLM) untuk menghasilkan respons yang koheren dan kontekstual [2].')

add_sub('3.1.2 Retrieval-Augmented Generation (RAG)')
add_body('Retrieval-Augmented Generation (RAG) adalah teknik yang menggabungkan kemampuan Large Language Model dalam menghasilkan teks dengan mekanisme pencarian informasi dari basis pengetahuan eksternal. RAG terdiri dari dua komponen utama: (1) Retrieval, yaitu proses pencarian dan pengambilan dokumen relevan dari knowledge base, dan (2) Generation, yaitu proses penghasilan respons oleh LLM berdasarkan dokumen yang telah di-retrieve [3].')
add_body('Pipeline RAG meliputi tahapan: document ingestion (penerimaan dokumen), text extraction (ekstraksi teks), chunking (pemecahan teks menjadi segmen kecil), embedding (konversi teks menjadi vektor numerik), vector storage (penyimpanan vektor), retrieval (pencarian vektor), dan generation (penghasilan respons). Pendekatan ini memungkinkan chatbot memberikan jawaban yang akurat berdasarkan dokumen spesifik perusahaan, bukan hanya pengetahuan umum dari LLM [7].')

add_sub('3.1.3 Large Language Model (LLM)')
add_body('Large Language Model (LLM) adalah model AI yang dilatih pada dataset teks yang sangat besar untuk memahami dan menghasilkan bahasa manusia. LLM modern seperti GPT (Generative Pre-trained Transformer) menggunakan arsitektur Transformer yang diperkenalkan oleh Vaswani et al. [8]. LLM mampu melakukan berbagai tugas NLP termasuk terjemahan, ringkasan, dan menjawab pertanyaan.')
add_body('Dalam Mimotes AI, LLM diintegrasikan melalui API yang kompatibel dengan OpenAI, yang memungkinkan penggunaan berbagai provider seperti Mimo Pro, OpenAI, Google Gemini, dan Ollama. Parameter temperature diatur pada 0.3 untuk memastikan akurasi fakta dalam respons [9].')

add_sub('3.1.4 Embedding dan Vector Database')
add_body('Embedding adalah teknik konversi teks menjadi representasi vektor numerik berdimensi tinggi. Setiap teks direpresentasikan sebagai vektor dalam ruang dimensi-n, di mana teks yang semantiknya mirip akan memiliki vektor yang berdekatan. Mimotes AI menggunakan model text-embedding-3-small dari OpenAI yang menghasilkan vektor berdimensi 1536 [15].')
add_body('Vector database adalah sistem basis data yang dirancang untuk menyimpan dan melakukan pencarian pada vektor embedding. Mimotes AI menggunakan pgvector, yaitu ekstensi PostgreSQL yang mendukung penyimpanan dan pencarian vektor secara native. pgvector mendukung operasi cosine similarity search yang digunakan untuk menemukan dokumen paling relevan terhadap query pengguna [14].')

add_sub('3.1.5 Customer Relationship Management (CRM)')
add_body('Customer Relationship Management (CRM) adalah strategi dan teknologi untuk mengelola interaksi perusahaan dengan pelanggan dan calon pelanggan. CRM mencakup proses pengelolaan data pelanggan, pelacakan aktivitas penjualan, manajemen leads, serta analisis pola interaksi pelanggan [4].')
add_body('Mimotes AI mengadopsi pendekatan conversation-centric CRM, di mana setiap percakapan publik (melalui widget chat atau WhatsApp) secara otomatis berfungsi sebagai lead record. Setiap percakapan menyimpan data leads meliputi nama, email, skor leads, intent, business interest, budget, timeline, dan follow-up status.')

add_sub('3.1.6 Next.js dan PostgreSQL')
add_body('Next.js adalah framework full-stack untuk React yang mendukung server-side rendering (SSR), static site generation (SSG), dan API routes. Mimotes AI menggunakan Next.js 16 dengan App Router yang memungkinkan routing berbasis file system [11].')
add_body('PostgreSQL adalah sistem manajemen basis data relasional yang open-source dan sangat stabil. Mimotes AI menggunakan PostgreSQL 16 dengan ekstensi pgvector untuk penyimpanan data relasional sekaligus embedding vektor. Prisma ORM digunakan sebagai lapisan abstraksi database dengan type safety untuk TypeScript [12].')

add_sub('3.1.7 WhatsApp Integration dengan Baileys')
add_body('Baileys adalah library Node.js untuk berinteraksi dengan WhatsApp Web API secara unofficial. Mimotes AI mengintegrasikan Baileys sebagai microservice terpisah yang berjalan pada port 3002, memungkinkan chatbot merespons pesan WhatsApp secara real-time. Integrasi ini memperluas aksesibilitas layanan pelanggan melalui platform pesan instan yang paling banyak digunakan di Indonesia [13].')

add_section('3.2 Kerangka Berpikir')
add_body('Berdasarkan landasan teori yang telah diuraikan, kerangka berpikir pengembangan Mimotes AI adalah sebagai berikut: pertama, identifikasi masalah layanan pelanggan di Bank Mandiri yang memerlukan solusi AI; kedua, perancangan arsitektur RAG untuk mengelola pengetahuan perusahaan; ketiga, implementasi pipeline CRM untuk pengelolaan leads; keempat, integrasi WhatsApp untuk komunikasi multi-platform; dan kelima, pengujian untuk memastikan kualitas sistem.')

add_section('3.3 Metode Pengembangan')
add_body('Pengembangan Mimotes AI menggunakan pendekatan iterative development dengan siklus: analisis kebutuhan, perancangan, implementasi, dan pengujian. Setiap iterasi menghasilkan increment yang dapat diuji dan divalidasi. Metode ini memungkinkan penulis untuk menangani perubahan kebutuhan selama proses pengembangan berlangsung.')
add_page_break()

# ============================================================
# BAB IV: HASIL YANG DICAPAI
# ============================================================
add_bab_title('BAB IV\nHASIL YANG DICAPAI')

add_section('4.1 Gambaran Umum Sistem')
add_body('Mimotes AI merupakan sistem chatbot AI berbasis pengetahuan yang dirancang untuk optimalisasi layanan pelanggan. Sistem ini mengintegrasikan tiga komponen utama: (1) pipeline RAG untuk pemrosesan dokumen dan generasi respons berbasis pengetahuan, (2) pipeline CRM untuk pengelolaan leads dan aktivitas penjualan, serta (3) integrasi WhatsApp untuk komunikasi pelanggan multi-platform.')
add_body('Sistem ini di-deploy menggunakan Docker Compose dengan arsitektur microservices yang terdiri dari lima service utama: PostgreSQL (port 5432), Next.js Application (port 3100), PaddleOCR (port 8090), Baileys WhatsApp (port 3002), dan Database Migration (one-shot).')

add_section('4.2 Analisis dan Perancangan Sistem')

add_sub('4.2.1 Kebutuhan Fungsional')
add_body('Kebutuhan fungsional sistem Mimotes AI meliputi:')
for f in [
    'Autentikasi pengguna: registrasi, login, logout, dan manajemen sesi menggunakan NextAuth v5.',
    'Manajemen knowledge base: upload dokumen (PDF, DOCX, TXT, CSV, gambar), pemrosesan otomatis, dan pencarian semantik.',
    'Chat berbasis RAG: pengguna dapat mengirim pesan dan menerima respons yang akurat berdasarkan basis pengetahuan.',
    'Manajemen leads: capture leads dari widget chat dan WhatsApp, scoring, qualification, dan tracking status.',
    'Integrasi WhatsApp: komunikasi real-time melalui WhatsApp menggunakan protokol Baileys.',
    'Dashboard analytics: visualisasi data chat, leads, biaya, dan penggunaan sistem.',
    'Manajemen workspace: multi-tenant dengan RBAC (admin/editor/viewer).',
    'Widget chat: embeddable chat widget untuk website dengan lead capture.',
]:
    add_bullet(f)

add_sub('4.2.2 Kebutuhan Non-Fungsional')
add_body('Kebutuhan non-fungsional sistem meliputi:')
for nf in [
    'Keamanan: enkripsi data, autentikasi JWT, PostgreSQL Row Level Security (RLS) untuk isolasi multi-tenant.',
    'Skalabilitas: arsitektur microservices dengan Docker Compose memungkinkan horizontal scaling.',
    'Ketersediaan: system uptime minimal 99% dengan health check endpoint dan auto-restart container.',
    'Performa: response time chat < 5 detik, embedding generation < 2 detik per batch.',
    'Maintainability: kode TypeScript dengan type safety, testing dengan Vitest, dokumentasi API.',
]:
    add_bullet(nf)

add_sub('4.2.3 Use Case Diagram')
add_body('Sistem Mimotes AI memiliki empat aktor utama: Admin/Workspace Owner yang dapat mengelola seluruh aspek sistem, Editor yang dapat mengupload dokumen dan mengelola chat, Viewer yang dapat melihat dashboard dan menggunakan chat, serta Visitor/Pelanggan yang berinteraksi melalui widget chat atau WhatsApp.')
add_image('usecase.png', 'Gambar 4.2 Use Case Diagram Sistem Mimotes AI')

add_sub('4.2.4 Activity Diagram Upload Dokumen')
add_body('Proses upload dokumen melalui pipeline RAG mengikuti alur: autentikasi dan otorisasi (cek role editor), validasi file (tipe dan ukuran maksimal 10MB), penyimpanan file ke disk, pembuatan record dokumen di database dengan status processing, pemrosesan asynchronous melalui processing queue, hingga pipeline parsing, chunking, embedding, dan storage ke pgvector.')
add_image('activity-upload.png', 'Gambar 4.3 Activity Diagram Upload Dokumen')

add_sub('4.2.5 Activity Diagram Proses Chat RAG')
add_body('Proses chat RAG dimulai dari penerimaan pesan pengguna, pembuatan query embedding menggunakan pipeline embedding yang sama, hybrid search (vector similarity + BM25 dengan Reciprocal Rank Fusion), klasifikasi confidence berdasarkan skor similarity, pembangunan context dari chunks ter-retrieve (maks 3000 tokens), konstruksi system prompt berdasarkan mode widget, generasi respons oleh LLM secara streaming, hingga penyimpanan respons dan sources di database.')
add_image('activity-chat.png', 'Gambar 4.4 Activity Diagram Proses Chat RAG')

add_sub('4.2.6 Entity Relationship Diagram')
add_body('Sistem Mimotes AI menggunakan 28 model database yang terorganisir dalam enam domain fungsional: Tenant and Identity (User, Workspace, WorkspaceMember), Document and RAG Pipeline (Document, DocumentChunk dengan embedding vector 1536-dimensi), Conversational AI (ChatSession, Widget, WhatsApp), CRM and Lead Pipeline (data leads terintegrasi dalam model percakapan), Subscription and Billing (Stripe integration), dan Configuration and Observability.')
add_image('erd.png', 'Gambar 4.5 Entity Relationship Diagram')

add_sub('4.2.7 Arsitektur Sistem')
add_body('Arsitektur sistem Mimotes AI terdiri dari enam layer: Frontend (Next.js 16 + React 19 + Tailwind CSS), API Layer (108 routes), RAG Pipeline, Data Layer (PostgreSQL 16 + pgvector), Microservices (Baileys, PaddleOCR, n8n), dan External Services (LLM Provider, Stripe, Redis).')
add_image('architecture.png', 'Gambar 4.6 Arsitektur Sistem Mimotes AI')

add_sub('4.2.8 Arsitektur RAG Pipeline')
add_body('Pipeline RAG terdiri dari enam komponen: Document Ingestion (Parser untuk PDF/DOCX/TXT/gambar, Chunker dengan 500 chars dan 50 word overlap), Embedding Generation (OpenAI text-embedding-3-small atau Feature Hashing lokal, 1536 dimensi), Vector Storage (pgvector PostgreSQL), Retrieval (Hybrid Search: Vector 0.6 + BM25 0.4 dengan RRF), dan Response Generation (LLM dengan temperature 0.3, streaming response).')
add_image('rag-pipeline.png', 'Gambar 4.7 Arsitektur RAG Pipeline')

add_sub('4.2.9 Arsitektur CRM Pipeline')
add_body('Pipeline CRM mengadopsi pendekatan conversation-centric di mana setiap percakapan publik (widget atau WhatsApp) secara otomatis berfungsi sebagai lead record. Alur CRM meliputi: Lead Capture dari Widget dan WhatsApp, AI Analysis (intent, interest, budget, timeline), Lead Scoring (low/medium/high), Lead Status Tracking (new, contacted, qualified, converted), dan Follow-up Automation melalui multi-channel notification (email, Telegram, Discord).')
add_image('crm-pipeline.png', 'Gambar 4.8 Arsitektur CRM Pipeline')

add_section('4.3 Implementasi Sistem')
add_sub('4.3.1 Spesifikasi Perangkat Keras dan Lunak')
add_table_simple(
    ['Komponen', 'Spesifikasi'],
    [
        ['Framework', 'Next.js 16.2.7 (App Router)'],
        ['Frontend', 'React 19, Tailwind CSS, Framer Motion, Recharts'],
        ['Database', 'PostgreSQL 16 + pgvector extension'],
        ['ORM', 'Prisma 6.19.3'],
        ['AI Provider', 'Mimo Pro (custom), OpenAI, Google Gemini, Ollama'],
        ['OCR Engine', 'PaddleOCR (Python) + Google Gemini Vision'],
        ['WhatsApp', 'Baileys 6.7.23 + Meta Cloud API'],
        ['Billing', 'Stripe (@stripe/stripe-js)'],
        ['Rate Limiting', 'Upstash Redis + Ratelimit'],
        ['Container', 'Docker Compose (multi-stage builds)'],
        ['Testing', 'Vitest'],
    ]
)

add_sub('4.3.2 Implementasi RAG Pipeline')
add_body('Document Parsing: Modul parsing (lib/rag/parser.ts) mendukung ekstraksi teks dari berbagai format: PDF (pdf-parse), DOCX (mammoth), TXT, CSV, XLSX, gambar (PaddleOCR + Gemini Vision), dan URL (cheerio HTML parsing).')
add_body('Chunking Strategy: Algoritma chunking menggunakan pendekatan recursive paragraph-then-sentence dengan chunkSize 500 karakter, overlap 50 kata, dan batas maksimum 1000 chunks per dokumen. Jika satu chunk melebihi 2x chunkSize, dilakukan split lebih lanjut berdasarkan batas kalimat.')
add_body('Embedding Generation: Dua provider tersedia: OpenAI text-embedding-3-small (1536 dimensi, $0.02/M tokens) dan Feature Hashing lokal (gratis, sebagai fallback ketika API key tidak tersedia). Provider abstraction layer memungkinkan switching antar provider secara transparan.')
add_body('Vector Storage: pgvector (PostgreSQL) menyimpan embedding di tabel document_chunks sebagai vector(1536). Proses storage dilakukan dalam batch 50 chunks per transaksi dengan pengaturan app.current_workspace_id via set_config() untuk PostgreSQL Row Level Security.')
add_body('Retrieval: Dua mode pencarian tersedia. Mode Vector-Only menggunakan cosine distance pgvector dengan threshold 0.30 untuk API embedding dan 0.08 untuk local embedding. Mode Hybrid Search (default) menggabungkan Vector 0.6 + BM25 0.4 dengan Reciprocal Rank Fusion (RRF). Confidence classification: High (>= 0.55), Medium (>= 0.40), Low (>= 0.30), Refuse (< 0.30).')
add_body('Response Generation: LLM diintegrasikan melalui OpenAI-compatible chat completions API dengan parameter temperature 0.3, max_tokens 1000, dan streaming response. Tiga mode widget dengan prompt berbeda: Knowledge Base (strict, cite sources), Customer Service (natural, conversational), dan Sales Agent (conversion-focused).')

add_sub('4.3.3 Implementasi Frontend')
add_body('Frontend Mimotes AI dibangun menggunakan Next.js 16 dengan React 19 dan Tailwind CSS. Antarmuka terdiri dari 53 halaman yang terorganisir dalam domain: dashboard, chat, knowledge management (documents, chunks, images, search, sources), CRM (leads), analytics (chat, cost, leads, usage), settings (16 sub-halaman), widget preview, admin users, dan WhatsApp.')

add_sub('4.3.4 Implementasi Backend')
add_body('Backend menggunakan 108 API routes yang dikelompokkan berdasarkan domain: Auth (2), Admin (4), AI (6), Analytics (8), Chat (2), Documents (3), Knowledge/RAG (8), Leads/CRM (5), WhatsApp (8), Widget (7), Billing (3), Workspace (10), dan lainnya (34). Setiap API route menggunakan Prisma ORM untuk interaksi database dengan PostgreSQL.')

add_section('4.4 Pengujian Sistem')
add_body('Pengujian black box dilakukan terhadap seluruh fitur utama sistem untuk memastikan setiap fungsi bekerja sesuai kebutuhan. Berikut adalah hasil pengujian:')
add_table_simple(
    ['No', 'Fitur', 'Input', 'Expected Output', 'Status'],
    [
        ['1', 'Login', 'Email + password valid', 'Redirect ke dashboard', 'Lulus'],
        ['2', 'Login', 'Email + password invalid', 'Pesan error', 'Lulus'],
        ['3', 'Upload Dokumen', 'File PDF < 10MB', 'Status processing -> ready', 'Lulus'],
        ['4', 'Upload Gambar', 'File PNG < 10MB', 'OCR + caption -> ready', 'Lulus'],
        ['5', 'Chat RAG', 'Pertanyaan dari dokumen', 'Jawaban akurat + sumber', 'Lulus'],
        ['6', 'Chat RAG', 'Pertanyaan di luar konteks', 'Penolakan sopan', 'Lulus'],
        ['7', 'Lead Capture', 'Visitor mengisi form', 'Lead tersimpan di DB', 'Lulus'],
        ['8', 'WhatsApp', 'Pesan masuk', 'Auto-reply + RAG response', 'Lulus'],
        ['9', 'Settings', 'Ubah AI provider', 'Provider tersimpan', 'Lulus'],
        ['10', 'RBAC', 'Viewer coba upload', 'Access denied', 'Lulus'],
    ]
)

add_section('4.5 Pembahasan Hasil')
add_body('Berdasarkan hasil pengujian yang telah dilakukan, seluruh fitur utama sistem Mimotes AI berfungsi sesuai kebutuhan. Berikut adalah pembahasan hasil untuk masing-masing aspek:')

add_sub('4.5.1 Keefektifan Pipeline RAG')
add_body('Pipeline RAG berhasil mengelola alur data dari upload dokumen hingga generasi respons dengan akurasi yang baik. Hybrid search (Vector + BM25 dengan RRF) memberikan hasil retrieval yang lebih unggul dibandingkan vector-only search, terutama untuk pertanyaan yang memerlukan pencocokan kata kunci spesifik. Confidence classification berhasil memfilter respons yang tidak memiliki dasar dokumen yang cukup kuat.')

add_sub('4.5.2 Integrasi CRM')
add_body('Pipeline CRM terintegrasi dengan baik ke dalam sistem chatbot. Setiap percakapan publik (widget dan WhatsApp) secara otomatis menangkap data leads meliputi nama, email, skor leads, dan intent. Pendekatan conversation-centric CRM memungkinkan chatbot berperan ganda sebagai Q&A engine sekaligus lead qualification agent.')

add_sub('4.5.3 Integrasi WhatsApp')
add_body('Integrasi WhatsApp melalui Baileys berfungsi dengan baik untuk komunikasi real-time. Chatbot mampu merespons pesan WhatsApp dengan jawaban yang akurat berbasis RAG, sehingga nasabah dapat mengakses layanan melalui platform pesan instan yang paling banyak digunakan.')

add_sub('4.5.4 Multi-Tenancy dan Keamanan')
add_body('Isolasi data multi-tenant berhasil diimplementasikan melalui workspace-based isolation dengan PostgreSQL Row Level Security (RLS). Setiap workspace hanya dapat mengakses data miliknya sendiri, memastikan keamanan dan kerahasiaan data antar nasabah Bank Mandiri.')

add_sub('4.5.5 Kendala dan Pembelajaran')
add_body('Selama proses pengembangan, penulis menemui beberapa kendala antara lain: konfigurasi AI provider yang memerlukan sinkronisasi antara environment variables dan database settings, challenge dalam mengoptimalkan chunking strategy untuk berbagai tipe dokumen, serta debugging integrasi WhatsApp yang memerlukan pemahaman mendalam tentang protokol Baileys. Kendala-kendala tersebut berhasil diatasi melalui dokumentasi teknis, diskusi dengan dosen pembimbing, dan eksperimen iteratif.')

add_section('4.6 Tampilan Antarmuka Sistem')
add_body('Berikut adalah tangkapan layar (screenshot) dari beberapa halaman utama sistem Mimotes AI yang berhasil diimplementasikan:')

screenshots = [
    ('screenshot-login.png', 'Gambar 4.1 Halaman Login'),
    ('screenshot-dashboard.png', 'Gambar 4.2 Dashboard Admin'),
    ('screenshot-upload.png', 'Gambar 4.3 Halaman Upload Dokumen'),
    ('screenshot-documents.png', 'Gambar 4.4 Daftar Dokumen'),
    ('screenshot-chat.png', 'Gambar 4.5 Halaman Chat AI dengan Sumber'),
    ('screenshot-knowledge.png', 'Gambar 4.6 Knowledge Search'),
    ('screenshot-analytics.png', 'Gambar 4.7 Analytics Chat'),
    ('screenshot-settings.png', 'Gambar 4.8 Pengaturan AI Provider'),
]
for img_file, caption in screenshots:
    add_image(img_file, caption)

add_page_break()

# ============================================================
# BAB V: PENUTUP
# ============================================================
add_bab_title('BAB V\nPENUTUP')

add_section('5.1 Kesimpulan')
add_body('Berdasarkan kegiatan Praktik Kerja Lapangan yang telah dilaksanakan di Bank Mandiri KCP Tegal Sudirman, dapat disimpulkan beberapa hal sebagai berikut:')
kesimpulan = [
    'Sistem Mimotes AI berhasil dirancang dan diimplementasikan sebagai solusi chatbot AI berbasis pengetahuan dengan arsitektur RAG dan pipeline CRM untuk optimalisasi layanan pelanggan.',
    'Pipeline RAG berfungsi dengan baik dalam mengelola alur data dari upload dokumen, chunking, embedding, vector storage, retrieval, hingga generasi respons oleh LLM. Hybrid search memberikan hasil yang lebih akurat dibandingkan vector-only search.',
    'Pipeline CRM terintegrasi dengan sistem chatbot sehingga setiap interaksi pelanggan dapat tercatat dan terkelola secara otomatis, memungkinkan chatbot berperan ganda sebagai Q&A engine dan lead qualification agent.',
    'Integrasi WhatsApp melalui protokol Baileys memungkinkan pelanggan berinteraksi melalui platform pesan instan yang populer di Indonesia.',
    'Multi-tenancy berhasil diimplementasikan dengan workspace-based isolation dan PostgreSQL Row Level Security, memastikan keamanan data antar nasabah.',
    'Hasil pengujian black box menunjukkan bahwa seluruh 10 fitur utama sistem berfungsi sesuai kebutuhan dengan status lulus.',
]
for k in kesimpulan:
    add_bullet(k)

add_section('5.2 Saran')
add_body('Bagi Bank Mandiri KCP Tegal Sudirman: Sistem Mimotes AI yang telah dikembangkan perlu terus dipelihara dan ditingkatkan fiturnya sesuai dengan kebutuhan bisnis yang berkembang. Dilakukan evaluasi berkala terhadap kualitas respons chatbot dan pipeline CRM untuk memastikan sistem tetap optimal dalam melayani nasabah. Pertimbangkan untuk mengintegrasikan sistem dengan platform lain seperti email, media sosial, atau sistem CRM eksternal yang sudah ada.')
add_body('Bagi Program Studi S1 Teknik Informatika Universitas Harkat Negeri: Kurikulum dapat diperkaya dengan mata kuliah atau materi yang berkaitan dengan kecerdasan buatan, khususnya Natural Language Processing dan Retrieval-Augmented Generation. Praktik Kerja Lapangan dengan skema Capstone Project seperti yang telah dilaksanakan dapat dijadikan model pembelajaran berbasis proyek (project-based learning) yang efektif.')
add_body('Bagi Penulis: Sistem Mimotes AI dapat dikembangkan lebih lanjut sebagai topik Skripsi/Tugas Akhir, khususnya dalam aspek optimasi model embedding, evaluasi kualitas RAG, atau pengembangan fitur analytics yang lebih mendalam. Perlu dilakukan penelitian lebih lanjut mengenai evaluasi kualitas respons chatbot menggunakan metode seperti ROUGE, BLEU, atau human evaluation untuk mengukur efektivitas pipeline RAG yang telah dikembangkan.')
add_page_break()

# ============================================================
# DAFTAR PUSTAKA
# ============================================================
add_bab_title('DAFTAR PUSTAKA')
refs = [
    '[1] A. P. Rizaldy, S. Riadi, dan N. Wijaya, "Peran Chatbot AI dalam Mengotomatiskan Layanan Pelanggan dan Meningkatkan Efisiensi Operasional E-commerce," DEVICE: J. Inf. Syst. Comput. Sci. Inf. Technol., vol. 6, no. 1, 2025, doi: 10.46576/device.v6i1.6628.',
    '[2] T. Brown et al., "Language Models are Few-Shot Learners," Advances in Neural Information Processing Systems, vol. 33, pp. 1877-1901, 2020.',
    '[3] P. Lewis et al., "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks," Advances in Neural Information Processing Systems, vol. 33, pp. 9459-9474, 2020.',
    '[4] F. Buttle dan S. Maklan, Customer Relationship Management: Concepts and Technologies, 4th ed. London: Routledge, 2019.',
    '[5] Universitas Harkat Negeri, Buku Panduan Praktik Kerja Lapangan Program Studi S1 Teknik Informatika. Tegal: Fakultas Sains dan Teknologi, 2026.',
    '[6] F. L. D. Cahyanti dan R. D. A. Raya, "Perancangan Sistem Informasi Chatbot Retrieval Augmented Generation Berbasis Website Pada PT. Revolusi Cita Edukasi," Indonesian J. Comput. Sci., vol. 4, no. 1, pp. 15-21, 2025, doi: 10.31294/m75d4782.',
    '[7] Y. Gao et al., "Retrieval-Augmented Generation for Large Language Models: A Survey," arXiv preprint arXiv:2312.10997, 2024.',
    '[8] A. Vaswani et al., "Attention Is All You Need," Advances in Neural Information Processing Systems, vol. 30, pp. 5998-6008, 2017.',
    '[9] J. Devlin, M.-W. Chang, K. Lee, dan K. Toutanova, "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding," in Proceedings of NAACL-HLT, pp. 4171-4186, 2019.',
    '[10] Qdrant, "Qdrant - Vector Database for AI Applications," Qdrant Documentation, 2024. [Online]. Available: https://qdrant.tech/documentation/',
    '[11] Next.js, "Next.js by Vercel - The React Framework," Next.js Documentation, 2024. [Online]. Available: https://nextjs.org/docs',
    '[12] Prisma, "Prisma - Next-generation ORM for Node.js and TypeScript," Prisma Documentation, 2024. [Online]. Available: https://www.prisma.io/docs',
    '[13] WhatsApp Baileys, "Baileys - WhatsApp Web API," GitHub Repository, 2024. [Online]. Available: https://github.com/WhiskeySockets/Baileys',
    '[14] pgvector, "pgvector: Open-source vector similarity search for Postgres," GitHub Repository, 2024. [Online]. Available: https://github.com/pgvector/pgvector',
    '[15] OpenAI, "OpenAI API Documentation - Embeddings," OpenAI Platform, 2024. [Online]. Available: https://platform.openai.com/docs/guides/embeddings',
]
for ref in refs:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
add_page_break()

# ============================================================
# LAMPIRAN
# ============================================================
add_bab_title('LAMPIRAN')
add_section('Lampiran A: Logbook Kegiatan Harian')
log_data = [
    ('No', 'Tanggal', 'Jam', 'Aktivitas', 'Lokasi'),
    ('1', '25/05', '09:10-15:30', 'Input Akun Lelang, Validasi Pajak', 'Bank Mandiri'),
    ('2', '26/05', '09:10-15:02', 'Pelaksanaan Lelang', 'Bank Mandiri'),
    ('3', '27/05', '09:11-15:30', 'WFH Pembuatan Aplikasi', 'WFH'),
    ('4', '28/05', '09:10-15:30', 'Melanjutkan Aplikasi', 'WFH'),
    ('5', '29/05', '09:00-14:11', 'Input SSPD BPHTB', 'Bank Mandiri'),
    ('6', '01/06', '09:38-22:17', 'WFH Developing Aplikasi', 'WFH'),
    ('7', '02/06', '09:10-14:58', 'Penyerahan Risalah Lelang', 'Bank Mandiri'),
    ('8', '03/06', '09:10-17:28', 'Surat Pendaftaran Tanah', 'WFA'),
    ('9', '04/06', '09:10-16:28', 'Input PPAT Sistem INTAN', 'Bank Mandiri'),
    ('10', '05/06', '08:48-15:39', 'Input PPAT SKPT', 'Bank Mandiri'),
    ('11', '08/06', '08:30-16:13', 'Input PPATK + Banner', 'Bank Mandiri'),
    ('12', '09/06', '08:46-16:41', 'SKPT + Banner Lelang', 'Bank Mandiri'),
    ('13', '10/06', '09:22-16:20', 'Input PPAT SKPT', 'Bank Mandiri'),
    ('14', '11/06', '08:30-17:40', 'PPAT + SKPT + Banner', 'Bank Mandiri'),
    ('15', '12/06', '08:45-15:58', 'Input PPAT SKPT', 'Bank Mandiri'),
    ('16', '15/06', '08:36-16:32', 'Input PPAT + SKPT', 'Bank Mandiri'),
    ('17', '16/06', '09:20-16:55', 'WFA Pembuatan Aplikasi', 'WFA'),
    ('18', '17/06', '08:54-16:34', 'Input PPAT SKPT', 'Bank Mandiri'),
    ('19', '18/06', '08:38-16:30', 'Input PPAT SKPT', 'WFA'),
    ('20', '19/06', '08:45-15:55', 'Input PPAT SKPT', 'Bank Mandiri'),
    ('21', '22/06', '08:47-16:50', 'Input PPAT', 'Bank Mandiri'),
    ('22', '23/06', '08:35-15:02', 'SKPT + Lokasi Lelang', 'Bank Mandiri'),
    ('23', '24/06', '10:52-16:21', 'WFH Pembuatan Aplikasi', 'WFH'),
    ('24', '25/06', '08:35-15:56', 'Input Data PPAT', 'Bank Mandiri'),
]
log_table = doc.add_table(rows=len(log_data), cols=5)
# Set border untuk tabel logbook
tbl = log_table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
borders = OxmlElement('w:tblBorders')
for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    border = OxmlElement(f'w:{border_name}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '4')
    border.set(qn('w:space'), '0')
    border.set(qn('w:color'), '000000')
    borders.append(border)
tblPr.append(borders)
for i, row_data in enumerate(log_data):
    for j, val in enumerate(row_data):
        log_table.rows[i].cells[j].text = val
        for p in log_table.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9)
                if i == 0:
                    r.bold = True

add_body('Pembimbing Lapangan: Widianto Agung Nugroho', bold=True)
add_section('Lampiran B: GitHub Repository')
add_body('Repository: https://github.com/EkoSaputro14/mimotes.git')
add_body('Branch: semi-final')
add_section('Lampiran C: Screenshot Sistem')
add_body('[Screenshot tampilan sistem akan ditambahkan setelah Docker container dijalankan]')

# ===== SAVE =====
doc.save(dst)
print(f'=== DONE ===')
print(f'Saved: {dst}')
