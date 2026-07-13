#!/usr/bin/env python3
"""
Rebuild v5: Complete clean rebuild.
1. Copy template front matter (cover, persetujuan, pengesahan, kata pengantar)
2. Clear EVERYTHING after that
3. Add Daftar Isi (empty), Daftar Tabel (empty), Daftar Gambar (empty)
4. Add BAB I-V from source with CORRECT style mapping
5. Add Daftar Pustaka, Lampiran
6. Insert images
"""
import os
import copy
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = 'C:/Users/SMANSA/AppData/Local/hermes/cache/documents/doc_4ce6543a497c_LAPORAN PKL (1).docx'
SOURCE = 'LAPORAN_PKL_v14_UML.docx'
OUTPUT = 'LAPORAN_PKL_v15_Template.docx'
DIAGRAMS_DIR = 'diagrams'

NEW_TITLE = "RANCANG BANGUN SISTEM CHATBOT AI BERBASIS PENGETAHUAN DENGAN RETRIEVAL-AUGMENTED GENERATION DAN PIPELINE CRM UNTUK OPTIMALISASI LAYANAN PELANGGAN"
NEW_STUDENT = "Eko Saputro"
NEW_NIM = "23215050"
NEW_SUPERVISOR = "Widianto Agung Nugroho"
OLD_STUDENT = "Moh. Arif Prasetyo"
OLD_NIM = "23215043"

print("📄 Loading...")
tpl = Document(TEMPLATE)
src = Document(SOURCE)

# ===== Find where front matter ends (Daftar Isi heading) =====
dai_idx = None
for i, para in enumerate(tpl.paragraphs):
    if para.style.name == 'Heading 1' and 'Daftar Isi' in para.text:
        dai_idx = i
        break

print(f"  Front matter ends at paragraph {dai_idx}")

# ===== Remove everything from Daftar Isi onwards =====
body = tpl.element.body
children = list(body)

# Find the Daftar Isi element
dai_elem_idx = None
for i, child in enumerate(children):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        pPr = child.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                text = ''.join(t.text or '' for t in child.iter(qn('w:t'))).strip()
                if 'Daftar Isi' in text:
                    dai_elem_idx = i
                    break

print(f"  Daftar Isi element index: {dai_elem_idx}")

if dai_elem_idx is not None:
    # Keep sectPr if it exists at the end
    sectPr = body.find(qn('w:sectPr'))
    
    # Remove everything from Daftar Isi onwards (except sectPr)
    to_remove = []
    for child in children[dai_elem_idx:]:
        if child is not sectPr:
            to_remove.append(child)
    
    for elem in to_remove:
        body.remove(elem)
    
    print(f"  Removed {len(to_remove)} elements")

# ===== Get style definitions from template =====
# Find the style ID for Heading 1, 2, 3 in the template
style_map = {}
for para in tpl.paragraphs[:dai_idx]:
    if para.style and para.style.name.startswith('Heading'):
        level = para.style.name.replace('Heading ', '')
        # Get the XML style ID
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                style_id = pStyle.get(qn('w:val'), '')
                style_map[f'Heading{level}'] = style_id
                style_map[f'Heading {level}'] = style_id

print(f"  Style map: {style_map}")

# ===== Extract source content =====
print("\n📝 Extracting source...")
src_sections = {}
current_key = None
current_paras = []

def classify_heading(text):
    first_line = text.strip().split('\n')[0].strip()
    if 'BAB V' in first_line: return 'BAB_V'
    if 'BAB IV' in first_line: return 'BAB_IV'
    if 'BAB III' in first_line: return 'BAB_III'
    if 'BAB II' in first_line: return 'BAB_II'
    if 'BAB I' in first_line: return 'BAB_I'
    if 'DAFTAR PUSTAKA' in first_line: return 'DAFTAR_PUSTAKA'
    if 'LAMPIRAN' in first_line: return 'LAMPIRAN'
    return None

for para in src.paragraphs:
    text = para.text.strip()
    style_name = para.style.name if para.style else 'Normal'
    is_h1 = style_name in ('Heading 1', 'Heading1')
    
    if is_h1:
        if current_key:
            src_sections[current_key] = current_paras
        current_paras = []
        current_key = classify_heading(text)
        continue
    
    if current_key and text:
        # Determine style
        if style_name in ('Heading 2', 'Heading2'):
            mapped_style = style_map.get('Heading 2', 'Heading2')
        elif style_name in ('Heading 3', 'Heading3'):
            mapped_style = style_map.get('Heading 3', 'Heading3')
        else:
            mapped_style = 'Normal'
        
        current_paras.append({
            'text': text,
            'style_id': mapped_style,
        })

if current_key:
    src_sections[current_key] = current_paras

for key, paras in src_sections.items():
    print(f"  {key}: {len(paras)} paragraphs")

# ===== Add new content =====
print("\n📝 Adding content...")

def add_heading(doc, text, level):
    """Add a heading paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    p.text = text
    return p

def add_normal(doc, text):
    """Add a normal paragraph."""
    p = doc.add_paragraph()
    p.text = text
    return p

# Add Daftar Isi (empty for user to fill)
add_heading(tpl, 'Daftar Isi', 1)
add_heading(tpl, 'Daftar  Tabel', 1)
add_heading(tpl, 'Daftar Gambar', 1)

# Add each BAB
for key in ['BAB_I', 'BAB_II', 'BAB_III', 'BAB_IV', 'BAB_V', 'DAFTAR_PUSTAKA', 'LAMPIRAN']:
    if key not in src_sections:
        continue
    
    paras = src_sections[key]
    
    # Add heading
    heading_texts = {
        'BAB_I': 'BAB I',
        'BAB_II': 'BAB II',
        'BAB_III': 'BAB III.',
        'BAB_IV': 'BAB IV',
        'BAB_V': 'BAB V',
        'DAFTAR_PUSTAKA': 'DAFTAR PUSTAKA',
        'LAMPIRAN': 'LAMPIRAN',
    }
    
    add_heading(tpl, heading_texts[key], 1)
    
    # Add subheading for BAB I
    if key == 'BAB_I':
        add_normal(tpl, 'PENDAHULUAN')
    
    # Add content
    for p in paras:
        style_id = p['style_id']
        text = p['text']
        
        if style_id in ('Heading2', 'Judul2'):
            # It's a heading 2
            add_heading(tpl, text, 2)
        elif style_id in ('Heading3', 'Judul3'):
            add_heading(tpl, text, 3)
        else:
            add_normal(tpl, text)

print("  ✅ Content added")

# ===== Fix front matter =====
print("\n📝 Fixing front matter...")
for para in tpl.paragraphs:
    if "SISTEM PAKAR" in para.text:
        for run in para.runs:
            if "SISTEM PAKAR" in run.text or "RANCANG BANGUN" in run.text:
                run.text = ""
        for run in para.runs:
            if not run.text.strip():
                run.text = NEW_TITLE
                break
    for run in para.runs:
        if OLD_STUDENT in run.text:
            run.text = run.text.replace(OLD_STUDENT, NEW_STUDENT)
        if OLD_NIM in run.text:
            run.text = run.text.replace(OLD_NIM, NEW_NIM)
        if "Nama Dosen Pembimbing" in run.text:
            run.text = run.text.replace("Nama Dosen Pembimbing", NEW_SUPERVISOR)
        if "NIPY." in run.text:
            run.text = run.text.replace("NIPY.", "")

print("  ✅ Done")

# ===== Save =====
print("\n💾 Saving...")
tpl.save(OUTPUT)

# ===== Insert images =====
print("\n🖼️ Inserting images...")
doc2 = Document(OUTPUT)

CAPTION_IMG = {
    "Use Case Diagram": "usecase.png",
    "Activity Diagram Upload": "activity-upload.png",
    "Activity Diagram Proses Chat": "activity-chat.png",
    "Sequence Diagram Chat RAG": "sequence-chat-rag.png",
    "ERD Domain Identity": "erd-a-identity.png",
    "ERD Domain RAG": "erd-b-rag.png",
    "ERD Domain Chat": "erd-c-chat-crm.png",
    "ERD Domain Billing": "erd-d-billing.png",
    "ERD Ringkasan": "erd-summary.png",
    "Arsitektur Sistem Mimotes": "architecture.png",
    "Arsitektur RAG": "rag-pipeline.png",
    "Arsitektur CRM": "crm-pipeline.png",
}

img_count = 0
for i, para in enumerate(doc2.paragraphs):
    text = para.text.strip()
    if not text.startswith('Gambar 4.'):
        continue
    for keyword, img_file in CAPTION_IMG.items():
        if keyword in text:
            img_path = os.path.join(DIAGRAMS_DIR, img_file)
            if not os.path.exists(img_path):
                continue
            try:
                new_para = doc2.add_paragraph()
                new_para.alignment = 1
                new_para.add_run().add_picture(img_path, width=Inches(5.5))
                parent = para._element.getparent()
                pos = list(parent).index(para._element)
                new_elem = new_para._element
                parent.remove(new_elem)
                parent.insert(pos, new_elem)
                img_count += 1
                print(f"  ✅ {img_file}")
            except Exception as e:
                print(f"  ❌ {img_file}: {e}")
            break

doc2.save(OUTPUT)
print(f"\n  Images: {img_count}")

# ===== Verify =====
doc3 = Document(OUTPUT)
total_imgs = sum(
    len(run._element.findall(qn('w:drawing')))
    for para in doc3.paragraphs
    for run in para.runs
)
old_words = ['Sistem Pakar', 'Dinas Pertanian', 'Brebes', 'SIPATAN', 'Certainty Factor', 'CF Kombinasi', 'Padi', 'Hama']
old_found = [w for w in old_words if any(w in p.text for p in doc3.paragraphs)]

print(f"\n📊 Final:")
print(f"  Paragraphs: {len(doc3.paragraphs)}")
print(f"  Images: {total_imgs}")
print(f"  Old content: {'⚠️ ' + str(old_found) if old_found else '✅ Clean'}")
print(f"  Size: {os.path.getsize(OUTPUT)/1024:.0f} KB")
