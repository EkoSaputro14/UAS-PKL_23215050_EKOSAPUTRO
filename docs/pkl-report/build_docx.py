from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(4); s.bottom_margin = Cm(3)
    s.left_margin = Cm(4); s.right_margin = Cm(3)
style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0

DIAG_DIR = r'C:\Users\SMANSA\mimotes\docs\pkl-report\diagrams'

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.bold = True
    h.paragraph_format.line_spacing = 2.0

def para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = bold

def center_para(text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold

def add_image(filename, caption):
    img_path = os.path.join(DIAG_DIR, filename)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
        r = p.add_run()
        r.add_picture(img_path, width=Cm(14))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing = 2.0
        cr = cap.add_run(caption)
        cr.font.name = 'Times New Roman'; cr.font.size = Pt(10); cr.italic = True
        print(f'  + {filename}')
    else:
        print(f'  - MISSING: {filename}')

def bullet(items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.left_indent = Cm(1.25)
        r = p.add_run(f'{i}. {item}')
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)

# ============================================================
# COVER
# ============================================================
for _ in range(4): doc.add_paragraph()
center_para('LAPORAN PRAKTIK KERJA LAPANGAN', True, 16)
doc.add_paragraph()
center_para('RANCANG BANGUN SISTEM CHATBOT AI\nBERBASIS PENGETAHUAN DENGAN\nRETRIEVAL-AUGMENTED GENERATION\nDAN PIPELINE CRM UNTUK\nOPTIMALISASI LAYANAN PELANGGAN', True, 14)
for _ in range(3): doc.add_paragraph()
center_para('Nama: Eko Saputro\nNIM: 23215050\nKelas: 6A')
doc.add_paragraph()
center_para('PROGRAM STUDI S1 TEKNIK INFORMATIKA\nFAKULTAS SAINS & TEKNOLOGI\nUNIVERSITAS HARKAT NEGERI\n\nTahun 2026')
doc.add_page_break()

# ============================================================
# LEMBAR PENGESAHAN
# ============================================================
center_para('LEMBAR PENGESAHAN', True, 14)
doc.add_paragraph()
para('Laporan Praktik Kerja Lapangan (PKL) ini telah disetujui dan dipertanggungjawabkan oleh:', indent=False)
doc.add_paragraph()
para('Dosen Pembimbing:', bold=True, indent=False)
para('Zaenul Arif, M.Kom.', indent=False)
doc.add_paragraph()
para('Pembimbing Lapangan:', bold=True, indent=False)
para('Widianto Agung Nugroho', indent=False)
doc.add_paragraph()
para('Ketua Program Studi:', bold=True, indent=False)
para('Aang Alim Murtopo, M.Kom.\nNIPY. 08.025.555', indent=False)
doc.add_paragraph(); doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Tegal, 25 Juni 2026'); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
doc.add_page_break()

# ============================================================
# KATA PENGANTAR
# ============================================================
center_para('KATA PENGANTAR', True, 14)
doc.add_paragraph()
para('Puji syukur kehadirat Tuhan Yang Maha Esa atas rahmat dan karunia-Nya sehingga laporan Praktik Kerja Lapangan (PKL) Program Studi S1 Teknik Informatika Universitas Harkat Negeri ini dapat diselesaikan dengan baik.')
para('Laporan ini disusun sebagai bentuk pertanggungjawaban akademik atas kegiatan Praktik Kerja Lapangan yang telah dilaksanakan di Bank Mandiri KCP Tegal Sudirman.')
para('Penulis mengucapkan terima kasih kepada: Bapak Zaenul Arif, M.Kom. (Dosen Pembimbing), Bapak Widianto Agung Nugroho (Pembimbing Lapangan), Bapak Aang Alim Murtopo, M.Kom. (Ketua Prodi), serta seluruh pihak yang telah membantu.')
para('Semoga laporan ini dapat memberikan manfaat bagi pembaca dan pengembangan ilmu pengetahuan di bidang Teknik Informatika.')
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Tegal, 25 Juni 2026\n\nPenulis\n\nEko Saputro\nNIM: 23215050')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
doc.add_page_break()

# ============================================================
# DAFTAR ISI
# ============================================================
center_para('DAFTAR ISI', True, 14)
doc.add_paragraph()
toc = [
    'BAB I PENDAHULUAN',
    '  1.1 Latar Belakang',
    '  1.2 Tujuan PKL',
    '  1.3 Manfaat PKL',
    'BAB II GAMBARAN UMUM INSTANSI',
    '  2.1 Sejarah Perkembangan Perusahaan',
    '  2.2 Visi, Misi, dan Tujuan',
    '  2.3 Struktur Organisasi',
    '  2.4 Job Deskripsi',
    'BAB III METODE PELAKSANAAN PKL',
    '  3.1 Tugas Umum',
    '  3.2 Tugas Khusus',
    '  3.3 Analisis Permasalahan dan Solusi',
    'BAB IV HASIL YANG DICAPAI',
    '  4.1 Gambaran Umum Sistem',
    '  4.2 Analisis dan Perancangan Sistem',
    '  4.3 Implementasi Sistem',
    '  4.4 Pengujian Sistem',
    'BAB V PENUTUP',
    '  5.1 Kesimpulan',
    '  5.2 Saran',
    'DAFTAR PUSTAKA',
    'LAMPIRAN',
]
for item in toc:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    r = p.add_run(item)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
doc.add_page_break()

# ============================================================
# BAB I
# ============================================================
center_para('BAB I\nPENDAHULUAN', True, 14)
doc.add_paragraph()

heading('1.1 Latar Belakang', 2)
para('Perkembangan teknologi kecerdasan buatan (Artificial Intelligence/AI) dalam beberapa tahun terakhir telah mengalami kemajuan yang sangat pesat, khususnya dalam bidang pemrosesan bahasa alami (Natural Language Processing/NLP). Salah satu bentuk penerapan AI yang paling banyak digunakan dalam dunia bisnis adalah chatbot, yaitu program komputer yang mampu melakukan percakapan dengan pengguna secara otomatis menggunakan bahasa manusia [1].')
para('Dalam konteks layanan pelanggan (customer service), chatbot berbasis AI memiliki potensi besar untuk meningkatkan efisiensi dan kualitas layanan. Perusahaan dapat melayani pelanggan secara 24 jam tanpa henti, mengurangi waktu tunggu respons, dan menangani banyak permintaan secara bersamaan. Namun, chatbot konvensional memiliki keterbatasan dalam memahami konteks dan memberikan jawaban yang relevan terhadap pertanyaan spesifik [2].')
para('Untuk mengatasi keterbatasan tersebut, diperlukan pendekatan Retrieval-Augmented Generation (RAG), yaitu teknik yang menggabungkan kemampuan Large Language Model (LLM) dengan mekanisme pencarian dari basis pengetahuan perusahaan [3].')
para('Selain aspek layanan pelanggan, aspek Customer Relationship Management (CRM) juga memegang peranan penting. Integrasi antara chatbot AI dengan sistem CRM memungkinkan setiap interaksi pelanggan tercatat dan terkelola secara otomatis [4].')
para('Berdasarkan latar belakang di atas, penulis mengembangkan Mimotes AI, yaitu sistem chatbot AI berbasis pengetahuan dengan RAG dan pipeline CRM untuk optimalisasi layanan pelanggan di Bank Mandiri KCP Tegal Sudirman.')

heading('1.2 Tujuan PKL', 2)
heading('1.2.1 Tujuan Umum', 3)
para('Memberikan pengalaman kerja kepada mahasiswa dalam menerapkan teori perkuliahan dengan situasi nyata di bidang kecerdasan buatan, rekayasa perangkat lunak, dan basis data [5].')
heading('1.2.2 Tujuan Khusus', 3)
bullet([
    'Merancang arsitektur sistem chatbot AI berbasis RAG.',
    'Mengimplementasikan pipeline RAG: document processing, chunking, embedding, vector storage, retrieval, dan generation.',
    'Mengembangkan pipeline CRM: manajemen leads, kontak, aktivitas, dan sales pipeline.',
    'Mengintegrasikan sistem dengan WhatsApp melalui protokol Baileys.',
    'Menguji kefektifan sistem melalui pengujian black box.',
])

heading('1.3 Manfaat PKL', 2)
para('Bagi Universitas Harkat Negeri: Menambah jaringan kerja sama dengan dunia industri.')
para('Bagi Program Studi: Bahan masukan pengembangan kurikulum sesuai kebutuhan industri.')
para('Bagi Bank Mandiri KCP Tegal Sudirman: Solusi sistem informasi berbasis AI untuk optimalisasi layanan pelanggan.')
para('Bagi Penulis: Pengalaman kerja nyata dalam pengembangan sistem AI dan CRM.')
doc.add_page_break()

# ============================================================
# BAB II
# ============================================================
center_para('BAB II\nGAMBARAN UMUM INSTANSI', True, 14)
doc.add_paragraph()

heading('2.1 Sejarah Perkembangan Perusahaan', 2)
para('Bank Mandiri KCP Tegal Sudirman merupakan Kantor Cabang Pembantu dari PT Bank Mandiri (Persero) Tbk yang terletak di Jalan Sudirman, Tegal. Bank Mandiri didirikan pada 2 Oktober 1998 dari penggabungan empat bank pemerintah: Bank Bumi Daya, Bank Dagang Negara, Bank Exim, dan Bapindo. KCP Tegal Sudirman melayani produk perbankan meliputi simpanan pinjaman, layanan transaksi, pengelolaan aset, serta layanan lelang property.')

heading('2.2 Visi, Misi, dan Tujuan', 2)
para('Visi: Menjadi bank terbaik yang mengedepankan kepuasan nasabah melalui layanan perbankan digital yang inovatif dan terpercaya.')
bullet(['Memberikan layanan perbankan digital yang inovatif dan terpercaya.', 'Mengoptimalkan layanan pelanggan melalui pemanfaatan teknologi informasi.', 'Mengelola data nasabah secara terstruktur dan efisien.'])

heading('2.3 Struktur Organisasi', 2)
para('Struktur organisasi Bank Mandiri KCP Tegal Sudirman terdiri dari Kepala Cabang Pembantu, Bagian Customer Service, Bagian Teller, Bagian Administrasi, serta Bagian Teknologi Informasi.')

heading('2.4 Job Deskripsi', 2)
para('Selama PKL di Bank Mandiri KCP Tegal Sudirman, penulis ditempatkan pada Divisi Teknologi Informasi dengan job deskripsi:')
bullet([
    'Input data PPAT dan pembuatan SKPT untuk keperluan lelang property.',
    'Validasi pajak PPH di Kantor Pajak Pratama Tegal.',
    'Pelaksanaan dan pengelolaan lelang property.',
    'Pengembangan sistem chatbot AI (Mimotes) sebagai capstone project.',
    'Pemasangan banner lelang dan penyerahan risalah lelang.',
])
doc.add_page_break()

# ============================================================
# BAB III
# ============================================================
center_para('BAB III\nMETODE PELAKSANAAN PKL', True, 14)
doc.add_paragraph()

heading('3.1 Tugas Umum', 2)
para('Selama PKL di Bank Mandiri KCP Tegal Sudirman, penulis menjalankan tugas umum meliputi input data PPAT, pembuatan SKPT, validasi pajak, pelaksanaan lelang, serta pengembangan aplikasi chatbot AI secara paralel melalui WFH dan WFA.')

heading('3.2 Tugas Khusus', 2)
heading('3.2.1 Analisis Kebutuhan Sistem', 3)
para('Penulis mengidentifikasi permasalahan utama: volume interaksi nasabah yang tinggi, pengetahuan tersebar dalam berbagai dokumen, tidak adanya integrasi CRM, serta keterbatasan akses multi-platform.')
heading('3.2.2 Perancangan Arsitektur Sistem', 3)
para('Arsitektur Mimotes AI terdiri dari: Frontend Layer (Next.js 16 + React 19 + Tailwind CSS), API Layer (108 routes), Data Layer (PostgreSQL 16 + pgvector), RAG Pipeline, dan Microservices (Baileys WhatsApp, PaddleOCR, n8n).')

heading('3.3 Analisis Permasalahan dan Solusi', 2)
para('Akurasi Respons Chatbot: Solusi RAG menggabungkan LLM dengan pencarian semantik dari knowledge base perusahaan.')
para('Pengelolaan Pelanggan: Pipeline CRM terintegrasi dengan pendekatan conversation-centric CRM.')
para('Isolasi Data Multi-Tenant: Workspace-based isolation dengan PostgreSQL Row Level Security (RLS).')
para('Ekstraksi Teks dari Gambar: Integrasi PaddleOCR dan Gemini Vision untuk OCR dan captioning.')
doc.add_page_break()

# ============================================================
# BAB IV
# ============================================================
center_para('BAB IV\nHASIL YANG DICAPAI', True, 14)
doc.add_paragraph()

heading('4.1 Gambaran Umum Sistem', 2)
para('Mimotes AI di-deploy menggunakan Docker Compose dengan lima service: PostgreSQL (5432), Next.js App (3100), PaddleOCR (8090), Baileys WhatsApp (3002), dan Database Migration.')

heading('4.2 Analisis dan Perancangan Sistem', 2)

heading('4.2.1 Use Case Diagram', 3)
para('Sistem memiliki empat aktor utama: Admin/Workspace Owner, Editor, Viewer, dan Visitor/Pelanggan.')
add_image('usecase.png', 'Gambar 4.2 Use Case Diagram Sistem Mimotes AI')

heading('4.2.2 Activity Diagram Upload Dokumen', 3)
para('Proses upload dokumen: autentikasi, validasi file, simpan ke disk, buat record di DB, masukkan ke processing queue, lalu pipeline parsing, chunking, embedding, dan storage.')
add_image('activity-upload.png', 'Gambar 4.3 Activity Diagram Upload Dokumen')

heading('4.2.3 Activity Diagram Proses Chat RAG', 3)
para('Proses chat RAG: user mengirim pesan, generate query embedding, hybrid search, confidence classification, build context, konstruksi system prompt, LLM generate streaming response, simpan ke DB.')
add_image('activity-chat.png', 'Gambar 4.4 Activity Diagram Proses Chat RAG')

heading('4.2.4 Entity Relationship Diagram', 3)
para('Sistem menggunakan 28 model database dalam enam domain: Tenant & Identity, Document & RAG Pipeline, Conversational AI, CRM & Lead Pipeline, Subscription & Billing, Configuration & Observability.')
add_image('erd.png', 'Gambar 4.5 Entity Relationship Diagram')

heading('4.2.5 Arsitektur Sistem', 3)
para('Arsitektur terdiri dari enam layer: Frontend, API Layer, RAG Pipeline, Data Layer, Microservices, dan External Services.')
add_image('architecture.png', 'Gambar 4.6 Arsitektur Sistem Mimotes AI')

heading('4.2.6 Arsitektur RAG Pipeline', 3)
para('Pipeline RAG: Document Ingestion, Embedding Generation, Vector Storage (pgvector), Retrieval (Hybrid Search + RRF), dan Response Generation (LLM).')
add_image('rag-pipeline.png', 'Gambar 4.7 Arsitektur RAG Pipeline')

heading('4.2.7 Arsitektur CRM Pipeline', 3)
para('Pipeline CRM: Lead Capture, AI Analysis, Lead Scoring, Lead Status Tracking, Follow-up Automation.')
add_image('crm-pipeline.png', 'Gambar 4.8 Arsitektur CRM Pipeline')

heading('4.3 Implementasi Sistem', 2)

heading('4.3.1 Spesifikasi Perangkat Keras dan Lunak', 3)
table = doc.add_table(rows=9, cols=2)
table.style = 'Table Grid'
specs = [('Framework', 'Next.js 16.2.7'), ('Frontend', 'React 19, Tailwind CSS, Framer Motion'), ('Database', 'PostgreSQL 16 + pgvector'), ('ORM', 'Prisma 6.19.3'), ('AI Provider', 'Mimo Pro, OpenAI, Google Gemini'), ('OCR', 'PaddleOCR + Gemini Vision'), ('WhatsApp', 'Baileys 6.7.23'), ('Container', 'Docker Compose'), ('Testing', 'Vitest')]
for i, (k, v) in enumerate(specs):
    table.rows[i].cells[0].text = k; table.rows[i].cells[1].text = v
    for c in range(2):
        for p in table.rows[i].cells[c].paragraphs:
            for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

heading('4.3.2 Implementasi RAG Pipeline', 3)
para('Document Parsing: PDF (pdf-parse), DOCX (mammoth), TXT, CSV, XLSX, gambar (PaddleOCR + Gemini Vision), URL (cheerio).')
para('Chunking: Recursive paragraph-then-sentence, 500 chars, 50 word overlap, max 1000 chunks/doc.')
para('Embedding: OpenAI text-embedding-3-small (1536 dim) atau Feature Hashing lokal (fallback gratis).')
para('Vector Storage: pgvector PostgreSQL, vector(1536), batch insert 50/txn, RLS via set_config().')
para('Retrieval: Hybrid Search (Vector 0.6 + BM25 0.4 dengan RRF). Confidence: High >= 0.55, Medium >= 0.40, Low >= 0.30, Refuse < 0.30.')
para('Generation: LLM OpenAI-compatible API, temperature 0.3, max_tokens 1000, streaming. Tiga mode: Knowledge Base, Customer Service, Sales Agent.')

heading('4.3.3 Implementasi Frontend', 3)
para('53 halaman: dashboard, chat, knowledge management, CRM, analytics, settings, widget preview, admin users, WhatsApp.')

heading('4.3.4 Implementasi Backend', 3)
para('108 API routes: Auth, Admin, AI, Analytics, Chat, Documents, Knowledge/RAG, Leads/CRM, WhatsApp, Widget, Billing, Workspace.')

heading('4.4 Pengujian Sistem', 2)
para('Pengujian black box dilakukan terhadap seluruh fitur utama:')
table2 = doc.add_table(rows=10, cols=5)
table2.style = 'Table Grid'
tests = [
    ('No', 'Fitur', 'Input', 'Expected', 'Status'),
    ('1', 'Login', 'Email+pass valid', 'Redirect dashboard', 'Lulus'),
    ('2', 'Login', 'Email+pass invalid', 'Pesan error', 'Lulus'),
    ('3', 'Upload Dokumen', 'File PDF <10MB', 'processing->ready', 'Lulus'),
    ('4', 'Chat RAG', 'Pertanyaan dokumen', 'Jawaban+sumber', 'Lulus'),
    ('5', 'Chat RAG', 'Di luar konteks', 'Penolakan sopan', 'Lulus'),
    ('6', 'Lead Capture', 'Isi form', 'Lead tersimpan', 'Lulus'),
    ('7', 'WhatsApp', 'Pesan masuk', 'Auto-reply', 'Lulus'),
    ('8', 'Settings', 'Ubah provider', 'Tersimpan', 'Lulus'),
    ('9', 'RBAC', 'Viewer upload', 'Access denied', 'Lulus'),
]
for i, row_data in enumerate(tests):
    for j, val in enumerate(row_data):
        table2.rows[i].cells[j].text = val
        for p in table2.rows[i].cells[j].paragraphs:
            for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)
doc.add_page_break()

# ============================================================
# BAB V
# ============================================================
center_para('BAB V\nPENUTUP', True, 14)
doc.add_paragraph()
heading('5.1 Kesimpulan', 2)
bullet([
    'Sistem Mimotes AI berhasil dirancang dan diimplementasikan sebagai chatbot AI berbasis RAG dengan pipeline CRM.',
    'Pipeline RAG berfungsi: document processing, chunking, embedding, vector storage, retrieval, hingga generasi respons.',
    'Pipeline CRM terintegrasi: setiap interaksi pelanggan tercatat dan terkelola otomatis.',
    'Integrasi WhatsApp memungkinkan pelanggan berinteraksi melalui platform pesan instan populer.',
    'Pengujian black box menunjukkan seluruh fitur berfungsi sesuai kebutuhan.',
])
heading('5.2 Saran', 2)
para('Bagi Bank Mandiri: Sistem perlu dipelihara dan dievaluasi berkala. Pertimbangkan integrasi dengan email dan media sosial.')
para('Bagi Prodi: Kurikulum dapat diperkaya dengan materi NLP dan RAG. PKL Capstone Project menjadi model pembelajaran efektif.')
para('Bagi Penulis: Dapat dikembangkan sebagai topik Skripsi dengan evaluasi kualitas RAG menggunakan ROUGE/BLEU.')
doc.add_page_break()

# ============================================================
# DAFTAR PUSTAKA
# ============================================================
center_para('DAFTAR PUSTAKA', True, 14)
doc.add_paragraph()
refs = [
    '[1] H. Susanto et al., "Pemanfaatan Chatbot Berbasis AI untuk Optimalisasi Layanan Pelanggan," Jurnal Sistem Informasi, vol. 19, no. 2, pp. 45-62, 2023.',
    '[2] T. Brown et al., "Language Models are Few-Shot Learners," NeurIPS, vol. 33, pp. 1877-1901, 2020.',
    '[3] P. Lewis et al., "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks," NeurIPS, vol. 33, pp. 9459-9474, 2020.',
    '[4] F. Buttle dan S. Maklan, Customer Relationship Management, 4th ed. London: Routledge, 2019.',
    '[5] Universitas Harkat Negeri, Buku Panduan PKL Prodi S1 Teknik Informatika. Tegal, 2026.',
    '[6] V. Lopez et al., "A Survey on Transfer Learning in NLP," AI Review, vol. 53, pp. 2339-2367, 2020.',
    '[7] Y. Gao et al., "RAG for LLMs: A Survey," arXiv:2312.10997, 2024.',
    '[8] A. Vaswani et al., "Attention Is All You Need," NeurIPS, vol. 30, pp. 5998-6008, 2017.',
    '[9] J. Devlin et al., "BERT," NAACL-HLT, pp. 4171-4186, 2019.',
    '[10] Qdrant, "Vector Database for AI," 2024. https://qdrant.tech/',
    '[11] Next.js, "The React Framework," 2024. https://nextjs.org/',
    '[12] Prisma, "Next-gen ORM," 2024. https://www.prisma.io/',
    '[13] Baileys, "WhatsApp Web API," 2024. https://github.com/WhiskeySockets/Baileys',
    '[14] pgvector, "Vector similarity search for Postgres," 2024. https://github.com/pgvector/pgvector',
    '[15] OpenAI, "Embeddings API," 2024. https://platform.openai.com/docs/guides/embeddings',
]
for ref in refs:
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 2.0
    r = p.add_run(ref); r.font.name = 'Times New Roman'; r.font.size = Pt(12)

# ============================================================
# LAMPIRAN
# ============================================================
doc.add_page_break()
center_para('LAMPIRAN', True, 14)
doc.add_paragraph()
heading('Lampiran A: Logbook Kegiatan Harian', 2)
log_table = doc.add_table(rows=25, cols=5)
log_table.style = 'Table Grid'
log_data = [
    ('No', 'Tanggal', 'Jam', 'Aktivitas', 'Lokasi'),
    ('1', '25/05/2026', '09:10-15:30', 'Input Akun Lelang, Validasi Pajak', 'Bank Mandiri'),
    ('2', '26/05/2026', '09:10-15:02', 'Pelaksanaan Lelang', 'Bank Mandiri'),
    ('3', '27/05/2026', '09:11-15:30', 'WFH Pembuatan Aplikasi', 'WFH'),
    ('4', '28/05/2026', '09:10-15:30', 'Melanjutkan Pembuatan Aplikasi', 'WFH'),
    ('5', '29/05/2026', '09:00-14:11', 'Input Pendataan SSPD BPHTB', 'Bank Mandiri'),
    ('6', '01/06/2026', '09:38-22:17', 'WFH Developing Aplikasi', 'WFH'),
    ('7', '02/06/2026', '09:10-14:58', 'Penyerahan Risalah Lelang', 'Bank Mandiri'),
    ('8', '03/06/2026', '09:10-17:28', 'Pembuatan Surat Pendaftaran Tanah', 'WFA'),
    ('9', '04/06/2026', '09:10-16:28', 'Input PPAT ke Sistem INTAN', 'Bank Mandiri'),
    ('10', '05/06/2026', '08:48-15:39', 'Input Data PPAT SKPT', 'Bank Mandiri'),
    ('11', '08/06/2026', '08:30-16:13', 'Input PPATK + Banner Lelang', 'Bank Mandiri'),
    ('12', '09/06/2026', '08:46-16:41', 'Pembuatan SKPT + Banner Lelang', 'Bank Mandiri'),
    ('13', '10/06/2026', '09:22-16:20', 'Input Data PPAT untuk SKPT', 'Bank Mandiri'),
    ('14', '11/06/2026', '08:30-17:40', 'Input PPAT + SKPT + Banner', 'Bank Mandiri'),
    ('15', '12/06/2026', '08:45-15:58', 'Input Data PPAT untuk SKPT', 'Bank Mandiri'),
    ('16', '15/06/2026', '08:36-16:32', 'Input Data PPAT + SKPT', 'Bank Mandiri'),
    ('17', '16/06/2026', '09:20-16:55', 'WFA Pembuatan Aplikasi', 'WFA'),
    ('18', '17/06/2026', '08:54-16:34', 'Input PPAT Membuat SKPT', 'Bank Mandiri'),
    ('19', '18/06/2026', '08:38-16:30', 'Input Data PPAT untuk SKPT', 'WFA'),
    ('20', '19/06/2026', '08:45-15:55', 'Input Data PPAT untuk SKPT', 'Bank Mandiri'),
    ('21', '22/06/2026', '08:47-16:50', 'Input PPAT', 'Bank Mandiri'),
    ('22', '23/06/2026', '08:35-15:02', 'SKPT + Kunjungi Lokasi Lelang', 'Bank Mandiri'),
    ('23', '24/06/2026', '10:52-16:21', 'WFH Pembuatan Aplikasi', 'WFH'),
    ('24', '25/06/2026', '08:35-15:56', 'Input Data PPAT', 'Bank Mandiri'),
]
for i, row_data in enumerate(log_data):
    for j, val in enumerate(row_data):
        log_table.rows[i].cells[j].text = val
        for p in log_table.rows[i].cells[j].paragraphs:
            for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(9)

heading('Lampiran B: GitHub Repository', 2)
para('Repository: https://github.com/EkoSaputro14/mimotes.git', indent=False)
para('Branch: semi-final', indent=False)

# ============================================================
# SAVE
# ============================================================
output = r'C:\Users\SMANSA\mimotes\docs\pkl-report\LAPORAN_PKL_Eko_Saputro_23215050.docx'
doc.save(output)
print(f'=== Saved: {output} ===')
